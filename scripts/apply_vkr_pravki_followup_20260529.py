from __future__ import annotations

import shutil
import zipfile
import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14


def set_run_font(run, *, size_pt: int = FONT_SIZE_PT) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def set_style_font(style, *, size_pt: int = FONT_SIZE_PT) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def is_heading(paragraph) -> bool:
    style = paragraph.style
    if style is None:
        return False
    style_id = getattr(style, "style_id", "") or ""
    style_name = getattr(style, "name", "") or ""
    return style_id.startswith("Heading") or style_name.startswith("Heading") or style_name.startswith("Заголовок")


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for section in doc.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        set_run_font(paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def replace_text_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if style is not None:
        new_paragraph.style = style
    run = new_paragraph.add_run(text)
    set_run_font(run)
    return new_paragraph


def apply_python_docx_edits() -> dict[str, int | str]:
    doc = Document(DOCX_PATH)
    report: dict[str, int | str] = {}

    for name in ("Heading 1", "Heading 2", "Heading 3", "Заголовок 1", "Заголовок 2", "Заголовок 3"):
        try:
            set_style_font(doc.styles[name])
            report[f"style_{name}"] = 1
        except Exception:
            pass

    heading_runs = 0
    for paragraph in iter_all_paragraphs(doc):
        if is_heading(paragraph):
            for run in paragraph.runs:
                set_run_font(run)
                heading_runs += 1
    report["heading_runs_font14"] = heading_runs

    table_runs = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)
                        table_runs += 1
    report["table_runs_font14"] = table_runs

    # Bibliography: keep current ГОСТ 34.602-2020 as source [4], and ГОСТ 34.602-89 after it.
    bibliography_start = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "Список использованных источников"),
        0,
    )
    gost_entries = [
        (idx, paragraph)
        for idx, paragraph in enumerate(doc.paragraphs[bibliography_start + 1 :], start=bibliography_start + 1)
        if paragraph.text.strip().startswith("ГОСТ 34.602")
    ]
    if len(gost_entries) >= 2:
        first_idx, first = gost_entries[0]
        second_idx, second = gost_entries[1]
        txt_first = first.text
        txt_second = second.text
        if "2020" not in txt_first and "2020" in txt_second:
            replace_paragraph_text(first, txt_second)
            replace_paragraph_text(second, txt_first)
            report["gost_entries_swapped"] = 1
        else:
            report["gost_entries_swapped"] = 0

    # Table 14: replace chapter labels with engineering deliverables.
    table14_rows_changed = 0
    deliverables = {
        "Обследование": "Отчет предпроектного обследования, перечень контролируемых узлов и исходных данных",
        "Концепция": "Концепция ПАК, схема потоков данных, состав функций",
        "Техническое задание": "Техническое задание на прототип ПАК PdM РП",
        "Техническое проектирование": "Технический проект архитектуры, данных, алгоритмов и интерфейсов",
        "Рабочее проектирование": "Рабочий прототип, программные модули, цифровая сцена и структура данных",
        "Апробация": "Протокол апробации, метрики RUL, расчет надежности и экономической эффективности",
        "Документирование": "Итоговый комплект пояснительной записки, приложений и материалов к защите",
    }
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.rows[0].cells) < 3:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        txt = table_text(table)
        if header[:3] == ["Этап", "Содержание работ", "Результат"] and "Глава 1" in txt:
            for row in table.rows[1:]:
                if len(row.cells) >= 3:
                    key = row.cells[0].text.strip()
                    if key in deliverables:
                        set_cell_text(row.cells[2], deliverables[key])
                        table14_rows_changed += 1
            break
    report["table14_rows_changed"] = table14_rows_changed

    # Duplicate strategy table in chapter 2: add predictive maintenance if absent.
    pdm_rows_added = 0
    for table in doc.tables:
        txt = table_text(table)
        if "Реактивная" in txt and "ППР" in txt and "По состоянию" in txt and "PdM" not in txt:
            row = table.add_row()
            values = [
                "Предиктивное обслуживание (PdM)",
                "Прогнозирование остаточного ресурса по телеметрии и диагностическим признакам.",
                "Требует цифрового контура данных и настройки модели, но позволяет заранее планировать ТО.",
            ]
            for cell, value in zip(row.cells, values):
                set_cell_text(cell, value)
            pdm_rows_added += 1
    report["pdm_rows_added"] = pdm_rows_added

    # Technical stack table: align ML row with actual implementation.
    ml_rows_changed = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2 and row.cells[0].text.strip() == "ML":
                set_cell_text(row.cells[1], "scikit-learn; XGBoost (резервный вариант для сравнения)")
                ml_rows_changed += 1
    report["ml_rows_changed"] = ml_rows_changed

    # Requirements table: add factual collector frequency requirement.
    freq_req_added = 0
    for table in doc.tables:
        txt = table_text(table)
        if "Группа требований" in txt and "Сбор данных" in txt and "Визуализация" in txt:
            if "не ниже 10 Гц" not in txt:
                row = table.add_row()
                set_cell_text(row.cells[0], "Частота записи")
                set_cell_text(row.cells[1], "Фактическая средняя частота внешнего коллектора должна быть не ниже 10 Гц.")
                freq_req_added += 1
            break
    report["freq_req_added"] = freq_req_added

    # Dashboard table: remove unsupported confidence interval wording.
    dashboard_rows_changed = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 3 and row.cells[0].text.strip() == "Прогноз":
                set_cell_text(row.cells[1], "Прогнозное значение RUL")
                dashboard_rows_changed += 1
    report["dashboard_rows_changed"] = dashboard_rows_changed

    paragraph_replacements = {
        "scikit-learn и XGBoost - для реализации базовых и ансамблевых моделей прогнозирования": (
            "scikit-learn - для реализованной модели MLPRegressor и базовых алгоритмов сравнения; "
            "XGBoost рассматривается как резервный вариант для дальнейшего сопоставления моделей прогнозирования"
        ),
        "ГОСТ 34.602–89": "ГОСТ 34.602–2020",
        "ГОСТ 34.602-89": "ГОСТ 34.602-2020",
    }
    paragraph_hits = 0
    for paragraph in iter_all_paragraphs(doc):
        # Keep bibliography entries handled separately; do not globally mutate them.
        if paragraph.text.strip().startswith("ГОСТ 34.602"):
            continue
        if "Источник" in paragraph.text and "ГОСТ 34.602" in paragraph.text:
            continue
        for old, new in paragraph_replacements.items():
            if old in paragraph.text:
                replace_text_in_paragraph(paragraph, old, new)
                paragraph_hits += 1
        # User decision: in-text references should not include page numbers.
        cleaned = re.sub(r"\[(\d+),\s*с\.\s*\d+\]", r"[\1]", paragraph.text)
        cleaned = re.sub(r"\[(\d+),\s*c\.\s*\d+\]", r"[\1]", cleaned)
        if cleaned != paragraph.text:
            replace_paragraph_text(paragraph, cleaned)
            paragraph_hits += 1
        if "ГОСТ 34.602–2020 [5]" in paragraph.text:
            replace_text_in_paragraph(paragraph, "ГОСТ 34.602–2020 [5]", "ГОСТ 34.602–2020 [4]")
            paragraph_hits += 1
    report["paragraph_replacements"] = paragraph_hits

    container_rows_changed = 0
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        if paragraph.text.strip().startswith("В проектируемой системе контейнеризация используется"):
            replace_paragraph_text(
                paragraph,
                "В реализованном прототипе контейнерно развернуты инфраструктурные сервисы хранения и "
                "визуализации, а сбор и аналитическая обработка выполняются Python-скриптами файлового "
                "конвейера. Контейнеризация collector и ml-service рассматривается как проектная структура "
                "дальнейшего развития.",
            )
            container_rows_changed += 1
        elif paragraph.text.strip() == "Проектный состав контейнеров:":
            replace_paragraph_text(paragraph, "Проектный состав контейнеров при дальнейшем развитии:")
            container_rows_changed += 1
        elif paragraph.text.strip().startswith("collector - Python-сервис"):
            replace_paragraph_text(
                paragraph,
                "collector - перспективный Python-сервис сбора телеметрии из CoppeliaSim или промышленного контроллера;",
            )
            container_rows_changed += 1
        elif paragraph.text.strip().startswith("ml-service - расчет"):
            replace_paragraph_text(
                paragraph,
                "ml-service - перспективный сервис расчета признаков, HI, RUL и рекомендаций по ТО;",
            )
            container_rows_changed += 1
        elif paragraph.text.strip().startswith("notebooks - исследовательская"):
            replace_paragraph_text(
                paragraph,
                "notebooks - вспомогательная исследовательская среда для анализа данных и подбора моделей;",
            )
            container_rows_changed += 1
        elif "CoppeliaSim -> collector -> influxdb -> ml-service" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                "Для текущей апробации фактическая цепочка имеет вид CoppeliaSim -> Python-коллектор -> "
                "CSV/JSONL -> расчет признаков/HI/RUL -> InfluxDB/Grafana. При промышленном развитии collector "
                "и ml-service могут быть вынесены в отдельные контейнеры.",
            )
            container_rows_changed += 1
        elif paragraph.text.strip() == "CoppeliaSim -> collector -> influxdb -> ml-service -> influxdb -> grafana.":
            replace_paragraph_text(paragraph, "CoppeliaSim -> Python pipeline -> InfluxDB -> Grafana.")
            container_rows_changed += 1
    report["container_paragraphs_changed"] = container_rows_changed

    # Chapter 5: clarify the simulated update rate versus factual collector rate.
    freq_note_added = 0
    for paragraph in doc.paragraphs:
        if (
            ("Фактическая средняя частота записи" in paragraph.text or "Средний шаг записи" in paragraph.text)
            and "10,77 Гц" in paragraph.text
            and not any("10,77 Гц является фактической средней частотой" in p.text for p in doc.paragraphs)
        ):
            insert_paragraph_after(
                paragraph,
                "Отличие от расчетных 25 Гц связано с тем, что 25 Гц относится к обновлению графиков "
                "внутри Lua-сценария CoppeliaSim, а 10,77 Гц является фактической средней частотой записи "
                "пакетов внешним Python-коллектором с учетом обмена через Remote API.",
                style=paragraph.style,
            )
            freq_note_added += 1
            break
    report["freq_note_added"] = freq_note_added

    # Chapter 6: add the detailed damage/RUL calculation before the conclusion if it is still only in conclusion.
    damage_calc_added = 0
    paragraphs_now = list(doc.paragraphs)
    conclusion_idx = next((idx for idx, p in enumerate(paragraphs_now) if p.text.strip() == "Заключение"), len(paragraphs_now))
    has_damage_before_conclusion = any(
        "эффективная скорость накопления повреждения составляет 5,23" in p.text
        for p in paragraphs_now[:conclusion_idx]
    )
    if not has_damage_before_conclusion:
        for paragraph in doc.paragraphs:
            if ("Средний шаг записи" in paragraph.text or "Фактическая средняя частота записи" in paragraph.text) and "10,77 Гц" in paragraph.text:
                insert_paragraph_after(
                    paragraph,
                    "Для иллюстрации связи коэффициента загрузки и деградационного индекса использован расчетный "
                    "сценарий: при коэффициентах нагруженности из НИРС эффективная скорость накопления повреждения "
                    "составляет 5,23 x 10^-6 1/цикл; после 10000 циклов D = 0,052, HI = 0,948; при текущем "
                    "повреждении D = 0,40 расчетный остаточный ресурс составляет около 114700 циклов. Этот расчет "
                    "носит иллюстративный характер и подлежит уточнению по реальной статистике отказов.",
                    style=paragraph.style,
                )
                damage_calc_added += 1
                break
    report["damage_calc_added"] = damage_calc_added

    conclusion_trimmed = 0
    paragraphs_for_trim = list(doc.paragraphs)
    conclusion_idx_for_trim = next(
        (idx for idx, p in enumerate(paragraphs_for_trim) if p.text.strip() == "Заключение"),
        len(paragraphs_for_trim),
    )
    for paragraph in paragraphs_for_trim[conclusion_idx_for_trim + 1 :]:
        if "114700 циклов" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                "Экономическая оценка для расчетного сценария показала годовой эффект 450000 руб. "
                "и срок окупаемости 1,0 год.",
            )
            conclusion_trimmed += 1
    report["conclusion_trimmed"] = conclusion_trimmed

    reliability_note_added = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if "Для прототипа принимаются целевые значения" in paragraph.text and not any(
            "количественная оценка надежности ограничена работоспособностью диагностического контура" in p.text
            for p in doc.paragraphs
        ):
            insert_paragraph_after(
                paragraph,
                "В рамках ВКР количественная оценка надежности ограничена работоспособностью диагностического "
                "контура, поскольку отсутствует статистика реальных отказов и восстановлений робота ABB IRB 660 "
                "на рассматриваемом участке. Для промышленного внедрения показатели безотказности и готовности "
                "должны уточняться по журналам ТОиР, аварийным простоям и данным контроллера.",
                style=paragraph.style,
            )
            reliability_note_added += 1
            break
    report["reliability_note_added"] = reliability_note_added

    economics_note_added = 0
    existing_economics_notes = [
        p for p in doc.paragraphs if "3 события/год принято как расчетное допущение" in p.text
    ]
    for p in existing_economics_notes:
        # Move the note out of the chapter conclusion if it was inserted there by an earlier pass.
        delete_paragraph(p)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Экономический эффект связан"):
            insert_paragraph_after(
                paragraph,
                "Число предотвращаемых событий 3 события/год принято как расчетное допущение для оценки "
                "экономического эффекта. При внедрении системы параметр должен уточняться по фактической "
                "статистике простоев предприятия и истории ремонтных воздействий.",
                style=paragraph.style,
            )
            economics_note_added += 1
            break
    report["economics_note_added"] = economics_note_added

    # Add a small clarification if the object/pallet count paragraph still needs it.
    n_object_note_added = 0
    if not any("Nобъект = 16 относится к переносимым объектам" in p.text for p in doc.paragraphs):
        for paragraph in doc.paragraphs:
            if "дополнительная единица соответствует рабочей копии паллеты" in paragraph.text:
                insert_paragraph_after(
                    paragraph,
                    "Следовательно, Nобъект = 16 относится к переносимым объектам, а Nсозд = 17 учитывает также "
                    "рабочую копию паллеты.",
                    style=paragraph.style,
                )
                n_object_note_added += 1
                break
    report["n_object_note_added"] = n_object_note_added

    doc.save(DOCX_PATH)
    return report


def replace_text_nodes(paragraph, old: str, new: str, ns: dict[str, str]) -> bool:
    nodes = paragraph.xpath(".//w:t | .//m:t", namespaces=ns)
    if not nodes:
        return False
    full = "".join(node.text or "" for node in nodes)
    if old not in full:
        return False
    updated = full.replace(old, new)
    nodes[0].text = updated
    for node in nodes[1:]:
        node.text = ""
    return True


def patch_ooxml() -> dict[str, int]:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    tmp = DOCX_PATH.with_suffix(".tmp.docx")
    report = {
        "quote_fixes": 0,
        "abstract_fixes": 0,
        "formula48": 0,
        "formula104": 0,
        "formula105": 0,
        "formula106": 0,
        "reliability_phrase": 0,
    }
    with zipfile.ZipFile(DOCX_PATH, "r") as zin:
        document_xml = zin.read("word/document.xml")
        root = etree.fromstring(document_xml)
        for paragraph in root.xpath(".//w:p", namespaces=ns):
            text = "".join(node.text or "" for node in paragraph.xpath(".//w:t | .//m:t", namespaces=ns))
            if not text:
                continue
            for old in (
                "ООО “Компания “Здоровая жизнь””.",
                "ООО “Компания “Здоровая жизнь””",
                "ООО «Компания “Здоровая жизнь”».",
            ):
                if old in text:
                    if replace_text_nodes(paragraph, old, "ООО «Компания “Здоровая жизнь”»", ns):
                        report["quote_fixes"] += 1
                        text = text.replace(old, "ООО «Компания “Здоровая жизнь”»")
            if "формируется в объеме основную часть" in text:
                old = "формируется в объеме основную часть, список литературы и приложения формата А4"
                new = "включает основную часть, список использованных источников и приложения, выполненные на листах формата А4"
                if replace_text_nodes(paragraph, old, new, ns):
                    report["abstract_fixes"] += 1
                    text = text.replace(old, new)
            if "Оценку надёжности, как технической системы" in text:
                old = "Оценку надёжности, как технической системы и анализ влияния"
                new = "Оценку работоспособности диагностического контура и расчетный анализ влияния"
                if replace_text_nodes(paragraph, old, new, ns):
                    report["reliability_phrase"] += 1
                    text = text.replace(old, new)
            if "Vчас" in text and "10242" in text and "(48)" in text:
                if replace_text_nodes(paragraph, "10242", "1024²", ns):
                    report["formula48"] += 1
            if "Tнабл" in text and "23,6" in text and "(104)" in text:
                if replace_text_nodes(paragraph, "23,6", "2059,05", ns):
                    report["formula104"] += 1
            if "Δtср" in text and "(105)" in text:
                changed = False
                for old, new in (("23,6", "2059,05"), ("472", "22173"), ("0,05", "0,0929")):
                    if replace_text_nodes(paragraph, old, new, ns):
                        changed = True
                if changed:
                    report["formula105"] += 1
            if "fнабл" in text and "(106)" in text:
                if replace_text_nodes(paragraph, "20", "10,77", ns):
                    report["formula106"] += 1

        new_document_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = new_document_xml
                zout.writestr(item, data)
    tmp.replace(DOCX_PATH)
    return report


def audit_docx() -> dict[str, int | str]:
    report: dict[str, int | str] = {}
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        report["zip_bad"] = str(zf.testzip())
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    checks = {
        "bad_quotes_left": "ООО “Компания “Здоровая жизнь””" in xml,
        "bad_abstract_left": "формируется в объеме основную часть" in xml,
        "formula48_bad_left": "10242" in xml,
        "formula105_old_left": "23,6 / 472" in xml,
        "formula106_old_left": "20 Гц" in xml and "fнабл" in xml,
    }
    for key, value in checks.items():
        report[key] = int(value)

    doc = Document(DOCX_PATH)
    table14_chapter_results_left = 0
    for table in doc.tables:
        if len(table.rows) and [cell.text.strip() for cell in table.rows[0].cells][:3] == [
            "Этап",
            "Содержание работ",
            "Результат",
        ]:
            table14_chapter_results_left = int(
                any("Глава" in row.cells[2].text for row in table.rows[1:] if len(row.cells) >= 3)
            )
            break
    report["table14_chapter_results_left"] = table14_chapter_results_left

    heading_count = 0
    heading_bad = 0
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            heading_count += 1
            for run in paragraph.runs:
                if run.text.strip() and run.font.size is not None and round(run.font.size.pt, 1) != FONT_SIZE_PT:
                    heading_bad += 1
    report["heading_count"] = heading_count
    report["heading_bad_direct_sizes"] = heading_bad

    table_runs = 0
    table_bad = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():
                            table_runs += 1
                            if run.font.size is not None and round(run.font.size.pt, 1) != FONT_SIZE_PT:
                                table_bad += 1
    report["table_text_runs"] = table_runs
    report["table_bad_direct_sizes"] = table_bad
    report["tables"] = len(doc.tables)
    report["paragraphs"] = len(doc.paragraphs)
    return report


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}.backup_before_pravki_followup_{datetime.now():%Y%m%d_%H%M%S}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)
    py_report = apply_python_docx_edits()
    xml_report = patch_ooxml()
    audit = audit_docx()
    print(f"backup={backup}")
    print("python_docx_report=" + repr(py_report))
    print("ooxml_report=" + repr(xml_report))
    print("audit=" + repr(audit))


if __name__ == "__main__":
    main()
