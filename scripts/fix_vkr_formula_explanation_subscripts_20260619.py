from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = PROJECT_ROOT / "вкр"
PREFIX_GDE = "\u0433\u0434\u0435"


UNDERSCORE_TOKEN_RE = re.compile(
    r"(?<![\w/])"
    r"(?P<base>[A-Za-zА-Яа-яЁёΑ-Ωα-ωµμ][A-Za-zА-Яа-яЁёΑ-Ωα-ωµμ0-9\u0302]*)"
    r"_"
    r"(?P<idx>\{[^}]+\}|[A-Za-zА-Яа-яЁёΑ-Ωα-ωµμ0-9]+(?:[,-][A-Za-zА-Яа-яЁёΑ-Ωα-ωµμ0-9]+)*)"
)


COMPACT_SUBSCRIPTS = {
    "Mmax,i": ("M", "max,i"),
    "Mдоп,i": ("M", "доп,i"),
    "Kз,i": ("K", "з,i"),
    "xкр,j": ("x", "кр,j"),
    "Iпризнаки": ("I", "признаки"),
    "Iсцена": ("I", "сцена"),
    "Iтелем": ("I", "телем"),
    "Sготов": ("S", "готов"),
    "Sапр": ("S", "апр"),
    "Kготовн": ("K", "готовн"),
    "Kпред": ("K", "пред"),
    "Kдан": ("K", "дан"),
    "Kфаз": ("K", "фаз"),
    "Mmax": ("M", "max"),
    "Mср": ("M", "ср"),
    "amax": ("a", "max"),
    "ωср": ("ω", "ср"),
    "HIкр": ("HI", "кр"),
    "Nкр": ("N", "кр"),
    "Nтек": ("N", "тек"),
    "Tпл": ("T", "пл"),
    "Tрез": ("T", "рез"),
    "AТО": ("A", "ТО"),
    "Iцикл": ("I", "цикл"),
    "Iтел": ("I", "тел"),
    "Iсц": ("I", "сц"),
    "Iц": ("I", "ц"),
    "Iпр": ("I", "пр"),
}

COMPACT_TOKEN_RE = re.compile(
    r"(?<![\w])("
    + "|".join(re.escape(token) for token in sorted(COMPACT_SUBSCRIPTS, key=len, reverse=True))
    + r")(?![\w])"
)


def find_working_docx() -> Path:
    candidates = [
        p
        for p in DOCX_DIR.glob("*.docx")
        if "2026" in p.name and "backup" not in p.name.lower() and not p.name.startswith("~$")
    ]
    if not candidates:
        raise FileNotFoundError("Working VKR DOCX was not found in the vkr folder")
    return candidates[0]


def set_run_font(run, *, subscript: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.subscript = subscript
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def strip_index_braces(index: str) -> str:
    if index.startswith("{") and index.endswith("}"):
        return index[1:-1]
    return index


def next_match(text: str, pos: int):
    matches = []
    m1 = UNDERSCORE_TOKEN_RE.search(text, pos)
    if m1:
        matches.append(("underscore", m1))
    m2 = COMPACT_TOKEN_RE.search(text, pos)
    if m2:
        matches.append(("compact", m2))
    if not matches:
        return None
    return min(matches, key=lambda item: item[1].start())


def add_text(paragraph, text: str, *, subscript: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, subscript=subscript)


def rebuild_with_subscripts(paragraph, text: str) -> int:
    paragraph.clear()
    pos = 0
    conversions = 0

    while True:
        found = next_match(text, pos)
        if found is None:
            add_text(paragraph, text[pos:])
            break

        kind, match = found
        add_text(paragraph, text[pos : match.start()])

        if kind == "underscore":
            base = match.group("base")
            index = strip_index_braces(match.group("idx"))
        else:
            token = match.group(1)
            base, index = COMPACT_SUBSCRIPTS[token]

        add_text(paragraph, base)
        add_text(paragraph, index, subscript=True)
        conversions += 1
        pos = match.end()

    return conversions


def main() -> None:
    docx_path = find_working_docx()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = docx_path.with_name(f"{docx_path.stem}.backup_before_formula_explanation_subscripts_{timestamp}.docx")
    shutil.copy2(docx_path, backup_path)

    doc = Document(docx_path)
    changed_paragraphs = 0
    conversions = 0
    touched = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        stripped = text.strip().lower()
        if not stripped.startswith(PREFIX_GDE):
            continue
        if "_" not in text and not COMPACT_TOKEN_RE.search(text):
            continue
        current = rebuild_with_subscripts(paragraph, text)
        if current:
            changed_paragraphs += 1
            conversions += current
            touched.append(idx)

    doc.save(docx_path)

    print(f"docx={docx_path}")
    print(f"backup={backup_path}")
    print(f"changed_paragraphs={changed_paragraphs}")
    print(f"subscript_conversions={conversions}")
    print("paragraph_indices=" + ",".join(str(i) for i in touched))


if __name__ == "__main__":
    main()
