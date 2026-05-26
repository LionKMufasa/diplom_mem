# -*- coding: utf-8 -*-
"""Add literature references to the current NIRS-8 DOCX.

The script deliberately edits only citation markers in existing paragraphs.
It keeps the user's final DOCX filename and writes a timestamped backup.
"""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


def find_nirs8_docx(root: Path) -> Path:
    candidates: list[tuple[float, int, Path]] = []
    for path in root.rglob("*.docx"):
        if path.name.startswith("~$"):
            continue
        try:
            stat = path.stat()
        except OSError:
            continue
        path_text = str(path.parent)
        if "2026" in path.name and "(" in path_text and "8" in path_text:
            candidates.append((stat.st_mtime, stat.st_size, path))
    if not candidates:
        raise FileNotFoundError("NIRS-8 DOCX was not found")
    return sorted(candidates, reverse=True)[0][2]


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def append_citation(paragraph, number: int) -> bool:
    marker = f"[{number}]"
    if marker in paragraph.text:
        return False

    run = None
    for candidate in reversed(paragraph.runs):
        if candidate.text.strip():
            run = candidate
            break
    if run is None:
        return False

    text = run.text
    stripped = text.rstrip()
    trailing_space = text[len(stripped) :]
    if re.search(r"[.!?]$", stripped):
        run.text = stripped[:-1].rstrip() + f" {marker}" + stripped[-1] + trailing_space
    else:
        run.text = stripped + f" {marker}" + trailing_space
    return True


def add_citation_by_prefix(doc: Document, prefix: str, number: int) -> tuple[bool, str]:
    for index, paragraph in enumerate(doc.paragraphs):
        normalized = " ".join(paragraph.text.split())
        if normalized.startswith(prefix):
            return append_citation(paragraph, number), f"{index}: {normalized[:90]}"
    return False, f"NOT FOUND: {prefix[:90]}"


def count_bibliography_entries(doc: Document) -> int:
    in_bibliography = False
    count = 0
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "Список литературы":
            in_bibliography = True
            continue
        if text == "Приложение":
            break
        if in_bibliography and text:
            count += 1
    return count


def main() -> None:
    root = Path.cwd()
    path = find_nirs8_docx(root)

    backup_dir = path.parent / "_backups"
    backup_dir.mkdir(exist_ok=True)
    backup_path = backup_dir / f"{path.stem}.backup_before_citations_{datetime.now():%Y%m%d_%H%M%S}.docx"
    shutil.copy2(path, backup_path)

    doc = Document(path)

    replacements = 0
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text.startswith("Предпроектное обследование выполняется") and "[5]" in text:
            if replace_in_runs(paragraph, "[5]", "[13]"):
                replacements += 1

    targets: list[tuple[str, int]] = [
        ("Современные автоматизированные производственные линии всё чаще строятся", 18),
        ("Традиционное планово-предупредительное обслуживание не всегда позволяет", 9),
        ("Объектом исследования является промышленный робот-паллетизатор ABB IRB 660", 12),
        ("Методическую основу работы составляют положения теории надёжности", 7),
        ("В рассматриваемом технологическом процессе робот выполняет захват", 13),
        ("Основными элементами промышленного робота, подверженными постепенной деградации", 8),
        ("На скорость деградации влияют масса переносимого груза", 7),
        ("Таким образом, даже при сравнительно малом повреждении", 14),
        ("Для построения модели целесообразно выбрать узлы", 15),
        ("Концептуальное проектирование модели начинается с анализа физических причин", 8),
        ("В редукторах деградация может выражаться в локальном повреждении", 11),
        ("На долговечность приводных узлов влияют амплитуда нагрузочного момента", 7),
        ("В рамках исследовательской модели факторы эксплуатации объединяются", 14),
        ("Модель деградации является центральным элементом предиктивного обслуживания", 9),
        ("Для оценки текущего состояния приводного узла используются диагностические признаки", 15),
        ("Данный показатель характеризует средний уровень нагруженности", 11),
        ("Пиковый момент позволяет фиксировать ударные режимы", 8),
        ("Этот показатель отражает энергетические затраты привода", 14),
        ("Увеличение отражает более резкие режимы движения", 15),
        ("Положительное значение при одинаковых условиях работы указывает", 9),
        ("Таким образом, представленные схемы показывают", 18),
        ("Для удобства интерпретации технического состояния механического узла", 5),
        ("Такое разделение позволяет использовать модель не только для фиксации", 6),
        ("Техническое задание составлено в соответствии со структурой ГОСТ 34.602-2020", 4),
        ("Объектом автоматизации является роботизированный участок паллетизации", 12),
        ("В составе объекта выделяются механические узлы", 8),
        ("Система в рамках настоящего ТЗ понимается как расчётная модель", 4),
        ("Требования к достоверности результатов состоят в том", 5),
        ("Источниками разработки являются задание на НИРС", 4),
        ("Для каждой оси и фазы цикла рассчитываются", 15),
        ("Такая модель не заменяет детальный расчёт ресурса", 16),
        ("Таким образом, при увеличении момента или ускорения", 14),
        ("При таком подходе показатель изменяется от 1 до 0", 16),
        ("Такой показатель удобен для использования в задачах предиктивного обслуживания", 22),
        ("Это означает, что при повышении момента относительно базового значения", 8),
        ("Увеличение отражает влияние более резких режимов движения", 15),
        ("Если , то среднеквадратичный момент увеличивается", 9),
        ("Критерий предельного состояния задаётся через достижение", 5),
        ("Для задач предиктивного обслуживания также вводятся предупредительный", 6),
        ("Рабочее проектирование описывает последовательность практического расчёта модели", 18),
        ("Полученное значение не является моментом на оси робота", 7),
        ("На первом этапе для выбранной оси и выбранной фазы", 15),
        ("Полученные признаки сохраняются в таблицу", 15),
        ("Для постоянного режима остаточный ресурс в циклах", 16),
        ("Полученный результат показывает, что при заданном режиме эксплуатации", 14),
        ("Цель апробации - проверить, что предложенная модель позволяет", 9),
        ("Полученное значение соответствует исправному состоянию узла", 5),
        ("Проверка критерия предельного состояния выполняется", 5),
        ("Остаточный ресурс определяется как число циклов", 16),
        ("Следовательно, остаточный ресурс уменьшается", 14),
        ("Апробация показала, что предложенная модель позволяет перейти", 9),
        ("В ходе научно-исследовательской работы разработана и исследована", 18),
        ("Исследованы основные механизмы износа приводных узлов", 8),
        ("Разработана математическая модель накопления повреждений", 16),
        ("В рабочем проектировании приведён алгоритм практического расчёта", 9),
    ]

    added: list[str] = []
    skipped: list[str] = []
    for prefix, number in targets:
        changed, info = add_citation_by_prefix(doc, prefix, number)
        if changed:
            added.append(f"[{number}] {info}")
        else:
            skipped.append(f"[{number}] {info}")

    doc.save(path)

    updated_doc = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in updated_doc.paragraphs)
    markers = re.findall(r"\[([^\]]+)\]", full_text)
    citation_numbers = [int(value) for value in markers if re.fullmatch(r"\d+", value)]
    invalid_markers = [value for value in markers if not re.fullmatch(r"\d+", value)]
    bibliography_count = count_bibliography_entries(updated_doc)

    print(f"updated={path}")
    print(f"backup={backup_path}")
    print(f"bibliography_count={bibliography_count}")
    print(f"replacements={replacements}")
    print(f"added={len(added)}")
    print(f"skipped={len(skipped)}")
    print(f"citation_mentions={len(citation_numbers)}")
    print(f"citation_numbers={sorted(set(citation_numbers))}")
    print(f"invalid_markers={invalid_markers}")
    too_high = [number for number in citation_numbers if number > bibliography_count]
    print(f"citations_over_bibliography_count={sorted(set(too_high))}")


if __name__ == "__main__":
    main()
