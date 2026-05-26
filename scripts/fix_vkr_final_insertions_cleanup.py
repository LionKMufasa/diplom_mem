from __future__ import annotations

from docx import Document

from fill_vkr_remaining_insertions import (
    FIG_DIR,
    add_picture_before,
    clear_paragraph,
    find_docx,
    remove_paragraph,
    set_run_font,
)


CONCLUSION_PARAGRAPHS = [
    "В выпускной квалификационной работе разработан и апробирован прототип программно-аппаратного контура предиктивного обслуживания промышленного робота-паллетизатора на базе цифровой модели в CoppeliaSim. Цель работы достигнута: построена логика сбора динамической телеметрии, сформированы диагностические признаки, реализованы расчет HI и прогноз остаточного ресурса, подготовлена визуализация состояния узлов и рекомендаций по техническому обслуживанию.",
    "В предпроектной части обоснована актуальность задачи для роботизированной паллетизации, где отказ привода или механического узла приводит к остановке связанной линии. По данным НИРС и расчетам ВКР один цикл включает 4 картонных листа и 12 упаковок массой 63 кг, длительность цикла принята 187 с, производительность составляет 231 упаковку/ч или 14,55 т/ч, а расчетный годовой грузопоток при двухсменной работе достигает 58212 т. Коэффициент загрузки робота по единичной упаковке равен 0,35, что не превышает паспортного ограничения выбранного паллетизирующего робота.",
    "В проектной части сформированы требования к системе по ГОСТ 34.602-89, разработана архитектура ПАК, описаны объекты цифровой сцены, состав телеметрии, структура хранения данных, алгоритмы формирования признаков и правила предупреждения. Практическая реализация включает CoppeliaSim-сцену, Lua-скрипт паллетизационного цикла, Python-конвейер обработки данных, расчетные таблицы HI/RUL, нейросетевую модель MLPRegressor и аналитическую панель Grafana.",
    "Апробация выполнена на наборе long_live_01. Получено 22174 пакета телеметрии и 88696 нормализованных строк по четырем осям; коэффициенты полноты данных и фазовой разметки составили 1,000. На основе 56 строк фазовых признаков сформировано 17920 RUL-оценок и 17920 нейросетевых прогнозов. Средняя тестовая ошибка MLPRegressor равна MAE = 1,173 цикла, RMSE = 1,442 цикла, R² = 0,994, что подтверждает пригодность выбранного подхода для учебного прототипа предиктивного обслуживания.",
    "В работе также приведены расчетные показатели деградационной модели: при коэффициентах нагруженности из НИРС эффективная скорость накопления повреждения составляет 5,23 x 10^-6 1/цикл, после 10000 циклов повреждение равно 0,052, индекс состояния HI = 0,948, а при текущем повреждении 0,40 расчетный остаточный ресурс составляет около 114700 циклов. Экономическая оценка для расчетного сценария показала годовой эффект 450000 руб. и срок окупаемости 1,0 год.",
    "Полученный результат может быть использован как основа для дальнейшей интеграции с реальным контроллером робота, промышленной базой временных рядов и эксплуатационными данными предприятия. Дальнейшее развитие работы целесообразно направить на сбор телеметрии с реального оборудования, уточнение деградационной модели по фактическим отказам, расширение набора диагностических признаков и настройку регламентов технического обслуживания по результатам RUL-прогноза.",
]


def find_exact(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Не найден абзац: {text}")


def fix_missing_rul_figure(doc: Document) -> int:
    fixed = 0
    for paragraph in list(doc.paragraphs):
        if "ВСТАВКА" in paragraph.text and "графики момента" in paragraph.text:
            add_picture_before(
                doc,
                paragraph,
                FIG_DIR / "torque_rms_by_axis.png",
                "Рисунок 11 - Среднеквадратический момент приводов по фазам набора long_live_01",
                width_cm=15.0,
            )
            remove_paragraph(paragraph)
            fixed += 1
        if "ВСТАВКА" in paragraph.text and "график фактического" in paragraph.text:
            add_picture_before(
                doc,
                paragraph,
                FIG_DIR / "rul_nn_actual_predicted_s3_motor1.png",
                "Рисунок 12 - Сравнение фактического и прогнозного RUL для сценария S3, motor1",
                width_cm=15.0,
            )
            remove_paragraph(paragraph)
            fixed += 1
    return fixed


def rewrite_conclusion(doc: Document) -> int:
    conclusion = find_exact(doc, "Заключение")
    bibliography = find_exact(doc, "Список литературы")

    removed = 0
    node = conclusion._p.getnext()
    while node is not None and node is not bibliography._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        removed += 1
        node = next_node

    for text in CONCLUSION_PARAGRAPHS:
        paragraph = doc.add_paragraph()
        bibliography._p.addprevious(paragraph._p)
        clear_paragraph(paragraph, text)
    return removed


def fix_table_formula_reference(doc: Document) -> int:
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "после перенумерации" in paragraph.text:
                        paragraph.text = "Срок окупаемости вычислен по формуле (113)"
                        for run in paragraph.runs:
                            set_run_font(run, size=9.8)
                        changed += 1
    return changed


def main() -> None:
    docx_path = find_docx()
    doc = Document(docx_path)
    fixed_figures = fix_missing_rul_figure(doc)
    removed_conclusion_nodes = rewrite_conclusion(doc)
    fixed_refs = fix_table_formula_reference(doc)
    doc.save(docx_path)
    placeholders = sum(1 for paragraph in doc.paragraphs if "ВСТАВКА" in paragraph.text)
    print(f"docx\t{docx_path}")
    print(f"fixed_figures\t{fixed_figures}")
    print(f"removed_conclusion_nodes\t{removed_conclusion_nodes}")
    print(f"fixed_refs\t{fixed_refs}")
    print(f"remaining_placeholders\t{placeholders}")


if __name__ == "__main__":
    main()
