from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
FONT = "Times New Roman"
CODE_FONT = "Courier New"


def set_run_font(run, size: float = 14, bold: bool | None = None, font: str = FONT) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, size: float = 14) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_paragraph_text(paragraph, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, size: float = 14) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    format_paragraph(paragraph, align=align, first_line=first_line, size=size)


def add_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph._parent.add_paragraph(style=style)
    paragraph._p.addnext(new_p._p)
    if text:
        set_paragraph_text(new_p, text)
    return new_p


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def fill_table(table, headers: list[str], rows: list[list[str]], *, font_size: float = 9.0) -> None:
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate([headers] + rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p._p.clear_content()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_run_font(run, size=font_size, bold=(r_idx == 0))


def add_table_after(anchor, caption: str, headers: list[str], rows: list[list[str]], *, font_size: float = 9.0):
    caption_p = add_paragraph_after(anchor, caption)
    format_paragraph(caption_p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Cm(15.8))
    caption_p._p.addnext(table._tbl)
    spacer = anchor._parent.add_paragraph()
    table._tbl.addnext(spacer._p)
    fill_table(table, headers, rows, font_size=font_size)
    return spacer


def add_code_listing(anchor, caption: str, code: str):
    caption_p = add_paragraph_after(anchor, caption)
    format_paragraph(caption_p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    previous = caption_p
    for line in code.strip("\n").splitlines():
        p = add_paragraph_after(previous)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line.rstrip() or " ")
        set_run_font(run, size=8.0, font=CODE_FONT)
        previous = p
    spacer = add_paragraph_after(previous)
    return spacer


def rewrite_main_references(doc: Document) -> int:
    replacements = {
        "приведен в приложении А.": "приведен в приложении Б.",
        "приведены в приложениях А-В.": "приведены в приложениях Б-Г и Ж.",
        "приложения А-В.": "приложения Б-Г и Ж.",
        "приведены в приложениях Б и Г.": "приведены в приложениях В и Д.",
        "раскрыта в приложении Б.": "раскрыта в приложении В.",
        "приведена в приложении В.": "приведена в приложении Г.",
        "Приложения А-Г содержат": "Приложения Б-Ж содержат",
        "Структура исходной записи и нормализованной строки приведена в приложении Б, а команды воспроизведения сбора и обработки данных - в приложении Г.": (
            "Структура исходной записи и нормализованной строки приведена в приложении В, "
            "команды воспроизведения сбора и обработки данных - в приложении Д, "
            "а фрагменты программного кода - в приложении Ж."
        ),
    }
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new = text
        for old, value in replacements.items():
            new = new.replace(old, value)
        if new != text:
            set_paragraph_text(paragraph, new)
            changed += 1
            continue
        if text.startswith("Полученный результат может быть использован") and "приложения Б-Ж" not in text:
            set_paragraph_text(
                paragraph,
                text
                + " Вспомогательные материалы, структура данных, алгоритмические фрагменты и листинги кода вынесены в приложения Б-Ж.",
            )
            changed += 1
    return changed


def find_appendix_start(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Приложение"):
            return paragraph
    return None


def clear_existing_appendices(doc: Document) -> None:
    start = find_appendix_start(doc)
    if start is None:
        return
    node = start._p
    while node is not None:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node


def add_appendix_heading(doc: Document, anchor, letter: str, title: str | None = None):
    if anchor is None:
        heading = doc.add_paragraph(style="Heading 3")
    else:
        heading = add_paragraph_after(anchor, style="Heading 3")
    set_paragraph_text(heading, f"Приложение {letter}", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    heading.style = "Heading 3"
    heading.paragraph_format.page_break_before = True
    if title:
        title_p = add_paragraph_after(heading, title)
        format_paragraph(title_p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        return title_p
    return heading


def rebuild_appendices(doc: Document) -> None:
    clear_existing_appendices(doc)

    anchor = add_appendix_heading(doc, None, "А")

    anchor = add_appendix_heading(doc, anchor, "Б", "Состав программной реализации ПАК предиктивного обслуживания")
    anchor = add_paragraph_after(
        anchor,
        "В приложении приведены сведения о файлах и конфигурациях, которые образуют воспроизводимый прототип ПАК. Эти материалы дополняют главу 5 и не перегружают основной текст описанием служебных путей и команд.",
    )
    anchor = add_table_after(
        anchor,
        "Таблица Б.1 - Состав программных модулей ПАК предиктивного обслуживания",
        ["Модуль", "Файл или объект", "Назначение"],
        [
            ["Цифровая сцена", "scenes/vkr_scena.ttt", "Модель роботизированной ячейки паллетизации и источник телеметрии"],
            ["Сценарий паллетизации", "scripts/coppeliasim/lua/final_scene_palletizing_cycle.lua", "Формирование цикла, фаз и состояния palletizingCycle"],
            ["Коллектор телеметрии", "scripts/coppeliasim/python/collect_final_scene_telemetry.py", "Запись пакетов CoppeliaSim Remote API в JSONL"],
            ["Нормализация", "scripts/data_pipeline/normalize_telemetry.py", "Раскрытие пакетов по осям и подготовка табличной телеметрии"],
            ["Контроль качества", "scripts/data_pipeline/validate_telemetry.py", "Проверка полноты данных и фазовой разметки"],
            ["Признаки", "scripts/data_pipeline/build_features.py", "Расчет статистических, энергетических и фазовых признаков"],
            ["Сценарии деградации", "scripts/data_pipeline/simulate_degradation.py", "Формирование режимов S0...S3"],
            ["Оценка RUL", "scripts/data_pipeline/estimate_rul.py", "Расчет HI, RUL и уровня риска"],
            ["ML-модель", "scripts/data_pipeline/train_rul_mlp.py", "Обучение MLPRegressor и расчет MAE/RMSE/R2"],
            ["Визуализация", "infra/pak/grafana/dashboards/vkr_pak_dashboard.json", "Панель мониторинга HI/RUL и предупреждений"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица Б.2 - Конфигурационные файлы инфраструктурного слоя",
        ["Компонент", "Файл", "Содержание"],
        [
            ["Docker Compose", "infra/pak/docker-compose.yml", "Локальный запуск InfluxDB и Grafana"],
            ["Источник данных Grafana", "infra/pak/grafana/provisioning/datasources/influxdb.yml", "Подключение Grafana к bucket vkr_pak"],
            ["Провижининг панелей", "infra/pak/grafana/provisioning/dashboards/dashboards.yml", "Автоматическая регистрация панели ПАК"],
            ["Панель оператора", "infra/pak/grafana/dashboards/vkr_pak_dashboard.json", "Графики момента, HI, RUL, ошибки и метрик качества"],
            ["Экспорт во временные ряды", "scripts/data_pipeline/export_to_influx.py", "Запись результатов обработки в InfluxDB"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица Б.3 - Контрольные артефакты апробации",
        ["Артефакт", "Значение", "Использование в РПЗ"],
        [
            ["data/telemetry/vkr_raw/long_live_01.jsonl", "22174 пакета", "Основной набор телеметрии"],
            ["data/telemetry/vkr_normalized/vkr_telemetry_normalized.csv", "88696 строк", "Проверка полноты данных"],
            ["data/features/vkr_features.csv", "56 строк признаков", "Расчет диагностических признаков"],
            ["data/experiments/vkr_degradation_features.csv", "17920 строк", "Сценарии деградации S0...S3"],
            ["data/results/vkr_rul_estimates.csv", "17920 строк", "Базовая оценка HI/RUL"],
            ["data/results/vkr_nn_rul_predictions.csv", "17920 строк", "Нейросетевой прогноз RUL"],
            ["data/results/vkr_nn_rul_metrics.csv", "MAE=1,173; RMSE=1,442; R2=0,994", "Оценка качества модели"],
        ],
    )

    anchor = add_appendix_heading(doc, anchor, "В", "Структура экспериментальных данных и расчетных файлов")
    anchor = add_paragraph_after(
        anchor,
        "Приложение В раскрывает структуру данных, на которые ссылаются подразделы 5.4, 5.5, 6.2 и 6.3. Оно позволяет проверить, какие поля сохраняются на входе и какие таблицы формируются после обработки.",
    )
    anchor = add_table_after(
        anchor,
        "Таблица В.1 - Структура исходного JSONL-пакета телеметрии",
        ["Поле", "Пример", "Назначение"],
        [
            ["time", "162,40", "Модельное время CoppeliaSim, с"],
            ["run_id", "long_live_01", "Идентификатор эксперимента"],
            ["scenario", "nominal", "Сценарий сбора исходной телеметрии"],
            ["cycle", "7", "Номер технологического цикла"],
            ["phase", "lift_before_pick", "Фаза паллетизационного цикла"],
            ["layer", "3", "Номер слоя паллеты"],
            ["item", "water_bundle_2", "Переносимый объект"],
            ["carrying", "true", "Признак переноса нагрузки"],
            ["axes", "motor1...motor4", "Массив измерений по приводам"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица В.2 - Структура нормализованной строки телеметрии",
        ["Поле", "Пример значения", "Смысл"],
        [
            ["time", "162,40", "Время симуляции"],
            ["run_id", "long_live_01", "Идентификатор прогона"],
            ["phase", "lift_before_pick", "Фаза цикла"],
            ["axis", "motor2", "Наблюдаемый привод"],
            ["q", "0,609417", "Положение оси"],
            ["omega", "-0,011277", "Скорость оси"],
            ["accel", "-0,225548", "Ускорение оси"],
            ["torque", "2699,12", "Момент на приводе"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица В.3 - Сводка обработки набора long_live_01",
        ["Показатель", "Значение", "Комментарий"],
        [
            ["Исходные пакеты", "22174", "Строки JSONL до раскрытия осей"],
            ["Нормализованные строки", "88696", "Четыре оси на каждый пакет"],
            ["Kдан", "1,000", "Неполные строки не обнаружены"],
            ["Kфаз", "1,000", "Фазовая разметка заполнена"],
            ["Количество фаз", "14", "Включая cycle_complete и pallet_outfeed"],
            ["Средний шаг записи", "0,0929 с", "Около 10,77 Гц"],
            ["Строки признаков", "56", "Фаза x ось"],
            ["Строки RUL/NN", "17920", "Сценарии S0...S3 и синтетические циклы"],
        ],
    )

    anchor = add_appendix_heading(doc, anchor, "Г", "Фрагменты алгоритмов обработки и прогнозирования")
    anchor = add_paragraph_after(
        anchor,
        "В приложении Г приведена компактная алгоритмическая форма тех процедур, которые в основной части описаны формулами и таблицами. Полные исходные файлы остаются в рабочем каталоге проекта.",
    )
    anchor = add_table_after(
        anchor,
        "Таблица Г.1 - Алгоритм конвейера обработки телеметрии",
        ["Шаг", "Операция", "Выход"],
        [
            ["1", "Считать JSONL-пакеты эксперимента", "Список записей run_id/time/phase/axes"],
            ["2", "Раскрыть массив axes в строки по отдельным приводам", "Нормализованная таблица motor1...motor4"],
            ["3", "Проверить обязательные поля и фазовую метку", "Kдан, Kфаз, список ошибок"],
            ["4", "Сгруппировать данные по run_id, phase и axis", "Окна для расчета признаков"],
            ["5", "Рассчитать torque_rms, torque_max, omega_max, accel_rms, energy", "Таблица признаков"],
            ["6", "Сформировать сценарии S0...S3", "Расширенный набор деградации"],
            ["7", "Рассчитать HI, RUL и риск", "Таблицы vkr_rul_estimates и метрик"],
            ["8", "Обучить MLPRegressor и экспортировать прогнозы", "NN-прогнозы и MAE/RMSE/R2"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица Г.2 - Логика расчета HI и RUL в программном прототипе",
        ["Этап", "Описание", "Связь с РПЗ"],
        [
            ["Нормирование признаков", "Значения момента, энергии и производных признаков приводятся к сопоставимому масштабу", "Формулы (81)-(85)"],
            ["Модель деградации", "Для сценариев S1...S3 задается рост коэффициента деградации alpha", "Таблица 25"],
            ["Индекс HI", "HI снижается при росте расчетной поврежденности и приближении к предельному состоянию", "Формулы (74)-(77)"],
            ["Оценка RUL", "Остаточный ресурс определяется как число циклов до достижения критического HI или fail_alpha", "Формулы (86)-(88)"],
            ["Риск", "Ось считается рискованной при малом RUL или низком HI", "Формула (107)"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица Г.3 - Параметры нейросетевой модели RUL",
        ["Параметр", "Значение", "Комментарий"],
        [
            ["Библиотека", "scikit-learn", "Используется модель MLPRegressor"],
            ["Скрытый слой", "16 нейронов", "Компактная модель для учебного прототипа"],
            ["Активация", "tanh", "Стабильна для нормированных признаков"],
            ["Оптимизатор", "adam", "Градиентное обучение без ручного задания правил"],
            ["Целевая переменная", "log1p(RUL)", "Снижает влияние больших значений ресурса"],
            ["Обучающая выборка", "14336 строк", "80 процентов расширенного набора"],
            ["Тестовая выборка", "3584 строки", "20 процентов расширенного набора"],
            ["Итерации до сходимости", "204", "Фактическое число итераций обучения"],
        ],
    )

    anchor = add_appendix_heading(doc, anchor, "Д", "Материалы воспроизведения и контроля результатов")
    anchor = add_paragraph_after(
        anchor,
        "Приложение Д фиксирует команды и ожидаемые результаты, по которым можно повторить обработку данных и проверить, что численные значения в главе 6 получены из файлов проекта.",
    )
    anchor = add_table_after(
        anchor,
        "Таблица Д.1 - Команды воспроизведения обработки данных",
        ["Шаг", "Команда", "Результат"],
        [
            ["Запуск полного файлового конвейера", "python scripts/data_pipeline/run_file_pipeline.py --inputs data/telemetry/vkr_raw/long_live_01.jsonl --run-id long_live_01", "Нормализация, признаки, деградация, RUL и ML-метрики"],
            ["Построение графиков", "python scripts/data_pipeline/make_vkr_figures.py", "SVG-графики для РПЗ и презентации"],
            ["Экспорт в InfluxDB", "powershell scripts/data_pipeline/run_pipeline_and_export.ps1", "Загрузка временных рядов и расчетных показателей"],
            ["Демонстрационный запуск", "powershell scripts/pak/run_pak_demo.ps1 -RunId long_live_01", "Сбор и отображение данных в режиме демонстрации"],
        ],
        font_size=8.6,
    )
    anchor = add_table_after(
        anchor,
        "Таблица Д.2 - Контрольные результаты воспроизведения",
        ["Контроль", "Ожидаемое значение", "Назначение"],
        [
            ["DOCX-сцена", "vkr_scena.ttt", "Единое имя финальной сцены"],
            ["Kдан", "1,000", "Проверка полноты телеметрии"],
            ["Kфаз", "1,000", "Проверка фазовой разметки"],
            ["Строки признаков", "56", "Проверка группировки phase x axis"],
            ["Строки деградации", "17920", "Проверка сценариев S0...S3"],
            ["MAE", "1,173 цикла", "Средняя тестовая ошибка MLPRegressor"],
            ["RMSE", "1,442 цикла", "Среднеквадратическая ошибка"],
            ["R2", "0,994", "Коэффициент детерминации"],
        ],
    )
    anchor = add_table_after(
        anchor,
        "Таблица Д.3 - Файлы рисунков, использованные в основной части",
        ["Файл", "Раздел РПЗ", "Содержание"],
        [
            ["reports/figures/vkr_practice/torque_rms_by_axis.svg", "6.3", "Среднеквадратический момент по осям и фазам"],
            ["reports/figures/vkr_practice/hi_curves_motor1.svg", "6.5", "Динамика HI для motor1 по сценариям"],
            ["reports/figures/vkr_practice/rul_nn_actual_predicted_s3_motor1.svg", "6.4", "Сравнение фактического и прогнозного RUL"],
            ["reports/figures/vkr_practice/pak_dashboard_summary.svg", "6.5", "Сводная панель мониторинга ПАК"],
            ["reports/figures/vkr_practice_png/*.png", "Презентация", "PNG-копии графиков для вставки в Word/PowerPoint"],
        ],
        font_size=8.8,
    )

    anchor = add_appendix_heading(doc, anchor, "Ж", "Листинги ключевых программных фрагментов")
    anchor = add_paragraph_after(
        anchor,
        "В приложении Ж приведены сокращенные листинги ключевых процедур. Они показывают принцип реализации, а полные версии файлов находятся в рабочем каталоге проекта.",
    )
    anchor = add_code_listing(
        anchor,
        "Листинг Ж.1 - Получение состояния цикла и телеметрии приводов",
        '''
STATE_PROPERTY = "customData.palletizingCycle"
MOTOR_NAMES = ("motor1", "motor2", "motor3", "motor4")

def try_unpack_cycle_state(sim, model):
    try:
        packed = sim.getBufferProperty(model, STATE_PROPERTY, {"noError": True})
    except Exception:
        return {}
    if not packed:
        return {}
    try:
        unpacked = normalize_packed_value(sim.unpackTable(packed))
    except Exception:
        return {}
    return unpacked if isinstance(unpacked, dict) else {}

def get_joint_torque(sim, joint):
    try:
        return safe_float(sim.getJointForce(joint))
    except Exception:
        return 0.0
''',
    )
    anchor = add_code_listing(
        anchor,
        "Листинг Ж.2 - Последовательность файлового конвейера обработки",
        '''
def main():
    args = parse_args()
    run_step(["normalize_telemetry.py", "--inputs", *args.inputs,
              "--run-id", args.run_id, "--scenario", args.scenario])
    run_step(["validate_telemetry.py"])
    run_step(["build_features.py"])
    run_step(["simulate_degradation.py", "--cycles", str(args.cycles)])
    run_step(["estimate_rul.py", "--cycles", str(args.cycles)])
    run_step(["train_rul_mlp.py", "--cycles", str(args.cycles)])
    run_step(["make_vkr_figures.py"])
    write_json("data/results/file_pipeline_run_summary.json", summary)
''',
    )
    anchor = add_code_listing(
        anchor,
        "Листинг Ж.3 - Расчет HI, RUL и рекомендации",
        '''
hi = max(0.0, min(1.0, 1.0 - alpha / fail_alpha))
if final_alpha <= 1e-12:
    actual = cycles
else:
    actual = max(0.0, (fail_alpha - alpha) / final_alpha * cycles)

risk = "normal"
recommendation = "continue_monitoring"
if hi <= hi_crit or pred < 10:
    risk = "high"
    recommendation = "plan_maintenance"
elif hi < 0.55 or pred < 25:
    risk = "warning"
    recommendation = "increase_monitoring"
''',
    )
    add_code_listing(
        anchor,
        "Листинг Ж.4 - Обучение MLPRegressor для прогноза RUL",
        '''
model = MLPRegressor(
    hidden_layer_sizes=(hidden_size,),
    activation="tanh",
    solver="adam",
    learning_rate_init=learning_rate,
    max_iter=epochs,
    random_state=seed,
    n_iter_no_change=50,
    tol=1e-6,
)
model.fit(x[train_mask], y[train_mask])
predictions = np.maximum(0.0, np.expm1(model.predict(x)))
''',
    )


def main() -> None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOCX.with_name(f"{DOCX.stem}.backup_before_appendix_shift_code_{timestamp}{DOCX.suffix}")
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    reference_changes = rewrite_main_references(doc)
    rebuild_appendices(doc)
    doc.save(DOCX)

    print(f"docx\t{DOCX}")
    print(f"backup\t{backup}")
    print(f"reference_changes\t{reference_changes}")
    print("appendices\tА(empty), Б, В, Г, Д, Ж")


if __name__ == "__main__":
    main()
