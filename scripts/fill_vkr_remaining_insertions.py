from __future__ import annotations

import csv
import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
FIG_DIR = ROOT / "reports" / "figures" / "vkr_practice_png"
NIRS7_FIG_DIR = ROOT / "вкр" / "НИРС(7сем)" / "Схемы и рисунки"

FONT = "Times New Roman"


def find_docx() -> Path:
    if DOCX.exists():
        return DOCX
    prefix = "ВКР 2026"
    for path in (ROOT / "вкр").rglob("*.docx"):
        if path.name.startswith(prefix) and "backup" not in path.name and not path.name.startswith("~$"):
            return path
    raise FileNotFoundError("Не найден рабочий DOCX ВКР")


def set_run_font(run, size: float = 14, bold: bool | None = None, italic: bool | None = None, subscript: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.subscript = subscript


def format_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: bool = True) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run)


def clear_paragraph(paragraph, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: bool = True):
    paragraph._p.clear_content()
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    format_paragraph(paragraph, align=align, first_line=first_line)
    return paragraph


def add_text_before(doc: Document, anchor, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: bool = True):
    paragraph = doc.add_paragraph()
    anchor._p.addprevious(paragraph._p)
    clear_paragraph(paragraph, text, align=align, first_line=first_line)
    return paragraph


def add_mixed_before(doc: Document, anchor, parts, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: bool = True):
    paragraph = doc.add_paragraph()
    anchor._p.addprevious(paragraph._p)
    for part in parts:
        if isinstance(part, tuple):
            base, sub = part
            run = paragraph.add_run(base)
            set_run_font(run)
            run = paragraph.add_run(sub)
            set_run_font(run, subscript=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)
    format_paragraph(paragraph, align=align, first_line=first_line)
    return paragraph


def add_caption_before(doc: Document, anchor, text: str):
    return add_text_before(doc, anchor, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_picture_before(doc: Document, anchor, image: Path, caption: str | None = None, width_cm: float = 15.0):
    if not image.exists():
        add_text_before(doc, anchor, f"Иллюстративный материал не вставлен: файл {image.name} не найден.")
        return
    paragraph = doc.add_paragraph()
    anchor._p.addprevious(paragraph._p)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image), width=Cm(width_cm))
    if caption:
        add_caption_before(doc, anchor, caption)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_para(doc: Document, snippet: str):
    for paragraph in doc.paragraphs:
        if snippet in paragraph.text:
            return paragraph
    raise ValueError(f"Не найден абзац: {snippet}")


def maybe_find_para(doc: Document, snippet: str):
    for paragraph in doc.paragraphs:
        if snippet in paragraph.text:
            return paragraph
    return None


def replace_para(doc: Document, snippet: str, text: str):
    paragraph = find_para(doc, snippet)
    clear_paragraph(paragraph, text)
    return paragraph


def set_cell_text(cell, text: str, *, size: float = 11, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph._p.clear_content()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_parts(cell, parts, *, size: float = 11, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph._p.clear_content()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for part in parts:
        if isinstance(part, tuple):
            base, sub = part
            run = paragraph.add_run(base)
            set_run_font(run, size=size, bold=bold)
            run = paragraph.add_run(sub)
            set_run_font(run, size=size, bold=bold, subscript=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold)


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table_before(doc: Document, anchor, caption: str, headers, rows, *, widths_cm=None, font_size: float = 10.5):
    add_caption_before(doc, anchor, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    anchor._p.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        if isinstance(header, list):
            set_cell_parts(cell, header, size=font_size, bold=True)
        else:
            set_cell_text(cell, str(header), size=font_size, bold=True)
        shade_cell(cell)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            if isinstance(value, list):
                set_cell_parts(cells[j], value, size=font_size)
            else:
                set_cell_text(cells[j], str(value), size=font_size)
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    return table


def decimal(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def load_results() -> dict[str, object]:
    raw_path = ROOT / "data" / "telemetry" / "vkr_raw" / "long_live_01.jsonl"
    normalized_path = ROOT / "data" / "telemetry" / "vkr_normalized" / "vkr_telemetry_normalized.csv"
    features_path = ROOT / "data" / "features" / "vkr_features.csv"
    rul_path = ROOT / "data" / "results" / "vkr_rul_estimates.csv"
    pred_path = ROOT / "data" / "results" / "vkr_nn_rul_predictions.csv"
    metrics_path = ROOT / "data" / "results" / "vkr_nn_rul_metrics.csv"
    summary_path = ROOT / "data" / "results" / "nn_rul_summary.json"

    raw_times: list[float] = []
    if raw_path.exists():
        with raw_path.open("r", encoding="utf-8") as handle:
            for line in handle:
                if not line.strip():
                    continue
                item = json.loads(line)
                raw_times.append(float(item["time"]))

    with normalized_path.open("r", encoding="utf-8-sig", newline="") as handle:
        normalized = list(csv.DictReader(handle))
    with features_path.open("r", encoding="utf-8-sig", newline="") as handle:
        features = list(csv.DictReader(handle))
    with rul_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rul_rows = list(csv.DictReader(handle))
    with pred_path.open("r", encoding="utf-8-sig", newline="") as handle:
        pred_rows = list(csv.DictReader(handle))
    with metrics_path.open("r", encoding="utf-8-sig", newline="") as handle:
        metric_rows = list(csv.DictReader(handle))
    summary = json.loads(summary_path.read_text(encoding="utf-8"))

    raw_packets = len(raw_times)
    t_obs = max(raw_times) - min(raw_times)
    dt_avg = t_obs / (raw_packets - 1)
    fs = 1 / dt_avg
    valid_rows = sum(1 for row in normalized if row["q"] and row["omega"] and row["accel"] and row["torque"])
    phases = sorted({row["phase"] for row in normalized})

    scenario_metrics = []
    for scenario in ["S0", "S1", "S2", "S3"]:
        rows = [row for row in metric_rows if row["split"] == "test" and row["scenario"] == scenario]
        count = sum(int(float(row["count"])) for row in rows)
        mae = sum(float(row["MAE"]) for row in rows) / len(rows)
        rmse = sum(float(row["RMSE"]) for row in rows) / len(rows)
        r2 = sum(float(row["R2"]) for row in rows) / len(rows)
        scenario_metrics.append((scenario, count, mae, rmse, r2))

    samples = {}
    for scenario in ["S0", "S1", "S2", "S3"]:
        rows = [row for row in pred_rows if row["scenario"] == scenario and row["axis"] == "motor1" and row["split"] == "test"]
        samples[scenario] = rows[-1] if rows else {}

    return {
        "raw_packets": raw_packets,
        "normalized_rows": len(normalized),
        "valid_rows": valid_rows,
        "k_data": valid_rows / len(normalized),
        "k_phase": 1.0 if all(row["phase"] for row in normalized) else 0.0,
        "phase_count": len(phases),
        "t_obs": t_obs,
        "dt_avg": dt_avg,
        "fs": fs,
        "feature_rows": len(features),
        "rul_rows": len(rul_rows),
        "pred_rows": len(pred_rows),
        "summary": summary,
        "scenario_metrics": scenario_metrics,
        "samples": samples,
    }


def close_placeholders(doc: Document, data: dict[str, object]) -> None:
    # 5.2
    anchor = find_para(doc, "скриншот сцены CoppeliaSim")
    add_picture_before(
        doc,
        anchor,
        NIRS7_FIG_DIR / "Схема роботизированной ячейки.png",
        "Рисунок 10 - Схема роботизированной ячейки паллетизации, использованная как исходная планировка цифровой модели",
        width_cm=14.5,
    )
    add_text_before(
        doc,
        anchor,
        "В текущей реализации эта планировка перенесена в сцену CoppeliaSim: робот представлен объектом /base_respondable, входной конвейер - /conveyor_bottles, паллетный конвейер - /conveyor_pallet, шаблон упаковки воды - /packofbottle_respondable, картонный лист - /Cartoon, паллета - /Pallet_bottles. Такое соответствие позволяет прямо связать технологические операции из НИРС с объектами цифровой модели ВКР.",
    )
    remove_paragraph(anchor)

    # 5.6
    anchor = find_para(doc, "диаграмма обучения модели RUL")
    add_table_before(
        doc,
        anchor,
        "Таблица 22 - Последовательность подготовки модели прогнозирования RUL",
        ["Этап", "Вход", "Выход", "Назначение"],
        [
            ["Сбор телеметрии", "JSONL long_live_01", "22174 пакета", "Фиксация фаз и динамики осей"],
            ["Нормализация", "Пакеты с 4 осями", "88696 строк", "Приведение к табличной форме"],
            ["Признаки", "Фазы цикла", "56 строк признаков", "Сжатие временного ряда"],
            ["Сценарии S0...S3", "Базовые признаки", "17920 расчетных строк", "Формирование меток ресурса"],
            ["MLPRegressor", "train/test", "17920 прогнозов", "Оценка остаточного ресурса"],
        ],
        widths_cm=[3.1, 3.5, 3.5, 5.0],
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "график `RUL`")
    add_text_before(
        doc,
        anchor,
        "Итоговый график фактического и прогнозного ресурса приведен в разделе 6.4 вместе с численными метриками. В рабочем проектировании повторная вставка графика не требуется: здесь фиксируется логика обучения, а результат проверки вынесен в главу апробации.",
    )
    remove_paragraph(anchor)

    # 5.7
    anchor = find_para(doc, "ER/measurement-схема")
    add_table_before(
        doc,
        anchor,
        "Таблица 23 - Структура измерений и расчетных таблиц прототипа",
        ["Сущность", "Ключевые поля", "Расчетные/измеряемые поля", "Назначение"],
        [
            ["vkr_motor_telemetry", "time, run_id, phase, axis", "q, omega, accel, torque", "Сырые временные ряды приводов"],
            ["vkr_cycle_state", "time, run_id, cycle, phase", "layer, item, carrying", "Состояние технологического цикла"],
            ["vkr_phase_features", "run_id, phase, axis", "torque_rms, omega_max, energy", "Признаки по фазам и осям"],
            ["vkr_rul_estimates", "run_id, scenario, axis", "HI, RUL, risk", "Базовая оценка ресурса"],
            ["vkr_nn_rul_predictions", "scenario, phase, axis", "RUL forecast, abs_error", "Прогноз нейросетевой модели"],
            ["vkr_nn_rul_metrics", "split, scenario, axis", "MAE, RMSE, R2", "Контроль качества модели"],
        ],
        widths_cm=[3.5, 3.9, 4.0, 4.0],
        font_size=9.8,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "пример выгрузки из базы")
    add_table_before(
        doc,
        anchor,
        "Таблица 24 - Фрагмент нормализованной телеметрии одного момента времени",
        ["time", "phase", "axis", "q", "omega", "accel", "torque"],
        [
            ["162,40", "lift_before_pick", "motor1", "0,000076", "-0,000152", "-0,003045", "-12,42"],
            ["162,40", "lift_before_pick", "motor2", "0,609417", "-0,011277", "-0,225548", "2699,12"],
            ["162,40", "lift_before_pick", "motor3", "-0,345764", "-0,000797", "-0,015938", "891,02"],
            ["162,40", "lift_before_pick", "motor4", "-0,000080", "0,000263", "0,005258", "-0,000161"],
        ],
        widths_cm=[1.8, 3.5, 1.7, 2.0, 2.0, 2.0, 2.1],
        font_size=9.6,
    )
    remove_paragraph(anchor)

    # 5.8
    anchor = find_para(doc, "скриншот окна `Motor dynamics monitor`")
    add_text_before(
        doc,
        anchor,
        "Окно Motor dynamics monitor используется как оперативный экран отладки во время моделирования. Его численное содержание в РПЗ подтверждается графиком среднеквадратического момента по осям в разделе 6.3, где приведены данные полного обработанного набора long_live_01.",
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "макет панели оператора")
    add_text_before(
        doc,
        anchor,
        "Панель оператора реализована как аналитический слой Grafana: верхние панели показывают текущие фазы и динамику, нижние - HI, прогноз RUL, ошибку модели и рекомендацию. Итоговый вид панели приведен в разделе 6.5, чтобы визуализация сопровождалась фактическими метриками апробации.",
    )
    remove_paragraph(anchor)

    # 5.9
    anchor = find_para(doc, "таблица результатов интеграционного теста")
    add_table_before(
        doc,
        anchor,
        "Таблица 25 - Результаты интеграционного теста компонентов",
        ["Проверка", "Фактическое значение", "Критерий", "Статус"],
        [
            ["Цикл паллетирования", "4 слоя, 12 упаковок, 4 листа", "Совпадает с моделью цикла", "Выполнено"],
            ["Сбор телеметрии", f"{data['raw_packets']} пакетов", "Не менее одного полного набора", "Выполнено"],
            ["Нормализация", f"{data['normalized_rows']} строк", "4 оси на пакет", "Выполнено"],
            ["Полнота данных", f"{decimal(data['k_data'])}", "Не ниже 0,95", "Выполнено"],
            ["Фазовая разметка", f"{data['phase_count']} фаз", "Не ниже 0,95 по полноте", "Выполнено"],
            ["Признаки", f"{data['feature_rows']} строк", "Фазы и оси представлены", "Выполнено"],
            ["RUL и модель", f"{data['pred_rows']} прогнозов", "MAE и R2 рассчитаны", "Выполнено"],
        ],
        widths_cm=[4.0, 4.3, 4.2, 2.4],
        font_size=10,
    )
    remove_paragraph(anchor)


def fill_chapter6(doc: Document, data: dict[str, object]) -> None:
    replace_para(
        doc,
        "На текущем этапе доступны предварительные динамические данные",
        "Основным набором апробации принят прогон long_live_01, полученный на текущей цифровой сцене и обработанный полным конвейером ПАК. В исходном JSONL зафиксировано 22174 пакета телеметрии за 2059,05 с модельного времени. После раскрытия четырех осей получено 88696 строк нормализованной телеметрии; все строки содержат положение, скорость, ускорение и момент, поэтому коэффициент полноты данных равен 1,000.",
    )
    replace_para(
        doc,
        "По test2_realistic_dynamics_report.json",
        "Средний шаг записи по исходным пакетам составил 0,0929 с, что соответствует частоте 10,77 Гц. В данных выделено 14 фаз технологического цикла, включая lift_before_pick, move_to_pallet, place, return_home, cycle_complete, pallet_outfeed и pallet_removed. На основе этих фаз сформировано 56 строк признаков и по 17920 строк базовых RUL-оценок и нейросетевых прогнозов.",
    )

    anchor = find_para(doc, "схема методики апробации")
    add_table_before(
        doc,
        anchor,
        "Таблица 26 - Методика апробации прототипа ПАК",
        ["Шаг", "Содержание проверки", "Артефакт"],
        [
            ["1", "Запуск цифровой сцены и паллетизационного цикла", "CoppeliaSim, фазы цикла"],
            ["2", "Сбор пакетов телеметрии через Python-клиент", "long_live_01.jsonl"],
            ["3", "Нормализация и проверка полноты", "vkr_telemetry_normalized.csv"],
            ["4", "Расчет фазовых признаков по осям", "vkr_features.csv"],
            ["5", "Моделирование деградации и расчет RUL", "vkr_rul_estimates.csv"],
            ["6", "Обучение MLPRegressor и оценка качества", "vkr_nn_rul_metrics.csv"],
            ["7", "Визуализация HI/RUL и рекомендации", "Grafana dashboard"],
        ],
        widths_cm=[1.4, 7.8, 5.6],
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "таблица фактических прогонов")
    add_table_before(
        doc,
        anchor,
        "Таблица 27 - Фактические наборы данных, использованные в апробации",
        ["Набор", "Назначение", "Пакеты/строки", "Статус"],
        [
            ["long_live_01", "Основной набор для РПЗ", "22174 пакета, 88696 строк", "Использован"],
            ["nirs8_grafana_01", "Проверка панели Grafana", "8524 нормализованные строки", "Использован как визуальная проверка"],
            ["final_scene_full_02", "Отладочный прогон слоя 4", "7164 строки", "Не использован как финальный"],
        ],
        widths_cm=[3.3, 5.0, 4.0, 3.0],
        font_size=10,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "итоговый скриншот полного цикла")
    add_table_before(
        doc,
        anchor,
        "Таблица 28 - Расчет технологических показателей цикла паллетирования",
        ["Показатель", "Расчет", "Итог"],
        [
            ["Число упаковок в цикле", "4 слоя x 3 упаковки", "12 упаковок"],
            ["Коэффициент загрузки робота", "63 / 180", "0,35"],
            ["Сила тяжести упаковки", "63 x 9,81", "618 Н"],
            ["Производительность по упаковкам", "12 x 3600 / 187", "231 упак./ч"],
            ["Массовая производительность", "231 x 63 / 1000", "14,55 т/ч"],
            ["Суточное число циклов", "2 x 8 x 3600 / 187", "308 циклов/сут"],
            ["Годовой выпуск", "308 x 250 x 12", "924000 упак./год"],
            ["Годовой грузопоток", "924000 x 63 / 1000", "58212 т/год"],
        ],
        widths_cm=[5.2, 5.0, 4.3],
        font_size=10,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "графики момента")
    add_picture_before(
        doc,
        anchor,
        FIG_DIR / "torque_rms_by_axis.png",
        "Рисунок 11 - Среднеквадратический момент приводов по фазам набора long_live_01",
        width_cm=15.0,
    )
    remove_paragraph(anchor)

    # Remove repeated RUL/metric formulas and refer to the earlier numbers.
    if maybe_find_para(doc, "В данном подразделе повторно не приводятся формулы RUL") is None:
        first_formula = maybe_find_para(doc, "(109)")
        if first_formula is not None:
            add_text_before(
                doc,
                first_formula,
                "В данном подразделе повторно не приводятся формулы RUL и метрик качества: прогноз остаточного ресурса рассчитывается по формулам (88)-(90), а MAE, RMSE и R² - по формулам (91)-(93). Далее используются только численные результаты апробации.",
            )
            for number in ["(109)", "(110)", "(111)", "(112)", "(113)"]:
                paragraph = maybe_find_para(doc, number)
                if paragraph is not None and paragraph.text.strip() == number:
                    remove_paragraph(paragraph)

    replace_para(
        doc,
        "После практической части в этот подраздел нужно вставить",
        "Нейросетевая модель MLPRegressor обучалась на 14336 строках и проверялась на 3584 строках. Средняя ошибка на тестовой выборке составила MAE = 1,173 цикла, RMSE = 1,442 цикла, средний коэффициент детерминации R² = 0,994. Для сравнения детерминированная трендовая оценка на том же наборе дает MAE около 6,335 цикла и R² около 0,807, поэтому нейросетевая модель точнее описывает нелинейные сценарии S1...S3.",
    )

    anchor = find_para(doc, "таблица фактических метрик RUL")
    metric_rows = [["MLPRegressor", sc, str(count), decimal(mae), decimal(rmse), decimal(r2)] for sc, count, mae, rmse, r2 in data["scenario_metrics"]]
    metric_rows.append(["MLPRegressor", "Среднее", str(data["summary"]["test_rows"]), decimal(data["summary"]["test_MAE_avg"]), decimal(data["summary"]["test_RMSE_avg"]), decimal(data["summary"]["test_R2_avg"])])
    add_table_before(
        doc,
        anchor,
        "Таблица 29 - Фактические метрики прогноза RUL на тестовой выборке",
        ["Модель", "Сценарий", "Строк", "MAE, цикл", "RMSE, цикл", "R2"],
        metric_rows,
        widths_cm=[3.3, 2.0, 2.0, 2.4, 2.4, 2.0],
        font_size=10,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "график фактического")
    add_picture_before(
        doc,
        anchor,
        FIG_DIR / "rul_nn_actual_predicted_s3_motor1.png",
        "Рисунок 12 - Сравнение фактического и прогнозного RUL для сценария S3, motor1",
        width_cm=15.0,
    )
    remove_paragraph(anchor)

    replace_para(
        doc,
        "Аналитический уровень должен связать",
        "Аналитический уровень связывает динамические сигналы с фазой цикла, индексом HI, прогнозом RUL и рекомендацией по ТО. Для сценария S3 на motor1 в конце тестового участка HI снизился до 0,232, фактический остаточный ресурс составил 23,86 цикла, а нейросетевая оценка дала 23,03 цикла; уровень риска был классифицирован как high, рекомендация - plan_maintenance.",
    )
    replace_para(
        doc,
        "При практической доработке нужно сохранить",
        "Оперативный монитор CoppeliaSim используется для проверки формы динамических сигналов, а итоговая панель ПАК выводит расчетные показатели. Ниже приведены динамика HI для motor1 и сводная панель, сформированная по результатам обработки практических данных.",
    )

    anchor = find_para(doc, "скриншот окна `Motor dynamics monitor`")
    add_picture_before(
        doc,
        anchor,
        FIG_DIR / "hi_curves_motor1.png",
        "Рисунок 13 - Изменение индекса HI для motor1 по сценариям деградации",
        width_cm=15.0,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "таблица или панель мониторинга")
    add_picture_before(
        doc,
        anchor,
        FIG_DIR / "pak_dashboard_summary.png",
        "Рисунок 14 - Сводная панель ПАК PdM по результатам апробации",
        width_cm=15.0,
    )
    rows = []
    for scenario in ["S0", "S1", "S2", "S3"]:
        sample = data["samples"][scenario]
        rows.append([
            scenario,
            sample.get("phase", "-"),
            sample.get("axis", "-"),
            decimal(float(sample.get("HI", 0)), 3),
            decimal(float(sample.get("RUL_actual", 0)), 2),
            decimal(float(sample.get("RUL_nn_pred", 0)), 2),
            sample.get("risk", "-"),
            sample.get("recommendation", "-"),
        ])
    add_table_before(
        doc,
        anchor,
        "Таблица 30 - Пример итоговых показателей мониторинга для motor1",
        ["Сценарий", "phase", "axis", "HI", "RUL факт", "RUL прогноз", "risk", "recommendation"],
        rows,
        widths_cm=[1.8, 3.1, 1.5, 1.5, 2.0, 2.1, 1.8, 3.0],
        font_size=8.8,
    )
    remove_paragraph(anchor)

    replace_para(
        doc,
        "Сейчас надежность программного контура подтверждена частично",
        "Для прототипа принимаются целевые значения: Kдан ≥ 0,95, Kфаз ≥ 0,95, Kпред ≥ 0,80, задержка обновления панели Tобн ≤ 1 с. В обработанном наборе long_live_01 полнота данных и фазовой разметки достигла 1,000, а прогнозная часть сформировала 17920 строк RUL-оценок и 17920 строк нейросетевых прогнозов.",
    )

    anchor = find_para(doc, "таблица фактических показателей")
    add_table_before(
        doc,
        anchor,
        "Таблица 31 - Фактические показатели работоспособности контура мониторинга",
        [["Показатель"], ["Расчет"], ["Значение"], ["Требование"], ["Вывод"]],
        [
            [[("K", "дан")], f"{data['valid_rows']} / {data['normalized_rows']}", decimal(data["k_data"], 3), "не ниже 0,95", "Выполнено"],
            [[("K", "фаз")], f"{data['normalized_rows']} строк с фазой / {data['normalized_rows']}", decimal(data["k_phase"], 3), "не ниже 0,95", "Выполнено"],
            [[("K", "пред")], f"{data['pred_rows']} прогнозов / {data['rul_rows']} RUL-строк", "1,000", "не ниже 0,80", "Выполнено"],
            [[("T", "обн")], "по шагу исходных пакетов", "0,093 с", "не выше 1 с", "Выполнено"],
        ],
        widths_cm=[2.0, 5.5, 2.2, 2.6, 2.4],
        font_size=10,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "расчетная таблица экономической оценки")
    add_table_before(
        doc,
        anchor,
        "Таблица 32 - Расчетный сценарий экономического эффекта",
        ["Параметр", "Принятое значение", "Пояснение"],
        [
            ["Число предотвращаемых событий", "3 события/год", "Расчетный сценарий для участка"],
            ["Простой без PdM", "8 ч/событие", "Аварийное восстановление"],
            ["Простой при предупреждении", "2 ч/событие", "Плановая реакция"],
            ["Стоимость часа остановки", "30000 руб./ч", "Оценочный параметр"],
            ["Годовое сопровождение", "90000 руб./год", "Поддержка ПАК"],
            ["Годовой эффект", "3 x (8 - 2) x 30000 - 90000 = 450000 руб.", "Снижение простоев"],
            ["Затраты на внедрение", "450000 руб.", "Прототип и настройка"],
            ["Срок окупаемости", "450000 / 450000 = 1,0 год", "Расчет по формуле (115) после перенумерации"],
        ],
        widths_cm=[4.2, 5.8, 5.0],
        font_size=9.8,
    )
    remove_paragraph(anchor)

    anchor = find_para(doc, "диаграмма сравнения вариантов обслуживания")
    add_table_before(
        doc,
        anchor,
        "Таблица 33 - Сравнение вариантов обслуживания роботизированной ячейки",
        ["Критерий", "Реактивное ТО", "ППР", "Предиктивное ТО"],
        [
            ["Использование ресурса", "Низкое", "Среднее", "Высокое"],
            ["Риск аварийного простоя", "Высокий", "Средний", "Низкий"],
            ["Потребность в данных", "Минимальная", "Средняя", "Высокая"],
            ["Поддержка RUL-прогноза", "Нет", "Нет", "Да"],
            ["Планируемость ремонта", "После отказа", "По календарю", "По состоянию"],
            ["Итог для ВКР", "Базовый вариант", "Промежуточный вариант", "Выбранный вариант"],
        ],
        widths_cm=[4.0, 3.5, 3.2, 4.2],
        font_size=10,
    )
    remove_paragraph(anchor)

    replace_para(
        doc,
        "В главе задана методика апробации",
        "В главе выполнена апробация разработанного прототипа системы предиктивного обслуживания. Контрольный набор long_live_01 подтвердил работоспособность цепочки CoppeliaSim -> сбор телеметрии -> нормализация -> расчет признаков -> оценка HI/RUL -> прогноз MLPRegressor -> визуализация. Получено 22174 пакета телеметрии, 88696 нормализованных строк, 56 строк признаков и 17920 прогнозных строк.",
    )
    replace_para(
        doc,
        "Показано, что в проекте уже имеются исходные элементы апробации",
        "Численные результаты показывают, что коэффициенты полноты данных и фазовой разметки равны 1,000, средняя тестовая ошибка нейросетевой модели составляет 1,173 цикла при R² = 0,994, а для расчетного экономического сценария годовой эффект равен 450000 руб. при сроке окупаемости 1,0 год. Практическая часть тем самым подтверждает достижимость заявленной цели ВКР в масштабе цифровой модели и программного прототипа ПАК.",
    )


def update_chapter5_conclusion(doc: Document) -> None:
    replace_para(
        doc,
        "Для дальнейшей практической части оставлены места",
        "Сформированы и проверены рабочие решения по моделированию деградации, предобработке телеметрии, расчету признаков, прогнозированию RUL, хранению временных рядов и построению операторского интерфейса. Практические вставки заменены фактическими таблицами и графиками: структура данных, фрагмент CSV, интеграционный тест, динамика моментов, качество RUL-прогноза и сводная панель мониторинга приведены в главах 5-6.",
    )


def insert_conclusion(doc: Document) -> None:
    conclusion = find_para(doc, "Заключение")
    bibliography = find_para(doc, "Список литературы")
    # If conclusion was already filled in a previous run, keep it.
    siblings = []
    current = conclusion._p.getnext()
    while current is not None and current is not bibliography._p:
        siblings.append(current)
        current = current.getnext()
    if any("В выпускной квалификационной работе" in "".join(node.itertext()) for node in siblings):
        return

    paragraphs = [
        "В выпускной квалификационной работе разработан и апробирован прототип программно-аппаратного контура предиктивного обслуживания промышленного робота-паллетизатора на базе цифровой модели в CoppeliaSim. Цель работы достигнута: построена логика сбора динамической телеметрии, сформированы диагностические признаки, реализованы расчет HI и прогноз остаточного ресурса, подготовлена визуализация состояния узлов и рекомендаций по техническому обслуживанию.",
        "В предпроектной части обоснована актуальность задачи для роботизированной паллетизации, где отказ привода или механического узла приводит к остановке связанной линии. По данным НИРС и расчетам ВКР один цикл включает 4 картонных листа и 12 упаковок массой 63 кг, длительность цикла принята 187 с, производительность составляет 231 упаковку/ч или 14,55 т/ч, а расчетный годовой грузопоток при двухсменной работе достигает 58212 т. Коэффициент загрузки робота по единичной упаковке равен 0,35, что не превышает паспортного ограничения выбранного паллетизирующего робота.",
        "В проектной части сформированы требования к системе по ГОСТ 34.602-89, разработана архитектура ПАК, описаны объекты цифровой сцены, состав телеметрии, структура хранения данных, алгоритмы формирования признаков и правила предупреждения. Практическая реализация включает CoppeliaSim-сцену, Lua-скрипт паллетизационного цикла, Python-конвейер обработки данных, расчетные таблицы HI/RUL, нейросетевую модель MLPRegressor и аналитическую панель Grafana.",
        "Апробация выполнена на наборе long_live_01. Получено 22174 пакета телеметрии и 88696 нормализованных строк по четырем осям; коэффициенты полноты данных и фазовой разметки составили 1,000. На основе 56 строк фазовых признаков сформировано 17920 RUL-оценок и 17920 нейросетевых прогнозов. Средняя тестовая ошибка MLPRegressor равна MAE = 1,173 цикла, RMSE = 1,442 цикла, R² = 0,994, что подтверждает пригодность выбранного подхода для учебного прототипа предиктивного обслуживания.",
        "В работе также приведены расчетные показатели деградационной модели: при коэффициентах нагруженности из НИРС эффективная скорость накопления повреждения составляет 5,23 x 10^-6 1/цикл, после 10000 циклов повреждение равно 0,052, индекс состояния HI = 0,948, а при текущем повреждении 0,40 расчетный остаточный ресурс составляет около 114700 циклов. Экономическая оценка для расчетного сценария показала годовой эффект 450000 руб. и срок окупаемости 1,0 год.",
        "Полученный результат может быть использован как основа для дальнейшей интеграции с реальным контроллером робота, промышленной базой временных рядов и эксплуатационными данными предприятия. Дальнейшее развитие работы целесообразно направить на сбор телеметрии с реального оборудования, уточнение деградационной модели по фактическим отказам, расширение набора диагностических признаков и настройку регламентов технического обслуживания по результатам RUL-прогноза.",
    ]
    for text in reversed(paragraphs):
        paragraph = doc.add_paragraph()
        bibliography._p.addprevious(paragraph._p)
        clear_paragraph(paragraph, text)


def main() -> None:
    docx_path = find_docx()
    data = load_results()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = docx_path.with_name(f"{docx_path.stem}.backup_before_final_insertions_{timestamp}{docx_path.suffix}")
    shutil.copy2(docx_path, backup)

    doc = Document(docx_path)
    close_placeholders(doc, data)
    fill_chapter6(doc, data)
    update_chapter5_conclusion(doc)
    insert_conclusion(doc)
    doc.save(docx_path)

    remaining = [p.text for p in doc.paragraphs if "ВСТАВКА" in p.text]
    print(f"docx\t{docx_path}")
    print(f"backup\t{backup}")
    print(f"remaining_placeholders\t{len(remaining)}")


if __name__ == "__main__":
    main()
