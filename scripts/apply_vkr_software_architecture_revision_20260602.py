from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
FONT_NAME = "Times New Roman"


def set_run_font(run, *, size_pt: int = 14, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(key), FONT_NAME)


def format_paragraph(paragraph: Paragraph, *, heading: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, bold=True if heading else None)
    if heading:
        paragraph.style = "Heading 3"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def replace_paragraph_text(paragraph: Paragraph, text: str, *, heading: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, bold=True if heading else None)
    format_paragraph(paragraph, heading=heading)


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    *,
    style=None,
    italic: bool = False,
    centered: bool = False,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    run = new_paragraph.add_run(text)
    set_run_font(run, italic=italic)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    return new_paragraph


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def move_block_before(doc: Document, start_prefix: str, end_before_prefix: str, before_prefix: str) -> None:
    body = doc.element.body
    start_p = find_paragraph(doc, start_prefix)._element
    end_p = find_paragraph(doc, end_before_prefix)._element
    before_p = find_paragraph(doc, before_prefix)._element
    children = list(body)
    start_idx = children.index(start_p)
    end_idx = children.index(end_p)
    block = children[start_idx:end_idx]
    for element in block:
        body.remove(element)
    insert_idx = list(body).index(before_p)
    for offset, element in enumerate(block):
        body.insert(insert_idx + offset, element)


def delete_headings(doc: Document, prefixes: tuple[str, ...]) -> int:
    deleted = 0
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes):
            delete_paragraph(paragraph)
            deleted += 1
    return deleted


def replace_table_after_caption(doc: Document, caption_prefix: str, rows: list[list[str]]) -> None:
    target_caption = None
    old_table = None
    found_caption = False
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            if block.text.strip().startswith(caption_prefix):
                target_caption = block
                found_caption = True
        elif isinstance(block, Table) and found_caption:
            old_table = block
            break
    if target_caption is None or old_table is None:
        raise ValueError(f"table after caption not found: {caption_prefix}")

    new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    new_table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = new_table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, bold=True if r_idx == 0 else None)

    body = doc.element.body
    old_elem = old_table._element
    new_elem = new_table._element
    body.remove(new_elem)
    insert_idx = list(body).index(old_elem)
    body.insert(insert_idx, new_elem)
    body.remove(old_elem)


def apply_chapter2_revision(doc: Document) -> None:
    move_block_before(doc, "2.5. Выбор технологической реализации", "2.6. Формирование требований", "2.2. Общая структура системы")

    replace_paragraph_text(
        find_paragraph(doc, "2. Концептуальное проектирование"),
        "2. Концептуальное проектирование",
        heading=True,
    )
    replace_paragraph_text(
        doc.paragraphs[
            next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Концептуальное проектирование определяет"))
        ],
        "Концептуальное проектирование определяет исходный инструментальный стек, архитектурную схему и поток данных системы предиктивного обслуживания. В главе сначала обосновывается выбор программных средств, затем формируется общая архитектура ПАК, функциональная модель, архитектурный вариант и укрупненные требования к прототипу.",
    )

    delete_headings(
        doc,
        (
            "2.5.1.",
            "2.5.2.",
            "2.2.1.",
            "2.3.1.",
            "2.4.1.",
            "2.4.2.",
        ),
    )

    replace_paragraph_text(find_paragraph(doc, "2.5. Выбор технологической реализации"), "2.2. Обоснование выбора программных средств", heading=True)
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("CoppeliaSim выбран для имитации")),
        "Выбор программных средств выполнялся по критериям воспроизводимости эксперимента, доступности интерфейсов обмена данными, пригодности к обработке временных рядов, возможности визуального контроля и дальнейшего переноса на реальный источник телеметрии. Инструменты подбирались так, чтобы обеспечить полный путь данных: подготовку модели, имитацию паллетизационного цикла, сбор измерений, расчет признаков, хранение временных рядов, прогноз RUL и отображение результатов. Подробная подготовка цифровой модели робота, включая работу со STEP-моделью, настройку сборки в SolidWorks, экспорт в URDF и последующую доработку сцены в CoppeliaSim, приведена в главе 5.",
    )
    replace_paragraph_text(find_paragraph(doc, "Таблица 8 - Технологический стек системы"), "Таблица 6 - Обоснование выбора программных средств")
    replace_table_after_caption(
        doc,
        "Таблица 6 - Обоснование выбора программных средств",
        [
            ["Задача", "Выбранное средство", "Альтернативы", "Причина выбора", "Ограничение"],
            ["Имитация роботизированной ячейки", "CoppeliaSim", "Gazebo, Webots, ABB RobotStudio", "Remote API, доступ к объектам сцены и параметрам движения", "требуется ручная настройка модели"],
            ["Подготовка геометрии робота", "SolidWorks", "FreeCAD, Blender", "работа со STEP-сборками, сопряжениями и осями звеньев", "не заменяет настройку модели в симуляторе"],
            ["Перенос кинематической структуры", "URDF", "SDF, COLLADA, ручная сборка", "передача иерархии звеньев и суставов", "после импорта нужна доработка"],
            ["Сбор и обработка данных", "Python", "MATLAB, C#, Lua", "API-интеграция, JSONL/CSV, библиотеки анализа данных", "не является промышленным ПЛК-контуром"],
            ["Хранение временных рядов", "InfluxDB", "CSV, PostgreSQL", "временные метки, теги, выборки по циклам и фазам", "требует отдельного сервиса"],
            ["Визуализация", "Grafana", "Matplotlib, web-интерфейс", "готовые панели, графики и предупреждения", "зависит от структуры данных"],
            ["Прогноз RUL", "scikit-learn", "XGBoost, PyTorch", "воспроизводимая MLPRegressor-модель и метрики качества", "качество проверено на модельных данных"],
        ],
    )
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("На концептуальном уровне задача прогнозирования")),
        "Для блока прогнозирования выбрана регрессионная постановка задачи. Такая постановка соответствует data-driven подходу, при котором эксплуатация оборудования представляется временными рядами и обрабатывается методами машинного обучения [6]. В рабочей реализации основным инструментом выбран scikit-learn, поскольку он позволяет воспроизводимо обучать компактную MLPRegressor-модель и рассчитывать метрики качества без усложнения прототипа.",
    )
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("Для базовой проверки используется")),
        "Для базовой проверки допускается сравнение с линейной или ансамблевой моделью, а качество оценивается по MAE, RMSE и доле своевременных предупреждений:",
    )

    replace_paragraph_text(find_paragraph(doc, "2.2. Общая структура системы"), "2.3. Общая архитектура ПАК", heading=True)
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("Система строится как программно-аппаратный комплекс")),
        "После выбора программных средств система представляется как ПАК с логическими уровнями цифрового моделирования, сбора телеметрии, предобработки, хранения, аналитики HI/RUL, визуализации и поддержки решений по ТОиР. В рамках ВКР физический робот заменяется цифровой моделью, но логика обмена данными сохраняется как основа для будущей интеграции с реальным оборудованием. Выделение уровней моделирования, сбора, хранения, аналитики и визуализации согласуется с практикой цифровых двойников и систем мониторинга [44].",
    )
    replace_paragraph_text(find_paragraph(doc, "Таблица 6 - Компоненты архитектуры"), "Таблица 7 - Логические уровни архитектуры ПАК")
    replace_table_after_caption(
        doc,
        "Таблица 7 - Логические уровни архитектуры ПАК",
        [
            ["Уровень", "Назначение", "Входные данные", "Выходные данные", "Реализация"],
            ["Цифровая модель", "воспроизведение цикла паллетизации", "сцена, параметры цикла", "состояния объектов и приводов", "CoppeliaSim"],
            ["Сбор телеметрии", "получение измерений по осям", "состояние сцены и приводов", "JSONL-пакеты телеметрии", "Python Remote API"],
            ["Предобработка", "нормализация и привязка к фазам", "сырые измерения", "очищенные временные ряды", "Python/pandas"],
            ["Хранение", "накопление истории измерений", "телеметрия и расчетные признаки", "временные ряды по тегам", "InfluxDB"],
            ["Аналитика HI/RUL", "оценка состояния и ресурса", "признаки окон наблюдения", "HI, RUL, риск", "Python, scikit-learn"],
            ["Визуализация", "контроль состояния оператором", "показатели и прогнозы", "панели и предупреждения", "Grafana"],
        ],
    )

    replace_paragraph_text(find_paragraph(doc, "2.3. Функциональная модель системы"), "2.4. Функциональная модель и потоки данных", heading=True)
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("Функциональная модель включает шесть функций")),
        "Функциональная модель включает шесть ключевых функций: моделирование цикла, сбор телеметрии, предобработку, расчет признаков, прогнозирование состояния и визуализацию. В IDEF0-логике входами являются параметры робота, сценарий цикла и поток измерений; механизмами - CoppeliaSim, Python-модули, БД и ML-библиотеки; управляющими воздействиями - требования к частоте сбора, составу признаков и правилам предупреждений.",
    )

    replace_paragraph_text(find_paragraph(doc, "2.4. Анализ архитектуры системы"), "2.5. Анализ архитектурного варианта", heading=True)
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("Для ПАК рассматриваются три варианта")),
        "Для ПАК рассматриваются три архитектурных варианта: локальный монолит, распределенная система и гибридный вариант. Для ВКР выбран гибридный вариант: моделирование, сбор и первичная обработка выполняются локально, а хранение, аналитика и визуализация выделяются как самостоятельные компоненты. Это упрощает отладку, сохраняет воспроизводимость эксперимента и оставляет возможность масштабирования.",
    )
    replace_paragraph_text(find_paragraph(doc, "Таблица 7 - Сравнение архитектурных вариантов"), "Таблица 8 - Сравнение архитектурных вариантов")
    replace_table_after_caption(
        doc,
        "Таблица 8 - Сравнение архитектурных вариантов",
        [
            ["Вариант", "Содержание", "Преимущества", "Ограничения", "Вывод для ВКР"],
            ["Локальный монолит", "все функции в одном приложении", "простая первичная отладка", "сложно заменять компоненты", "подходит только для раннего макета"],
            ["Распределенная система", "каждая функция выделена в сервис", "масштабируемость и гибкость", "высокая сложность интеграции", "избыточна для учебного стенда"],
            ["Гибридный вариант", "модель и сбор локально, хранение и визуализация выделены", "баланс простоты и расширяемости", "требует явных интерфейсов обмена", "выбран для реализации"],
        ],
    )
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("Выбор архитектуры выполняется по критериям")),
        "Выбор архитектуры выполняется по критериям воспроизводимости экспериментов, доступности данных, модульности, возможности визуального контроля и расширяемости под реальный объект. Для сравнения вариантов можно использовать интегральную оценку:",
    )

    replace_paragraph_text(find_paragraph(doc, "2.6. Формирование требований к системе"), "2.6. Формирование требований к системе", heading=True)
    replace_paragraph_text(find_paragraph(doc, "2.7. Выводы по главе"), "2.7. Выводы по главе", heading=True)
    replace_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip().startswith("В главе сформирована концепция системы предиктивного обслуживания")),
        "В главе сформирована концепция системы предиктивного обслуживания робота-паллетизатора. Сначала обоснован выбор программных средств для подготовки модели, имитации, сбора данных, хранения, аналитики и визуализации. Затем определены логические уровни архитектуры ПАК, информационная модель измерения, функциональная цепочка обработки данных и гибридный архитектурный вариант, сочетающий цифровую модель, Python-контур, хранилище временных рядов, ML-модуль и операторскую панель.",
    )


def apply_chapter5_model_preparation(doc: Document) -> None:
    marker = "исходно была подготовлена STEP-модель робота"
    if any(marker in p.text for p in doc.paragraphs):
        return
    anchor = find_paragraph(doc, "Цифровая модель реализована в сцене")
    p1 = insert_paragraph_after(
        anchor,
        "Работа с моделью робота не сводилась к прямому импорту готового объекта в CoppeliaSim: исходно была подготовлена STEP-модель робота, содержащая геометрию звеньев, но не готовую управляемую структуру для имитационного эксперимента. Поэтому перед переносом в симулятор потребовалось восстановить и проверить состав сборки, взаимное положение звеньев, оси вращения и сопряжения в SolidWorks.",
    )
    p2 = insert_paragraph_after(
        p1,
        "После настройки сборки модель экспортировалась в формат URDF, который позволяет передать иерархию звеньев и суставов. При этом URDF-экспорт не устранил необходимость ручной доводки: после импорта в CoppeliaSim уточнялись имена объектов, структура /base_respondable, параметры сочленений motor1...motor4, dummy-связи, ограничения движения, respondable-формы, массы объектов и траектории паллетизационного цикла.",
    )
    p3 = insert_paragraph_after(
        p2,
        "Итерационная доработка была необходима из-за того, что импортированная CAD-модель сохраняла геометрию, но не гарантировала корректное поведение в задаче паллетизации. На практике отдельно проверялись замкнутые связи, достижимость точек захвата и укладки, устойчивость движения при sim.moveToConfig и отсутствие критических нарушений сцены при повторяемом цикле. Поэтому итоговая сцена рассматривается как адаптированный имитационный стенд, подготовленный для проверки контура PdM.",
    )
    insert_paragraph_after(
        p3,
        "[Место для вставки скриншота из SolidWorks: сборка робота, сопряжения и подготовка к экспорту в URDF]",
        italic=True,
        centered=True,
    )


def ooxml_renumber_formula_labels() -> dict[int, int]:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    tmp = DOCX_PATH.with_suffix(".tmp.docx")
    old_to_new: dict[int, int] = {}
    with zipfile.ZipFile(DOCX_PATH, "r") as zin:
        root = etree.fromstring(zin.read("word/document.xml"))
        next_num = 1
        for paragraph in root.xpath(".//w:p", namespaces=ns):
            if not paragraph.xpath(".//m:oMath | .//m:oMathPara", namespaces=ns):
                continue
            nodes = paragraph.xpath(".//w:t | .//m:t", namespaces=ns)
            full = "".join(node.text or "" for node in nodes)
            match = re.search(r"\((\d+)\)\s*$", full)
            if not match:
                continue
            old_num = int(match.group(1))
            new_label = f"({next_num})"
            old_to_new.setdefault(old_num, next_num)
            old_label = match.group(0).strip()
            if old_label != new_label:
                suffix_nodes = []
                suffix_text = ""
                for node in reversed(nodes):
                    suffix_nodes.append(node)
                    suffix_text = (node.text or "") + suffix_text
                    if len(suffix_text) >= len(old_label):
                        break
                label_nodes = list(reversed(suffix_nodes))
                total = "".join(node.text or "" for node in label_nodes)
                if total.endswith(old_label):
                    label_nodes[0].text = total[: -len(old_label)] + new_label
                    for node in label_nodes[1:]:
                        node.text = ""
            next_num += 1

        new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = new_xml
                zout.writestr(item, data)
    tmp.replace(DOCX_PATH)
    return old_to_new


def update_formula_reference_texts(doc: Document, old_to_new: dict[int, int]) -> int:
    changed = 0

    def repl_range(match: re.Match[str]) -> str:
        a = int(match.group(1))
        b = int(match.group(2))
        return f"({old_to_new.get(a, a)})-({old_to_new.get(b, b)})"

    def repl_single(match: re.Match[str]) -> str:
        n = int(match.group(1))
        return f"({old_to_new.get(n, n)})"

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "формул" not in text.lower() and "формуле" not in text.lower():
            continue
        new_text = re.sub(r"\((\d+)\)-\((\d+)\)", repl_range, text)
        new_text = re.sub(r"\((\d+)\)", repl_single, new_text)
        if new_text != text:
            replace_paragraph_text(paragraph, new_text)
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if "Формул" not in text and "формул" not in text:
                        continue
                    new_text = re.sub(r"\((\d+)\)-\((\d+)\)", repl_range, text)
                    new_text = re.sub(r"\((\d+)\)", repl_single, new_text)
                    if new_text != text:
                        replace_paragraph_text(paragraph, new_text)
                        changed += 1
    return changed


def audit_docx(doc: Document) -> dict[str, object]:
    report: dict[str, object] = {}
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        report["zip_bad"] = str(zf.testzip())

    main_end = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Приложение"))
    table_nums = []
    fig_nums = []
    for paragraph in doc.paragraphs[:main_end]:
        t = paragraph.text.strip()
        mt = re.match(r"^Таблица\s+(\d+)\s*[-–]", t)
        mf = re.match(r"^Рисунок\s+(\d+)\s*[-–]", t)
        if mt:
            table_nums.append(int(mt.group(1)))
        if mf:
            fig_nums.append(int(mf.group(1)))
    report["table_count"] = len(table_nums)
    report["table_max"] = max(table_nums) if table_nums else 0
    report["table_missing"] = [n for n in range(1, max(table_nums) + 1) if n not in set(table_nums)] if table_nums else []
    report["table_duplicates"] = sorted({n for n in table_nums if table_nums.count(n) > 1})
    report["visible_figure_numbers"] = fig_nums

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    formula_nums = []
    for paragraph in root.xpath(".//w:p", namespaces=ns):
        if not paragraph.xpath(".//m:oMath | .//m:oMathPara", namespaces=ns):
            continue
        full = "".join(node.text or "" for node in paragraph.xpath(".//w:t | .//m:t", namespaces=ns)).strip()
        m = re.search(r"\((\d+)\)$", full)
        if m:
            formula_nums.append(int(m.group(1)))
    report["formula_count"] = len(formula_nums)
    report["formula_max"] = max(formula_nums) if formula_nums else 0
    report["formula_missing"] = [n for n in range(1, max(formula_nums) + 1) if n not in set(formula_nums)] if formula_nums else []
    report["formula_duplicates"] = sorted({n for n in formula_nums if formula_nums.count(n) > 1})
    report["solidworks_placeholder"] = int(any("скриншота из SolidWorks" in p.text for p in doc.paragraphs))
    report["chapter2_order"] = [
        p.text.strip()
        for p in doc.paragraphs
        if re.match(r"^2\.\d+\. ", p.text.strip())
    ]
    return report


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}.backup_before_software_architecture_revision_{datetime.now():%Y%m%d_%H%M%S}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)

    doc = Document(DOCX_PATH)
    apply_chapter2_revision(doc)
    apply_chapter5_model_preparation(doc)
    for paragraph in doc.paragraphs:
        if re.match(r"^\d+(\.\d+)*\.", paragraph.text.strip()):
            format_paragraph(paragraph, heading=True)
    doc.save(DOCX_PATH)

    formula_map = ooxml_renumber_formula_labels()
    doc = Document(DOCX_PATH)
    formula_ref_changes = update_formula_reference_texts(doc, formula_map)
    doc.save(DOCX_PATH)

    audited = Document(DOCX_PATH)
    report = audit_docx(audited)
    formula_changes = {k: v for k, v in sorted(formula_map.items()) if k != v}
    print(f"backup={backup}")
    print(f"formula_changes={formula_changes}")
    print(f"formula_reference_changes={formula_ref_changes}")
    print(f"audit={report}")


if __name__ == "__main__":
    main()
