from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs)[::-1]:
        paragraph._p.remove(run._r)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_run_font(run, size: int = 14, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_num_pr(paragraph: Paragraph, num_id: str, ilvl: str = "0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def remove_num_pr(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)


def format_body(paragraph: Paragraph, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(35.4) if first_line else None
    for run in paragraph.runs:
        set_run_font(run, 14)


def format_heading(paragraph: Paragraph, level: int, numbered: bool = True) -> None:
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 14, bold=True)
    if numbered:
        set_num_pr(paragraph, "9", str(level - 1))
    else:
        remove_num_pr(paragraph)


def format_placeholder(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.line_spacing = 1.15
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, 12, italic=True)


def find_index(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if (paragraph.text or "").strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text!r}")


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\.\s*", "", text).strip()


def normalize_placeholder(text: str) -> str:
    cleaned = text.strip()
    cleaned = cleaned.removeprefix("[[").removesuffix("]]").strip()
    return f"[{cleaned}]"


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def set_cell_text(cell, text: str, header: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, 12, bold=header)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_widths(table, rows: list[list[str]]) -> None:
    col_count = len(rows[0])
    if col_count == 2:
        widths = [2.2, 4.6]
    elif col_count == 3:
        widths = [1.8, 2.45, 2.55]
    elif col_count == 4:
        widths = [1.45, 1.75, 1.85, 1.75]
    else:
        widths = [6.8 / col_count] * col_count
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]) -> Paragraph:
    doc = paragraph._parent
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.8))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, header=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    apply_table_widths(table, rows)
    paragraph._p.addnext(table._tbl)
    spacer = OxmlElement("w:p")
    table._tbl.addnext(spacer)
    spacer_para = Paragraph(spacer, paragraph._parent)
    spacer_para.paragraph_format.space_after = Pt(0)
    return spacer_para


def iter_content_blocks(markdown: str, stop_heading: str | None = None):
    lines = markdown.splitlines()
    if stop_heading:
        stop = next((idx for idx, line in enumerate(lines) if line.strip() == stop_heading), len(lines))
        lines = lines[:stop]
    if lines and lines[0].startswith("# "):
        lines = lines[1:]

    idx = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
            paragraph_buffer = []
            return ("paragraph", text)
        return None

    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            block = flush_paragraph()
            if block:
                yield block
            idx += 1
            continue

        if stripped.startswith("|"):
            block = flush_paragraph()
            if block:
                yield block
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            yield ("table", parse_table(table_lines))
            continue

        if stripped.startswith("### "):
            block = flush_paragraph()
            if block:
                yield block
            yield ("heading3", strip_heading_number(stripped[4:]))
            idx += 1
            continue

        if stripped.startswith("## "):
            block = flush_paragraph()
            if block:
                yield block
            yield ("heading2", strip_heading_number(stripped[3:]))
            idx += 1
            continue

        if stripped.startswith("- "):
            block = flush_paragraph()
            if block:
                yield block
            yield ("bullet", stripped[2:].strip())
            idx += 1
            continue

        if stripped.startswith("[[ВСТАВКА"):
            block = flush_paragraph()
            if block:
                yield block
            yield ("placeholder", normalize_placeholder(stripped))
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    block = flush_paragraph()
    if block:
        yield block


def set_update_fields_on_open(docx_path: Path) -> None:
    tmp_docx = docx_path.with_name(f"{docx_path.stem}.tmp_update_fields.docx")
    parser = etree.XMLParser(remove_blank_text=False)
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp_docx, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        names = set(src.namelist())
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "word/settings.xml":
                root = etree.fromstring(data, parser)
                update_fields = root.find("w:updateFields", NS)
                if update_fields is None:
                    update_fields = etree.Element(f"{{{W_NS}}}updateFields")
                    root.insert(0, update_fields)
                update_fields.set(f"{{{W_NS}}}val", "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            dst.writestr(info, data)
        if "word/settings.xml" not in names:
            root = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            update_fields = etree.SubElement(root, f"{{{W_NS}}}updateFields")
            update_fields.set(f"{{{W_NS}}}val", "true")
            dst.writestr(
                "word/settings.xml",
                etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
    shutil.move(str(tmp_docx), str(docx_path))


def replace_chapter(docx_path: Path, draft_path: Path, start_title: str, end_title: str) -> None:
    doc = Document(str(docx_path))
    start_idx = find_index(doc, start_title)
    end_idx = find_index(doc, end_title)

    start = doc.paragraphs[start_idx]
    clear_paragraph(start)
    start.add_run(start_title)
    format_heading(start, 1, numbered=True)

    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        delete_paragraph(paragraph)

    current = start
    markdown = draft_path.read_text(encoding="utf-8")
    for kind, value in iter_content_blocks(markdown, stop_heading="## Список рекомендуемых вставок для главы 2"):
        if kind == "paragraph":
            current = insert_paragraph_after(current, value, "Normal")
            format_body(current)
        elif kind == "heading2":
            current = insert_paragraph_after(current, value, "Heading 2")
            format_heading(current, 2, numbered=True)
        elif kind == "heading3":
            current = insert_paragraph_after(current, value, "Heading 3")
            format_heading(current, 3, numbered=True)
        elif kind == "bullet":
            current = insert_paragraph_after(current, value, "List Paragraph")
            set_num_pr(current, "10", "0")
            format_body(current, first_line=False)
        elif kind == "placeholder":
            current = insert_paragraph_after(current, value, "Normal")
            format_placeholder(current)
        elif kind == "table" and value:
            current = insert_table_after(current, value)

    doc.save(str(docx_path))
    set_update_fields_on_open(docx_path)


def main() -> None:
    if len(sys.argv) != 5:
        raise SystemExit("Usage: insert_vkr_chapter_from_md.py <docx> <draft.md> <start_title> <end_title>")
    replace_chapter(Path(sys.argv[1]), Path(sys.argv[2]), sys.argv[3], sys.argv[4])
    print(f"Inserted chapter {sys.argv[3]!r} into {sys.argv[1]}")


if __name__ == "__main__":
    main()
