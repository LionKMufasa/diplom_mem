from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
FONT = "Times New Roman"


def find_docx() -> Path:
    if DOCX.exists():
        return DOCX
    prefix = "ВКР 2026"
    candidates = [
        p
        for p in ROOT.rglob("*.docx")
        if p.name.startswith(prefix) and "backup" not in p.name and not p.name.startswith("~$")
    ]
    if not candidates:
        raise FileNotFoundError("Не найден рабочий DOCX ВКР")
    return max(candidates, key=lambda p: p.stat().st_mtime)


def set_run_font(run, size: float = 14, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT
    )
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, size: float = 14) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_paragraph_text(paragraph, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True) -> None:
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    set_run_font(run)
    format_paragraph(paragraph, align=align, first_line=first_line)


def add_paragraph_after(paragraph, text: str = "", style: str | None = None):
    doc = paragraph._parent
    new_p = doc.add_paragraph(style=style)
    paragraph._p.addnext(new_p._p)
    if text:
        set_paragraph_text(new_p, text)
    return new_p


def add_caption_after(paragraph, text: str):
    new_p = add_paragraph_after(paragraph, text)
    format_paragraph(new_p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    return new_p


def clear_and_fill_table(table, headers: list[str], rows: list[list[str]], *, font_size: float = 10.5) -> None:
    # Normalize row count.
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    # Normalize column count where possible by using the existing width.
    current_cols = len(table.columns)
    if current_cols < len(headers):
        for _ in range(len(headers) - current_cols):
            table.add_column(Cm(3))
    elif current_cols > len(headers):
        # Keep extra columns empty rather than doing risky XML surgery.
        headers = headers + [""] * (current_cols - len(headers))
        rows = [row + [""] * (current_cols - len(row)) for row in rows]

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate([headers] + rows):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p._p.clear_content()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_run_font(run, size=font_size, bold=(r_idx == 0))


def edit_all_text(doc: Document) -> int:
    replacements = {
        "83 листов основной части формата А4": "основную часть, список литературы и приложения формата А4",
        "vkr_scena.ttt": "vkr_scena.ttt",
        "final_scena_diplom.ttt": "vkr_scena.ttt",
        "pred_final.ttt": "vkr_scena.ttt",
        "scenes/final_scena_diplom.ttt": "scenes/vkr_scena.ttt",
        "scenes/pred_final.ttt": "scenes/vkr_scena.ttt",
        "формулам (88)-(90)": "формулам (86)-(88)",
        "формулам (91)-(93)": "формулам (89)-(91)",
        "Цлевое значение": "Целевое значение",
        "Будущая вставка": "Подтверждающий артефакт",
        "Будут вставлены после финальных прогонов": "Использованы в главах 5-6 по результатам обработки long_live_01",
        "collect_predictive_telemetry.py": "collect_final_scene_telemetry.py",
        "data/telemetry/test2_dynamics_monitor.csv": "data/telemetry/vkr_raw/long_live_01.jsonl",
        "Тестовые ряды координат, моментов, скоростей и ускорений": "Финальный набор телеметрии полного прогона",
        "Используются для проверки формата признаков": "Использован для расчета признаков и RUL",
        "Подготовлен для интеграции с очередью телеметрии": "Использован для записи JSONL-телеметрии",
    }
    changed = 0
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        text = paragraph.text
        new = text
        for old, value in replacements.items():
            new = new.replace(old, value)
        if new != text:
            set_paragraph_text(paragraph, new, align=paragraph.alignment or WD_ALIGN_PARAGRAPH.JUSTIFY)
            changed += 1
    return changed


def replace_captions(doc: Document) -> int:
    captions = {
        "Таблица 6": "Таблица 6 - Сравнение стратегий обслуживания на концептуальном этапе",
        "Таблица 19": "Таблица 19 - Сравнение метрик качества прогнозирования RUL",
        "Таблица 22": "Таблица 22 - Состав артефактов рабочего проекта ПАК предиктивного обслуживания",
        "Таблица 23": "Таблица 23 - Объекты сцены CoppeliaSim и их назначение в цифровой модели",
        "Таблица 24": "Таблица 24 - Фазы паллетизационного цикла и их диагностическое значение",
        "Таблица 25": "Таблица 25 - Сценарии моделирования деградации узлов робота",
        "Таблица 26": "Таблица 26 - Состав записи телеметрии привода",
        "Таблица 27": "Таблица 27 - Диагностические признаки, рассчитываемые по телеметрии",
        "Таблица 28": "Таблица 28 - Этапы подготовки и проверки модели прогнозирования RUL",
        "Таблица 30": "Таблица 30 - Проект логической структуры хранения временных рядов",
        "Таблица 33": "Таблица 33 - Состав виджетов операторской панели мониторинга",
        "Таблица 34": "Таблица 34 - Проверки интеграции компонентов ПАК",
        "Таблица 36": "Таблица 36 - Критерии апробации и подтверждающие артефакты",
        "Таблица 44": "Таблица 44 - Исходные оценки для интегрального сравнения вариантов обслуживания",
    }
    changed = 0
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        for prefix, value in captions.items():
            if text.startswith(prefix):
                if text != value:
                    set_paragraph_text(paragraph, value, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                    changed += 1
                break
    return changed


def replace_figure_captions(doc: Document) -> int:
    sequence = [
        "Рисунок 1 - общая схема розлива",
        "Рисунок 2 - функциональная схема процесса паллетизации",
        "Рисунок 3 - графики телеметрии в среде CoppeliaSim",
        "Рисунок 4 - робот ABB IRB 660-180/3.15",
        "Рисунок 5 - декомпозиция нулевого уровня",
        "Рисунок 6 - декомпозиция первого уровня",
        "Рисунок 7 - архитектура ПАК",
        "Рисунок 8 - цифровая модель РТК",
        "Рисунок 9 - сравнение сценариев износа",
        "Рисунок 10 - макет панели мониторинга Grafana",
        "Рисунок 11 - схема интеграции компонентов ПАК",
        "Рисунок 12 - Схема роботизированной ячейки паллетизации, использованная как исходная планировка цифровой модели",
        "Рисунок 13 - Контрольный пример сравнения фактического и прогнозного RUL на этапе обучения модели",
        "Рисунок 14 - Среднеквадратический момент приводов по фазам набора long_live_01",
        "Рисунок 15 - Сравнение фактического и прогнозного RUL для сценария S3, motor1",
        "Рисунок 16 - Изменение индекса HI для motor1 по сценариям деградации",
        "Рисунок 17 - Сводная панель ПАК PdM по результатам апробации",
    ]
    changed = 0
    idx = 0
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text.startswith("Рисунок ") and idx < len(sequence):
            if text != sequence[idx]:
                set_paragraph_text(paragraph, sequence[idx], align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                changed += 1
            idx += 1
    return changed


def split_formula_94_caption(doc: Document) -> bool:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "Таблица 31 - Структура измерений и расчетных таблиц прототипа" in text and "(94)" in text:
            for run in list(paragraph.runs):
                if "Таблица" in run.text or "31" == run.text or "Структура измерений" in run.text or run.text == "\n":
                    run._element.getparent().remove(run._element)
            add_caption_after(paragraph, "Таблица 31 - Структура измерений и расчетных таблиц прототипа")
            return True
    return False


def update_weighted_comparison_table(doc: Document) -> bool:
    for table in doc.tables:
        first = [cell.text.strip() for cell in table.rows[0].cells]
        joined = " ".join(first)
        if "Реактивное обслуживание" in joined and "Предиктивное" in joined:
            clear_and_fill_table(
                table,
                ["Критерий и вес", "Реактивное ТО", "ППР", "Предиктивное ТО"],
                [
                    ["Риск аварийного простоя, w=0,30", "1", "3", "5"],
                    ["Использование ресурса, w=0,25", "1", "3", "5"],
                    ["Простота внедрения, w=0,15", "5", "4", "3"],
                    ["Поддержка RUL-прогноза, w=0,20", "0", "1", "5"],
                    ["Планируемость ремонта, w=0,10", "1", "3", "5"],
                    ["Итоговый балл Qv", "1,65", "2,85", "4,70"],
                ],
                font_size=9.8,
            )
            return True
    return False


def find_appendix_heading(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Приложение"):
            return paragraph
    return None


def add_table_after(paragraph, caption: str, headers: list[str], rows: list[list[str]], *, font_size: float = 9.5):
    caption_p = add_caption_after(paragraph, caption)
    table = paragraph._parent.add_table(rows=1, cols=len(headers), width=Cm(15.5))
    caption_p._p.addnext(table._tbl)
    spacer = paragraph._parent.add_paragraph()
    table._tbl.addnext(spacer._p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    clear_and_fill_table(table, headers, rows, font_size=font_size)
    return spacer


def add_appendix(doc: Document) -> None:
    heading = find_appendix_heading(doc)
    if heading is None:
        heading = doc.add_paragraph("Приложение А. Дополнительные материалы по программной реализации ПАК")
        heading.style = "Heading 3"
    else:
        set_paragraph_text(
            heading,
            "Приложение А. Дополнительные материалы по программной реализации ПАК",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line=False,
        )
        heading.style = "Heading 3"
    heading.paragraph_format.page_break_before = True

    # Remove stale content after the appendix heading if this script is rerun.
    node = heading._p.getnext()
    while node is not None:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node

    p = add_paragraph_after(
        heading,
        "В приложении приведены дополнительные сведения, которые подтверждают воспроизводимость практической части, но перегружали бы основной текст РПЗ: состав программных модулей, контрольные файлы апробации, пример строки телеметрии и команды обработки данных.",
    )

    anchor = add_table_after(
        p,
        "Таблица А.1 - Состав программных модулей ПАК предиктивного обслуживания",
        ["Модуль", "Файл или объект", "Назначение"],
        [
            ["Цифровая сцена", "scenes/vkr_scena.ttt", "Модель роботизированной ячейки паллетизации"],
            ["Сценарий паллетизации", "scripts/coppeliasim/lua/final_scene_palletizing_cycle.lua", "Формирование цикла, фаз и данных состояния"],
            ["Коллектор телеметрии", "scripts/coppeliasim/python/collect_final_scene_telemetry.py", "Запись пакетов CoppeliaSim Remote API в JSONL"],
            ["Нормализация", "scripts/data_pipeline/normalize_telemetry.py", "Преобразование JSONL в табличную телеметрию по осям"],
            ["Контроль качества", "scripts/data_pipeline/validate_telemetry.py", "Проверка полноты данных и фазовой разметки"],
            ["Признаки", "scripts/data_pipeline/build_features.py", "Расчет статистических, энергетических и фазовых признаков"],
            ["Деградация", "scripts/data_pipeline/simulate_degradation.py", "Формирование сценариев S0...S3"],
            ["RUL", "scripts/data_pipeline/estimate_rul.py", "Расчет HI, RUL и уровня риска"],
            ["ML-модель", "scripts/data_pipeline/train_rul_mlp.py", "Обучение MLPRegressor и расчет MAE/RMSE/R2"],
            ["Визуализация", "infra/pak/grafana/dashboards/vkr_pak_dashboard.json", "Панель мониторинга HI/RUL и предупреждений"],
        ],
    )

    anchor = add_table_after(
        anchor,
        "Таблица А.2 - Контрольные артефакты апробации",
        ["Артефакт", "Значение", "Использование в РПЗ"],
        [
            ["data/telemetry/vkr_raw/long_live_01.jsonl", "22174 пакета", "Основной набор телеметрии"],
            ["data/telemetry/vkr_normalized/vkr_telemetry_normalized.csv", "88696 строк", "Проверка полноты данных"],
            ["data/features/vkr_features.csv", "56 строк признаков", "Расчет диагностических признаков"],
            ["data/results/vkr_rul_estimates.csv", "17920 строк", "Базовая оценка HI/RUL"],
            ["data/results/vkr_nn_rul_predictions.csv", "17920 строк", "Нейросетевой прогноз RUL"],
            ["data/results/vkr_nn_rul_metrics.csv", "MAE=1,173; RMSE=1,442; R2=0,994", "Оценка качества модели"],
        ],
    )

    anchor = add_table_after(
        anchor,
        "Таблица А.3 - Пример нормализованной записи телеметрии",
        ["Поле", "Пример значения", "Смысл"],
        [
            ["time", "162,40", "Время симуляции, с"],
            ["run_id", "long_live_01", "Идентификатор прогона"],
            ["phase", "lift_before_pick", "Фаза паллетизационного цикла"],
            ["axis", "motor2", "Наблюдаемый привод"],
            ["q", "0,609417", "Положение оси"],
            ["omega", "-0,011277", "Скорость оси"],
            ["accel", "-0,225548", "Ускорение оси"],
            ["torque", "2699,12", "Момент на приводе"],
        ],
    )

    add_table_after(
        anchor,
        "Таблица А.4 - Команды воспроизведения обработки данных",
        ["Шаг", "Команда", "Результат"],
        [
            ["Запуск полного конвейера", "python scripts/data_pipeline/run_file_pipeline.py", "Нормализация, признаки, деградация, RUL и ML-метрики"],
            ["Построение графиков", "python scripts/data_pipeline/make_vkr_figures.py", "SVG-графики для РПЗ и презентации"],
            ["Экспорт в InfluxDB", "powershell scripts/data_pipeline/run_pipeline_and_export.ps1", "Загрузка временных рядов и расчетных показателей"],
            ["Демонстрационный запуск", "powershell scripts/pak/run_pak_demo.ps1", "Сбор и отображение данных в режиме демонстрации"],
        ],
    )


def main() -> None:
    docx = find_docx()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = docx.with_name(f"{docx.stem}.backup_before_pdf_polish_{timestamp}{docx.suffix}")
    shutil.copy2(docx, backup)

    doc = Document(docx)
    text_changes = edit_all_text(doc)
    caption_changes = replace_captions(doc)
    figure_changes = replace_figure_captions(doc)
    split_formula = split_formula_94_caption(doc)
    weighted_table = update_weighted_comparison_table(doc)
    add_appendix(doc)
    doc.save(docx)

    print(f"docx\t{docx}")
    print(f"backup\t{backup}")
    print(f"text_changes\t{text_changes}")
    print(f"caption_changes\t{caption_changes}")
    print(f"figure_changes\t{figure_changes}")
    print(f"split_formula_94_caption\t{int(split_formula)}")
    print(f"weighted_table_updated\t{int(weighted_table)}")


if __name__ == "__main__":
    main()
