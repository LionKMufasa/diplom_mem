from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

from compress_vkr_filled_sections import (
    clear_paragraph,
    find_index,
    format_body,
    insert_paragraph_after,
    replace_section,
    set_num_pr,
    set_run_font,
    set_update_fields_on_open,
)
from fix_vkr_citations_bibliography_formulas_tz import BIBLIOGRAPHY_ENTRIES, improve_formulas


CITATION_RE = re.compile(r"\[(\d+),\s*с\.\s*(\d+)\]")


def set_plain_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, 14)


def shift_old_citations(doc: Document) -> tuple[int, int]:
    bibliography_idx = find_index(doc, "Список литературы")
    changed = 0
    old_source6_mentions = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal changed, old_source6_mentions
        num = int(match.group(1))
        page = match.group(2)
        if num == 6:
            old_source6_mentions += 1
            return match.group(0)
        if num > 6:
            changed += 1
            return f"[{num - 1}, с. {page}]"
        return match.group(0)

    for paragraph in doc.paragraphs[:bibliography_idx]:
        text = paragraph.text or ""
        if "[" not in text:
            continue
        new_text = CITATION_RE.sub(repl, text)
        if new_text != text:
            set_plain_text(paragraph, new_text)
            if not paragraph.style.name.startswith("Heading"):
                format_body(paragraph)
    return changed, old_source6_mentions


def rebuild_bibliography_without_source6(doc: Document) -> int:
    entries = [entry for entry in BIBLIOGRAPHY_ENTRIES if not entry.startswith("Галахарь А.С.")]
    start_idx = find_index(doc, "Список литературы")
    end_idx = find_index(doc, "Приложения")
    anchor = doc.paragraphs[start_idx]
    end = doc.paragraphs[end_idx]
    parent = anchor._element.getparent()
    start_pos = parent.index(anchor._element)
    end_pos = parent.index(end._element)
    for element in list(parent)[start_pos + 1 : end_pos]:
        parent.remove(element)

    current = anchor
    for entry in entries:
        current = insert_paragraph_after(current, entry, "Normal")
        set_num_pr(current, "6", "0")
        format_body(current, first_line=False)
    return len(entries)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: insert_vkr_chapter5_remove_source6.py <docx> <chapter5.md>")

    docx_path = Path(sys.argv[1])
    chapter5_path = Path(sys.argv[2])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = docx_path.with_name(
        f"{docx_path.stem}.backup_before_remove_source6_and_chapter5_{timestamp}{docx_path.suffix}"
    )
    shutil.copy2(docx_path, backup)

    doc = Document(str(docx_path))
    shifted, old_source6_mentions = shift_old_citations(doc)
    replace_section(doc, "Рабочее проектирование", "Апробация и оценка эффективности системы", chapter5_path, 1, True)
    doc.paragraphs[find_index(doc, "Рабочее проектирование")].paragraph_format.page_break_before = True
    bibliography_count = rebuild_bibliography_without_source6(doc)
    formula_changes = improve_formulas(doc)
    doc.save(str(docx_path))
    set_update_fields_on_open(docx_path)

    print(f"Updated VKR RPZ: {docx_path}")
    print(f"Backup: {backup}")
    print(f"Shifted citation mentions: {shifted}")
    print(f"Old source 6 citation mentions left for manual review: {old_source6_mentions}")
    print(f"Bibliography entries: {bibliography_count}")
    print(f"Formula replacements/format pass: {formula_changes}")


if __name__ == "__main__":
    main()
