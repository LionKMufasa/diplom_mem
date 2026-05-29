from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"

FONT_NAME = "Times New Roman"
HEADING_SIZE_PT = 14


def set_run_font(run, *, size_pt: int = HEADING_SIZE_PT, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def set_style_font(style, *, size_pt: int = HEADING_SIZE_PT, bold: bool = True) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), str(size_pt * 2))
    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(size_pt * 2))
    if bold and r_pr.find(qn("w:b")) is None:
        r_pr.append(OxmlElement("w:b"))
    if bold and r_pr.find(qn("w:bCs")) is None:
        r_pr.append(OxmlElement("w:bCs"))


def is_heading(paragraph) -> bool:
    if paragraph.style is None:
        return False
    style_id = getattr(paragraph.style, "style_id", "") or ""
    style_name = getattr(paragraph.style, "name", "") or ""
    return style_id.startswith("Heading") or style_name.startswith("Heading") or style_name.startswith("Заголовок")


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def find_bibliography_start(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if "Список" in text and "использованных источников" in text:
            return idx
    raise RuntimeError("Bibliography heading not found")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def renumber_citations_in_run_text(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        number = int(match.group(1))
        if number == 5:
            return "[4]"
        if number > 5:
            return f"[{number - 1}]"
        return match.group(0)

    return re.sub(r"\[(\d+)\]", repl, text)


def clear_part_to_single_empty_paragraph(part) -> Paragraph:
    element = part._element
    for child in list(element):
        element.remove(child)
    p = OxmlElement("w:p")
    element.append(p)
    return Paragraph(p, part)


def add_page_field(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size_pt=HEADING_SIZE_PT, bold=False)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    cached = OxmlElement("w:t")
    cached.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(cached)
    run._r.append(end)


def normalize_headers_footers(doc: Document) -> int:
    touched = 0
    for idx, section in enumerate(doc.sections):
        section.different_first_page_header_footer = idx == 0

        for header in (section.header, section.first_page_header, section.even_page_header):
            clear_part_to_single_empty_paragraph(header)
            touched += 1

        first_footer = clear_part_to_single_empty_paragraph(section.first_page_footer)
        touched += 1
        if idx != 0:
            add_page_field(first_footer)

        for footer in (section.footer, section.even_page_footer):
            p = clear_part_to_single_empty_paragraph(footer)
            add_page_field(p)
            touched += 1
    return touched


def apply_docx_edits() -> dict[str, int | str]:
    doc = Document(DOCX_PATH)
    report: dict[str, int | str] = {}

    heading3 = doc.styles["Heading 3"]
    for style_name in (
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
        "Title",
        "Subtitle",
    ):
        try:
            set_style_font(doc.styles[style_name], bold=True)
        except Exception:
            pass

    headings_changed = 0
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            paragraph.style = heading3
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                if run.text:
                    set_run_font(run, bold=True)
            headings_changed += 1
    report["headings_set_to_heading3"] = headings_changed

    bib_start = find_bibliography_start(doc)
    report["bibliography_start_before_delete"] = bib_start

    refs_5_before = 0
    refs_shifted = 0
    for paragraph in doc.paragraphs[:bib_start]:
        for run in paragraph.runs:
            if "[5]" in run.text:
                refs_5_before += run.text.count("[5]")
            new_text = renumber_citations_in_run_text(run.text)
            if new_text != run.text:
                refs_shifted += 1
                run.text = new_text
                set_run_font(run, size_pt=HEADING_SIZE_PT if is_heading(paragraph) else 14, bold=True if is_heading(paragraph) else None)
    report["refs_5_before"] = refs_5_before
    report["citation_runs_shifted"] = refs_shifted

    deleted_source5 = 0
    # Recompute after citation edits; source deletion changes paragraph indices.
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("ГОСТ 34.602") and "89" in text and "2020" not in text:
            delete_paragraph(paragraph)
            deleted_source5 += 1
            break
    report["deleted_source5_gost_1989"] = deleted_source5

    report["header_footer_parts_normalized"] = normalize_headers_footers(doc)

    doc.save(DOCX_PATH)
    return report


def ooxml_scrub_heading_sizes() -> dict[str, int]:
    """Force heading-style and heading-paragraph run sizes to 14 pt in raw OOXML."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tmp = DOCX_PATH.with_suffix(".tmp.docx")
    report = {"style_sizes_fixed": 0, "heading_run_sizes_fixed": 0}

    with zipfile.ZipFile(DOCX_PATH, "r") as zin:
        styles_root = etree.fromstring(zin.read("word/styles.xml"))
        for style in styles_root.xpath(".//w:style", namespaces=ns):
            style_id = style.get(qn("w:styleId"), "")
            name_el = style.find("w:name", namespaces=ns)
            name_val = name_el.get(qn("w:val"), "") if name_el is not None else ""
            if style_id.startswith("Heading") or name_val.startswith("heading") or name_val.startswith("Heading") or style_id in {
                "1",
                "2",
                "3",
                "Title",
                "Subtitle",
            }:
                r_pr = style.find("w:rPr", namespaces=ns)
                if r_pr is None:
                    r_pr = OxmlElement("w:rPr")
                    style.append(r_pr)
                for tag in ("w:sz", "w:szCs"):
                    el = r_pr.find(tag, namespaces=ns)
                    if el is None:
                        el = OxmlElement(tag)
                        r_pr.append(el)
                    if el.get(qn("w:val")) != "28":
                        el.set(qn("w:val"), "28")
                        report["style_sizes_fixed"] += 1

        document_root = etree.fromstring(zin.read("word/document.xml"))
        for paragraph in document_root.xpath(".//w:p", namespaces=ns):
            p_style = paragraph.find("w:pPr/w:pStyle", namespaces=ns)
            p_style_val = p_style.get(qn("w:val"), "") if p_style is not None else ""
            if p_style_val not in {"Heading1", "Heading2", "Heading3", "Heading4", "Heading5", "Heading6", "Heading7", "Heading8", "Heading9", "1", "2", "3"}:
                continue
            for run in paragraph.xpath(".//w:r", namespaces=ns):
                r_pr = run.find("w:rPr", namespaces=ns)
                if r_pr is None:
                    r_pr = OxmlElement("w:rPr")
                    run.insert(0, r_pr)
                for tag in ("w:sz", "w:szCs"):
                    el = r_pr.find(tag, namespaces=ns)
                    if el is None:
                        el = OxmlElement(tag)
                        r_pr.append(el)
                    if el.get(qn("w:val")) != "28":
                        el.set(qn("w:val"), "28")
                        report["heading_run_sizes_fixed"] += 1
                if r_pr.find("w:b", namespaces=ns) is None:
                    r_pr.append(OxmlElement("w:b"))
                if r_pr.find("w:bCs", namespaces=ns) is None:
                    r_pr.append(OxmlElement("w:bCs"))

        new_styles = etree.tostring(styles_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        new_document = etree.tostring(document_root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/styles.xml":
                    data = new_styles
                elif item.filename == "word/document.xml":
                    data = new_document
                zout.writestr(item, data)
    tmp.replace(DOCX_PATH)
    return report


def audit_docx() -> dict[str, int | str | list[int]]:
    doc = Document(DOCX_PATH)
    report: dict[str, int | str | list[int]] = {}
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        report["zip_bad"] = str(zf.testzip())

    heading_bad = 0
    heading_styles = {}
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            heading_styles[getattr(paragraph.style, "name", "")] = heading_styles.get(getattr(paragraph.style, "name", ""), 0) + 1
            if getattr(paragraph.style, "name", "") != "Heading 3":
                heading_bad += 1
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.size is not None and round(run.font.size.pt, 1) > 18:
                        heading_bad += 1
    report["heading_bad_style_or_size"] = heading_bad
    report["heading_styles"] = str(heading_styles)

    bib_start = find_bibliography_start(doc)
    refs = []
    for paragraph in doc.paragraphs[:bib_start]:
        refs.extend(int(x) for x in re.findall(r"\[(\d+)\]", paragraph.text))
    report["has_ref_5"] = int(5 in refs)
    report["max_ref"] = max(refs) if refs else 0
    report["source5_gost_1989_left"] = int(any(p.text.strip().startswith("ГОСТ 34.602") and "89" in p.text for p in doc.paragraphs[bib_start + 1 :]))

    headers_ok = 1
    footers_ok = 1
    for idx, section in enumerate(doc.sections):
        for header in (section.header, section.first_page_header, section.even_page_header):
            if len(header.paragraphs) != 1 or header.paragraphs[0].text != "":
                headers_ok = 0
        footer_parts = [
            (section.first_page_footer, idx == 0),
            (section.footer, False),
            (section.even_page_footer, False),
        ]
        for footer, should_be_empty in footer_parts:
            if len(footer.paragraphs) != 1:
                footers_ok = 0
                continue
            if should_be_empty:
                if footer.paragraphs[0].text != "":
                    footers_ok = 0
            else:
                # Field-coded page numbers are not visible in python-docx text.
                if footer.paragraphs[0].alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    footers_ok = 0
    report["headers_ok"] = headers_ok
    report["footers_ok"] = footers_ok
    report["sections"] = len(doc.sections)
    report["paragraphs"] = len(doc.paragraphs)
    report["tables"] = len(doc.tables)
    return report


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}.backup_before_normcontrol_headers_footers_sources_{datetime.now():%Y%m%d_%H%M%S}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)
    report = apply_docx_edits()
    scrub = ooxml_scrub_heading_sizes()
    audit = audit_docx()
    print(f"backup={backup}")
    print("report=" + repr(report))
    print("ooxml_scrub=" + repr(scrub))
    print("audit=" + repr(audit))


if __name__ == "__main__":
    main()
