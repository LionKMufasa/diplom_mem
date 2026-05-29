from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    with zipfile.ZipFile(DOCX) as zf:
        bad_member = zf.testzip()
    doc = Document(DOCX)
    paragraph_texts = [" ".join(p.text.split()) for p in doc.paragraphs]
    flat = "\n".join(paragraph_texts)
    appendix_headings = [
        text for text in paragraph_texts if re.fullmatch(r"Приложение [А-Я]", text)
    ]
    appendix_captions = [
        text for text in paragraph_texts if re.match(r"^Таблица [А-Я]\.\d+", text)
    ]
    expected_headings = ["Приложение А", "Приложение Б", "Приложение В", "Приложение Г"]
    expected_captions = [
        "Таблица А.1",
        "Таблица А.2",
        "Таблица А.3",
        "Таблица Б.1",
        "Таблица Б.2",
        "Таблица Б.3",
        "Таблица В.1",
        "Таблица В.2",
        "Таблица В.3",
        "Таблица Г.1",
        "Таблица Г.2",
        "Таблица Г.3",
    ]
    missing_headings = [item for item in expected_headings if item not in appendix_headings]
    missing_captions = [
        item for item in expected_captions if not any(text.startswith(item) for text in appendix_captions)
    ]
    stale_tokens = [
        token
        for token in [
            "customData.predictiveTelemetry.queue",
            "следующий практический шаг",
            "final_scena_diplom.ttt",
            "pred_final.ttt",
            "Будущая вставка",
            "Будут вставлены после финальных прогонов",
        ]
        if token in flat
    ]
    reference_markers = [
        "приложении А",
        "приложении Б",
        "приложении В",
        "приложении Г",
        "приложения А-В",
        "Приложения А-Г",
    ]
    present_references = [marker for marker in reference_markers if marker in flat]
    print(f"docx={DOCX}")
    print(f"zip_ok={bad_member is None}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"appendix_headings={appendix_headings}")
    print(f"missing_headings={missing_headings}")
    print(f"appendix_captions={appendix_captions}")
    print(f"missing_captions={missing_captions}")
    print(f"stale_tokens={stale_tokens}")
    print(f"present_references={present_references}")


if __name__ == "__main__":
    main()
