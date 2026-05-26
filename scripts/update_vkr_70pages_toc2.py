from __future__ import annotations

import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


ANNOTATION_TEXT = (
    "Пояснительная записка к выпускной квалификационной работе "
    "«Разработка системы предиктивного обслуживания узлов робота-паллетизатора "
    "на участке розлива продукции ООО «Компания “Здоровая жизнь”» "
    "формируется в объеме 70 листов основной части формата А4. Материалы, "
    "дополняющие основное содержание работы, включая крупные схемы, листинги, "
    "дополнительные таблицы, графики и иллюстрации, выносятся в приложения и "
    "в указанный объем основной части не включаются. В записке рассматриваются "
    "предпроектное обследование роботизированного участка паллетизации, "
    "проектирование программно-аппаратного комплекса сбора и анализа телеметрии, "
    "разработка цифровой модели робота, формирование модели деградации узлов и "
    "оценка эффективности предиктивного обслуживания."
)


def set_run_font(run, size: int = 14) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs)[::-1]:
        paragraph._p.remove(run._r)


def format_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(35.4)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 14)


def update_annotation(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for idx, paragraph in enumerate(doc.paragraphs):
        if (paragraph.text or "").strip() == "Аннотация":
            target = doc.paragraphs[idx + 1]
            clear_paragraph(target)
            run = target.add_run(ANNOTATION_TEXT)
            set_run_font(run, 14)
            format_body(target)
            doc.save(str(docx_path))
            return
    raise RuntimeError("Could not find annotation heading")


def update_toc_field_to_two_levels(docx_path: Path) -> None:
    tmp_docx = docx_path.with_name(f"{docx_path.stem}.tmp_toc2.docx")
    parser = etree.XMLParser(remove_blank_text=False)

    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp_docx, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        names = set(src.namelist())
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data, parser)
                for instr in root.findall(".//w:instrText", NS):
                    text = instr.text or ""
                    if "TOC" in text:
                        instr.text = text.replace('\\o "1-3"', '\\o "1-2"')
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            elif info.filename == "word/settings.xml":
                root = etree.fromstring(data, parser)
                update_fields = root.find("w:updateFields", NS)
                if update_fields is None:
                    update_fields = etree.Element(f"{{{W_NS}}}updateFields")
                    root.insert(0, update_fields)
                update_fields.set(f"{{{W_NS}}}val", "true")
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            dst.writestr(info, data)

        if "word/settings.xml" not in names:
            root = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            update_fields = etree.SubElement(root, f"{{{W_NS}}}updateFields")
            update_fields.set(f"{{{W_NS}}}val", "true")
            dst.writestr(
                "word/settings.xml",
                etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                ),
            )

    shutil.move(str(tmp_docx), str(docx_path))


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: update_vkr_70pages_toc2.py <docx>")
    docx_path = Path(sys.argv[1])
    update_annotation(docx_path)
    update_toc_field_to_two_levels(docx_path)
    print(f"Updated annotation and TOC field in {docx_path}")


if __name__ == "__main__":
    main()
