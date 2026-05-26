from __future__ import annotations

import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


BIBLIOGRAPHY = [
    "ГОСТ 27.002-2015. Надежность в технике. Термины и определения. М.: Стандартинформ, 2016. 28 с.",
    "ГОСТ 27.003-2016. Надежность в технике. Состав и общие правила задания требований по надежности. М.: Стандартинформ, 2018. 23 с.",
    "ГОСТ 27.301-95. Надежность в технике. Расчет надежности. Основные положения. Минск: Межгосударственный совет по стандартизации, метрологии и сертификации, 1995. 19 с.",
    "ГОСТ 34.602-2020. Информационные технологии. Комплекс стандартов на автоматизированные системы. Техническое задание на создание автоматизированной системы. М.: Российский институт стандартизации, 2021. 12 с.",
    "ABB Robotics. Product manual - IRB 660. ABB Robotics, 2026. 368 p.",
    "ABB Robotics. IRB 660: Taking palletizing to new heights. Datasheet. ABB Robotics, 2018. 2 p.",
    "Kang Z., Catal C., Tekinerdogan B. Remaining Useful Life (RUL) Prediction of Equipment in Production Lines Using Artificial Neural Networks // Sensors. 2021. Vol. 21, No. 3. Article 932. DOI: 10.3390/s21030932.",
    "Taşcı B., Omar A., Ayvaz S. Remaining useful lifetime prediction for predictive maintenance in manufacturing // Computers & Industrial Engineering. 2023. Vol. 184. Article 109566. DOI: 10.1016/j.cie.2023.109566.",
    "Baur M., Albertelli P., Monno M. A review of prognostics and health management of machine tools // The International Journal of Advanced Manufacturing Technology. 2020. DOI: 10.1007/s00170-020-05202-3.",
    "Gharib H., Kovács G. A Review of Prognostic and Health Management (PHM) Methods and Limitations for Marine Diesel Engines: New Research Directions // Machines. 2023. Vol. 11, No. 7. Article 695. DOI: 10.3390/machines11070695.",
    "Liu Y., Wen J., Wang G. A comprehensive overview of remaining useful life prediction: From traditional literature review to scientometric analysis // Machine Learning with Applications. 2025. Vol. 21. Article 100704. DOI: 10.1016/j.mlwa.2025.100704.",
    "Kumar S., Raj K.K., Cirrincione M., Cirrincione G., Franzitta V., Kumar R.R. A Comprehensive Review of Remaining Useful Life Estimation Approaches for Rotating Machinery // Energies. 2024. Vol. 17, No. 22. Article 5538. DOI: 10.3390/en17225538.",
    "Kritzinger W., Karner M., Traar G., Henjes J., Sihn W. Digital Twin in manufacturing: A categorical literature review and classification // IFAC-PapersOnLine. 2018. Vol. 51, No. 11. P. 1016-1022. DOI: 10.1016/j.ifacol.2018.08.474.",
    "Fuller A., Fan Z., Day C., Barlow C. Digital Twin: Enabling Technologies, Challenges and Open Research // IEEE Access. 2020. Vol. 8. P. 108952-108971. DOI: 10.1109/ACCESS.2020.2998358.",
    "Sharma A., Kosasih E., Zhang J., Brintrup A., Calinescu A. Digital Twins: State of the art theory and practice, challenges, and open research questions // Journal of Industrial Information Integration. 2022. Vol. 30. Article 100383. DOI: 10.1016/j.jii.2022.100383.",
    "Soori M., Arezoo B., Dastres R. Digital twin for smart manufacturing, A review // Sustainable Manufacturing and Service Economics. 2023. Vol. 2. Article 100017. DOI: 10.1016/j.smse.2023.100017.",
    "Zhang L., Chen X., Zhou W., Cheng T., Chen L., Guo Z., Han B., Lu L. Digital Twins for Additive Manufacturing: A State-of-the-Art Review // Applied Sciences. 2020. Vol. 10, No. 23. Article 8350. DOI: 10.3390/app10238350.",
    "Coppelia Robotics. CoppeliaSim User Manual [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/ (дата обращения: 05.05.2026).",
    "Coppelia Robotics. ZeroMQ Remote API [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/en/zmqRemoteApiOverview.htm (дата обращения: 05.05.2026).",
    "Coppelia Robotics. Remote API overview [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/en/remoteApiOverview.htm (дата обращения: 05.05.2026).",
    "CoppeliaSim ZMQ Remote API client. Python package [Электронный ресурс]. URL: https://pypi.org/project/coppeliasim-zmqremoteapi-client/ (дата обращения: 05.05.2026).",
    "Scikit-learn developers. Scikit-learn User Guide [Электронный ресурс]. URL: https://scikit-learn.org/stable/user_guide.html (дата обращения: 05.05.2026).",
    "Scikit-learn developers. Supervised learning [Электронный ресурс]. URL: https://scikit-learn.org/stable/supervised_learning.html (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Documentation [Электронный ресурс]. URL: https://xgboost.readthedocs.io/ (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Parameters [Электронный ресурс]. URL: https://xgboost.readthedocs.io/en/stable/parameter.html (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Python API Reference [Электронный ресурс]. URL: https://xgboost.readthedocs.io/en/stable/python/python_api.html (дата обращения: 05.05.2026).",
    "InfluxData. InfluxDB documentation [Электронный ресурс]. URL: https://docs.influxdata.com/ (дата обращения: 05.05.2026).",
    "InfluxData. InfluxDB key concepts v2 [Электронный ресурс]. URL: https://docs.influxdata.com/influxdb/v2/reference/key-concepts/ (дата обращения: 05.05.2026).",
    "InfluxData. Time Series Database explained [Электронный ресурс]. URL: https://www.influxdata.com/time-series-database/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Grafana documentation [Электронный ресурс]. URL: https://grafana.com/docs/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Dashboards documentation [Электронный ресурс]. URL: https://grafana.com/docs/grafana/latest/visualizations/dashboards/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Alerting documentation [Электронный ресурс]. URL: https://grafana.com/docs/grafana/latest/alerting/ (дата обращения: 05.05.2026).",
]


ABBREVIATIONS = [
    "API - Application Programming Interface, программный интерфейс приложения;",
    "HI - Health Index, интегральный показатель технического состояния узла;",
    "ML - Machine Learning, машинное обучение;",
    "MQTT - Message Queuing Telemetry Transport, протокол обмена телеметрическими сообщениями;",
    "PdM - Predictive Maintenance, предиктивное обслуживание;",
    "PHM - Prognostics and Health Management, прогнозирование и управление техническим состоянием;",
    "RUL - Remaining Useful Life, остаточный полезный ресурс;",
    "БД - база данных;",
    "ВКР - выпускная квалификационная работа;",
    "ГОСТ - государственный стандарт;",
    "ПАК - программно-аппаратный комплекс;",
    "ППР - планово-предупредительный ремонт;",
    "ТОиР - техническое обслуживание и ремонт;",
    "ЦД - цифровой двойник.",
]


INTRO_PARAGRAPHS = [
    "Современные производственные предприятия все в большей степени зависят от устойчивой работы автоматизированных технологических линий. В условиях роста требований к производительности, прослеживаемости и снижению себестоимости продукции простои оборудования становятся не только технической, но и экономической проблемой. Для участков с последовательной структурой технологического процесса отказ одного ключевого элемента способен остановить выпуск продукции на всей линии, поэтому повышение эксплуатационной надежности оборудования является одной из центральных задач инженерного проектирования.",
    "Традиционная организация технического обслуживания часто строится на планово-предупредительном подходе: обслуживание выполняется через заранее заданные интервалы времени или наработки. Такой подход прост в администрировании, однако не всегда отражает фактическое состояние оборудования. При недостаточно частом обслуживании возрастает риск внезапного отказа, а при чрезмерно частом - увеличиваются затраты на ремонтные операции, запасные части и вынужденные остановы. В терминах надежности, закрепленных в ГОСТ 27.002-2015, техническое состояние и безотказность должны рассматриваться как свойства объекта, изменяющиеся во времени под воздействием эксплуатационных факторов [1].",
    "Одним из направлений развития систем технического обслуживания является переход к обслуживанию по состоянию и предиктивному обслуживанию. Предиктивное обслуживание (Predictive Maintenance, PdM) предполагает сбор диагностических данных, оценку текущего состояния узлов и прогнозирование остаточного полезного ресурса. В отличие от обслуживания по календарному графику, такой подход позволяет принимать решение о ремонте на основании измеряемых признаков деградации и расчетного прогноза, что особенно важно для оборудования с высокой стоимостью простоя.",
    "Объектом исследования в выпускной квалификационной работе является роботизированный участок паллетизации продукции линии розлива ООО «Компания \"Здоровая жизнь\"». На данном участке промышленный робот-паллетизатор выполняет операции перемещения и укладки групповой упаковки на паллету. Паллетизация относится к завершающим операциям производственной линии: сбой робота приводит к нарушению отгрузочного ритма и накоплению продукции перед зоной укладки. Поэтому робот-паллетизатор рассматривается как критический элемент производственной системы.",
    "В качестве технической основы рассматривается промышленный робот ABB IRB 660-180/3.15. Семейство ABB IRB 660 относится к специализированным паллетизирующим роботам с четырехосевой кинематической схемой, большим рабочим радиусом и высокой грузоподъемностью, что подтверждается эксплуатационной документацией и техническими материалами производителя [5, 6]. Для задач предиктивного обслуживания наибольший интерес представляют механические узлы привода: редукторы, подшипниковые опоры, сочленения и исполнительные механизмы осей, поскольку их деградация проявляется в изменении моментов, скоростей, вибрационных признаков, токовых нагрузок и температурных режимов.",
    "Предметом исследования являются методы построения системы предиктивного обслуживания узлов робота-паллетизатора на основе цифрового моделирования, сбора телеметрии и прогнозирования остаточного ресурса. В работе рассматривается не только алгоритм прогнозирования, но и связанная инженерная инфраструктура: модель робота в среде CoppeliaSim, контур получения телеметрии, хранение временных рядов, визуализация диагностических параметров и формирование признаков для модели деградации.",
    "Использование цифровой модели позволяет исследовать поведение роботизированной ячейки без вмешательства в работу реальной производственной линии. В литературе по цифровым двойникам различают цифровую модель, цифровую тень и цифровой двойник в зависимости от степени двусторонней связи с физическим объектом [13]. Для данной работы цифровая модель в CoppeliaSim применяется как инструмент воспроизведения паллетизационного цикла, генерации телеметрии и проверки логики диагностического контура. Такой подход согласуется с современными исследованиями цифровых двойников в производстве, где виртуальное представление оборудования используется для анализа сценариев, диагностики и поддержки принятия решений [14-16].",
    "Ключевой расчетной задачей системы является прогнозирование остаточного полезного ресурса (Remaining Useful Life, RUL). Обзоры по PHM и RUL показывают, что для сложных промышленных объектов применяются модельно-ориентированные, статистические, data-driven и гибридные подходы [10-12]. Для производственных линий и вращающихся машин перспективными являются методы машинного обучения, которые позволяют восстанавливать зависимость между телеметрическими признаками и стадией деградации по накопленным данным [7-9]. В рамках ВКР такой подход используется для построения экспериментальной модели оценки состояния узлов робота.",
    "Практическая реализация системы требует организации потока данных от модели робота к программным компонентам анализа. Среда CoppeliaSim предоставляет возможности моделирования робототехнических систем и взаимодействия с внешними приложениями через Remote API, включая ZeroMQ Remote API [18-20]. Это позволяет вынести сбор, обработку и анализ телеметрии в Python-клиент, а данные представить как временные ряды для последующего хранения и визуализации.",
    "Для хранения телеметрии целесообразно использовать специализированную базу временных рядов, поскольку диагностические параметры оборудования имеют естественную привязку ко времени, циклу и режиму работы. В качестве технологической основы в работе рассматривается InfluxDB, ориентированная на хранение измерений, полей, тегов и временных меток [27, 28]. Визуализация состояния оборудования и построение панелей мониторинга выполняются с использованием Grafana, которая предоставляет средства отображения временных рядов и настройки диагностических панелей [30, 31].",
    "Целью выпускной квалификационной работы является разработка и обоснование системы предиктивного обслуживания узлов робота-паллетизатора на основе цифрового моделирования, анализа телеметрических данных и прогнозирования остаточного ресурса.",
    "Для достижения поставленной цели необходимо решить следующие задачи:",
]


INTRO_TASKS = [
    "провести предпроектное обследование производственной линии розлива и определить роль робота-паллетизатора в технологическом процессе;",
    "проанализировать объект исследования, критические узлы робота и характерные механизмы их деградации;",
    "выполнить функциональную декомпозицию системы предиктивного обслуживания и определить потоки данных между ее компонентами;",
    "сформировать требования к программно-аппаратному комплексу сбора, хранения, обработки и визуализации телеметрии;",
    "разработать цифровую модель роботизированной ячейки в среде CoppeliaSim и реализовать цикл паллетизации;",
    "организовать сбор телеметрических параметров по приводам робота и подготовку набора данных для анализа деградации;",
    "разработать модель оценки технического состояния и прогнозирования остаточного ресурса узлов робота;",
    "провести апробацию разработанного решения, оценить влияние PdM на надежность и выполнить экономическое обоснование внедрения.",
]


INTRO_TAIL = [
    "Методическую основу работы составляют системный анализ производственного процесса, функциональная декомпозиция, имитационное моделирование робототехнической системы, методы обработки временных рядов, элементы теории надежности и методы машинного обучения для решения регрессионной задачи прогнозирования RUL.",
    "Научная новизна работы заключается в разработке интегрированного подхода к предиктивному обслуживанию робота-паллетизатора, объединяющего цифровую модель паллетизационного цикла, механизм формирования диагностической телеметрии, расчет показателя технического состояния и прогнозирование остаточного ресурса в рамках единого программно-аппаратного комплекса.",
    "Практическая значимость работы состоит в возможности применения предложенного подхода для снижения риска внеплановых простоев роботизированного участка, повышения прозрачности технического состояния узлов робота и обоснования перехода от планово-предупредительного обслуживания к обслуживанию по фактическому состоянию оборудования.",
    "Структура выпускной квалификационной работы включает введение, предпроектное обследование объекта, концептуальное проектирование системы, техническое задание, техническое и рабочее проектирование, апробацию и оценку эффективности, заключение, список литературы и приложения.",
]


STRUCTURE = [
    ("chapter", "Предпроектное обследование"),
    ("section", "Общая характеристика производственной линии"),
    ("section", "Анализ технологического процесса"),
    ("section", "Функциональная декомпозиция процесса"),
    ("section", "Анализ объекта исследования"),
    ("section", "Анализ отказов и деградации узлов"),
    ("section", "Обзор методов обслуживания оборудования"),
    ("section", "Обоснование выбора подхода PdM"),
    ("section", "Выводы по главе"),
    ("chapter", "Концептуальное проектирование"),
    ("section", "Цель и задачи проектирования"),
    ("section", "Общая структура системы"),
    ("section", "Функциональная модель системы"),
    ("section", "Потоки данных и взаимодействие компонентов"),
    ("section", "Анализ архитектуры системы"),
    ("section", "Выбор технологической реализации"),
    ("section", "Формирование требований к системе"),
    ("section", "Выводы по главе"),
    ("chapter", "Техническое задание"),
    ("section", "Общие сведения"),
    ("section", "Назначение и цели создания системы"),
    ("section", "Характеристика объекта автоматизации"),
    ("section", "Требования к функциям системы"),
    ("section", "Требования к видам обеспечения"),
    ("section", "Состав и содержание работ по созданию системы"),
    ("section", "Порядок контроля и приемки"),
    ("section", "Выводы по главе"),
    ("chapter", "Техническое проектирование"),
    ("section", "Архитектура программно-аппаратного комплекса"),
    ("section", "Проект цифровой модели робота"),
    ("section", "Моделирование деградации узлов"),
    ("section", "Подсистема сбора телеметрии"),
    ("section", "Подсистема предобработки данных"),
    ("section", "Математическая модель деградации"),
    ("section", "Формализация задачи прогнозирования RUL"),
    ("section", "Выбор и обоснование алгоритма машинного обучения"),
    ("section", "Подсистема хранения данных"),
    ("section", "Подсистема визуализации"),
    ("section", "Контейнеризация системы"),
    ("section", "Выводы по главе"),
    ("chapter", "Рабочее проектирование"),
    ("section", "Назначение и задачи рабочего проектирования"),
    ("section", "Реализация цифровой модели в CoppeliaSim"),
    ("section", "Реализация механизма деградации узлов"),
    ("section", "Реализация сбора телеметрии"),
    ("section", "Реализация предобработки и формирования признаков"),
    ("section", "Реализация алгоритма прогнозирования RUL"),
    ("section", "Реализация базы данных"),
    ("section", "Реализация интерфейса оператора"),
    ("section", "Интеграция компонентов системы"),
    ("section", "Выводы по главе"),
    ("chapter", "Апробация и оценка эффективности системы"),
    ("section", "Цель и методика апробации"),
    ("section", "Постановка эксперимента"),
    ("section", "Результаты моделирования"),
    ("section", "Оценка качества прогноза RUL"),
    ("section", "Анализ работы системы мониторинга"),
    ("section", "Оценка надежности системы"),
    ("section", "Экономическая оценка"),
    ("section", "Сравнительный анализ вариантов"),
    ("section", "Выводы по главе"),
]


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, text: str, start: int = 0) -> Paragraph:
    for paragraph in doc.paragraphs[start:]:
        if (paragraph.text or "").strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text!r}")


def find_index(doc: Document, text: str, start: int = 0) -> int:
    for idx, paragraph in enumerate(doc.paragraphs[start:], start):
        if (paragraph.text or "").strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text!r}")


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs)[::-1]:
        paragraph._p.remove(run._r)


def set_run_font(run, size: int = 14, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_body(paragraph: Paragraph, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Pt(35.4)
    for run in paragraph.runs:
        set_run_font(run, 14)


def format_heading(paragraph: Paragraph, level: int, numbered: bool = False) -> None:
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 14, bold=True)
    if numbered:
        set_num_pr(paragraph, "9", str(level - 1))


def set_num_pr(paragraph: Paragraph, num_id: str, ilvl: str = "0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def remove_num_pr(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)


def add_toc_field(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Содержание обновится при открытии документа в Word"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, 14)


def set_update_fields_on_open(docx_path: Path) -> None:
    tmp_docx = docx_path.with_name(f"{docx_path.stem}.tmp_update_fields.docx")
    parser = etree.XMLParser(remove_blank_text=False)

    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp_docx, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        names = set(src.namelist())
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "word/settings.xml":
                tree = etree.ElementTree(etree.fromstring(data, parser))
                root = tree.getroot()
                update_fields = root.find("w:updateFields", namespaces=NS)
                if update_fields is None:
                    update_fields = etree.Element(f"{{{W_NS}}}updateFields")
                    root.insert(0, update_fields)
                update_fields.set(f"{{{W_NS}}}val", "true")
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            dst.writestr(info, data)

        if "word/settings.xml" not in names:
            root = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            update_fields = etree.SubElement(root, f"{{{W_NS}}}updateFields")
            update_fields.set(f"{{{W_NS}}}val", "true")
            dst.writestr(
                "word/settings.xml",
                etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                ),
            )

    shutil.move(str(tmp_docx), str(docx_path))


def replace_between(doc: Document, start_text: str, end_text: str, new_items: list[tuple[str, str | None]]) -> None:
    start_idx = find_index(doc, start_text)
    end_idx = find_index(doc, end_text, start_idx + 1)
    for paragraph in list(doc.paragraphs[start_idx + 1 : end_idx]):
        delete_paragraph(paragraph)
    anchor = find_paragraph(doc, start_text)
    for text, style in reversed(new_items):
        insert_paragraph_after(anchor, text, style)


def update_annotation(doc: Document) -> None:
    start = find_index(doc, "Аннотация")
    paragraph = doc.paragraphs[start + 1]
    clear_paragraph(paragraph)
    run = paragraph.add_run(
        "Пояснительная записка к выпускной квалификационной работе "
        "«Разработка системы предиктивного обслуживания узлов робота-паллетизатора "
        "на участке розлива продукции ООО «Компания “Здоровая жизнь”» включает "
        "аннотацию, перечень сокращений, введение, шесть основных глав, заключение, "
        "список литературы и приложения. В работе рассматриваются предпроектное "
        "обследование роботизированного участка паллетизации, проектирование "
        "программно-аппаратного комплекса сбора и анализа телеметрии, разработка "
        "цифровой модели робота, формирование модели деградации узлов и оценка "
        "эффективности предиктивного обслуживания."
    )
    set_run_font(run, 14)
    format_body(paragraph)


def update_front_matter(doc: Document) -> None:
    toc_title = find_paragraph(doc, "Содержание")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        set_run_font(run, 14, bold=True)
    remove_num_pr(toc_title)

    replace_between(doc, "Содержание", "Перечень принятых сокращений", [("", None)])
    toc_paragraph = doc.paragraphs[find_index(doc, "Содержание") + 1]
    add_toc_field(toc_paragraph)

    abbrev_heading = find_paragraph(doc, "Перечень принятых сокращений")
    format_heading(abbrev_heading, 1, numbered=False)
    abbrev_heading.paragraph_format.page_break_before = True
    replace_between(doc, "Перечень принятых сокращений", "Введение", [(item, "Normal") for item in ABBREVIATIONS])
    for idx in range(find_index(doc, "Перечень принятых сокращений") + 1, find_index(doc, "Введение")):
        format_body(doc.paragraphs[idx], first_line=False)


def update_intro(doc: Document) -> None:
    intro_heading = find_paragraph(doc, "Введение")
    format_heading(intro_heading, 1, numbered=False)
    intro_heading.paragraph_format.page_break_before = True

    intro_items: list[tuple[str, str | None]] = [(text, "Normal") for text in INTRO_PARAGRAPHS]
    intro_items.extend((task, "List Paragraph") for task in INTRO_TASKS)
    intro_items.extend((text, "Normal") for text in INTRO_TAIL)

    first_structure = "Предпроектное обследование"
    replace_between(doc, "Введение", first_structure, intro_items)
    start = find_index(doc, "Введение") + 1
    end = find_index(doc, first_structure)
    for idx, paragraph in enumerate(doc.paragraphs[start:end]):
        if paragraph.style and paragraph.style.name == "List Paragraph":
            set_num_pr(paragraph, "10", "0")
            format_body(paragraph, first_line=False)
        else:
            remove_num_pr(paragraph)
            format_body(paragraph)


def rebuild_structure_and_bibliography(doc: Document) -> None:
    start_idx = find_index(doc, "Предпроектное обследование")
    for paragraph in list(doc.paragraphs[start_idx:]):
        delete_paragraph(paragraph)

    first_chapter = True
    for item_type, title in STRUCTURE:
        if item_type == "chapter":
            p = doc.add_paragraph(title, style="Heading 1")
            format_heading(p, 1, numbered=True)
            if first_chapter:
                p.paragraph_format.page_break_before = True
                first_chapter = False
        else:
            p = doc.add_paragraph(title, style="Heading 2")
            format_heading(p, 2, numbered=True)

    conclusion = doc.add_paragraph("Заключение", style="Heading 1")
    format_heading(conclusion, 1, numbered=False)
    conclusion.paragraph_format.page_break_before = True

    bibliography_heading = doc.add_paragraph("Список литературы", style="Heading 1")
    format_heading(bibliography_heading, 1, numbered=False)
    bibliography_heading.paragraph_format.page_break_before = True

    for entry in BIBLIOGRAPHY:
        p = doc.add_paragraph(entry, style="Normal")
        set_num_pr(p, "6", "0")
        format_body(p, first_line=False)

    appendix = doc.add_paragraph("Приложения", style="Heading 1")
    format_heading(appendix, 1, numbered=False)
    appendix.paragraph_format.page_break_before = True


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: update_vkr_rpz_lit_intro.py <docx>")

    docx_path = Path(sys.argv[1])
    doc = Document(str(docx_path))
    update_annotation(doc)
    update_front_matter(doc)
    update_intro(doc)
    rebuild_structure_and_bibliography(doc)
    doc.save(str(docx_path))
    set_update_fields_on_open(docx_path)
    print(f"Updated {docx_path}")


if __name__ == "__main__":
    main()
