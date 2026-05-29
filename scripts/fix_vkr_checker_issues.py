from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"


SMALL_HEADINGS_TO_REMOVE = {
    "1.3.1. Контекстная функция",
    "1.7.2. Границы разработки",
    "2.1.1. Цель системы",
    "2.2.2. Информационная модель наблюдения",
    "2.3.2. Расчет признаков",
    "2.6.1. Функциональные требования",
    "4.4.2. Частота и объем данных",
    "4.5.2. Формирование признаков",
    "4.7.1. Постановка задачи",
    "4.7.2. Обучающая выборка и целевая переменная",
    "4.9.1 Модель хранения временных рядов",
    "4.9.1. Модель хранения временных рядов",
    "4.9.2. Контроль целостности данных",
    "4.10.1. Состав панели мониторинга",
    "4.10.2. Правила предупреждений",
}


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = paragraph.part.document.styles["Normal"]
    new_paragraph.add_run(text)
    format_body_paragraph(new_paragraph)
    return new_paragraph


def format_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(1.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def find_paragraph(document: Document, fragment: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    return None


def body_text_before_bibliography(document: Document) -> str:
    result: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.upper().startswith("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
            break
        result.append(text)
    return "\n".join(result)


def add_after_if_absent(document: Document, anchor: str, marker: str, text: str) -> bool:
    if marker in body_text_before_bibliography(document):
        return False
    paragraph = find_paragraph(document, anchor)
    if paragraph is None:
        raise RuntimeError(f"Anchor not found: {anchor}")
    insert_paragraph_after(paragraph, text)
    return True


def main() -> None:
    backup = DOCX.with_name(
        f"{DOCX.stem}.backup_before_checker_fixes_{datetime.now():%Y%m%d_%H%M%S}.docx"
    )
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)

    removed: list[str] = []
    for paragraph in list(document.paragraphs):
        text = normalize_text(paragraph.text)
        is_heading = paragraph.style is not None and paragraph.style.name.startswith("Heading")
        if is_heading and (not text or text in SMALL_HEADINGS_TO_REMOVE):
            if text:
                removed.append(text)
            delete_paragraph(paragraph)

    additions = [
        (
            "Предпроектное обследование задает исходные данные",
            "Глава 1 дополнительно связывает",
            "Глава 1 дополнительно связывает производственную постановку с требованиями к будущей системе: от технологического цикла и контролируемых узлов выполняется переход к составу телеметрии, диагностическим признакам и ограничениям прототипа. Такой вводный блок нужен для того, чтобы последующие проектные решения не выглядели оторванными от объекта автоматизации.",
        ),
        (
            "Перспективным способом решения данной задачи является применение цифрового моделирования",
            "Термины надежности",
            "Термины надежности и предельного состояния используются в соответствии с ГОСТ 27.002-2015 [1, с. 3] и общими положениями расчета надежности по ГОСТ 27.301-95 [3, с. 2]. Требования к структуре технического задания согласованы с действующим ГОСТ 34.602-2020 [5, с. 4], а оценка программно-технической надежности опирается на подходы, рассмотренные в работе Е.М. Лаврищевой и соавторов [6, с. 96].",
        ),
        (
            "Термины надежности и предельного состояния",
            "Применение цифровых двойников",
            "Применение цифровых двойников и имитационных моделей в производственных системах рассматривается как способ связать физический объект, модель, данные и контур поддержки решений [21, с. 108953]. В обзорах также подчеркивается, что практическая ценность цифрового двойника раскрывается при совместном использовании технологической модели, потоков данных и аналитических алгоритмов [22], [23], [24].",
        ),
        (
            "ABB IRB 660–180/3.15 относится к специализированным паллетизирующим роботам",
            "Паспортные ограничения",
            "Паспортные ограничения, состав сервисной информации и требования к эксплуатации робота IRB 660 уточняются по руководству ABB Robotics [12, с. 45]. Эти сведения используются как верхняя граница допустимых нагрузок и как основание для выбора контролируемых узлов.",
        ),
        (
            "В практической части этот показатель должен быть рассчитан",
            "Выбор регрессионной постановки",
            "Выбор регрессионной постановки RUL соответствует исследованиям по прогнозированию остаточного ресурса производственного оборудования [14], [15, с. 155] и обзорам PHM для технологических машин [16, с. 2844], [17]. Современные обзоры RUL показывают, что для вращающихся и механических узлов применяются как классические модели машинного обучения, так и нейросетевые методы [18], [19, с. 3444], [42], [43], а для робототехнических систем актуальна связка телеметрии и алгоритмов реального времени [44].",
        ),
        (
            "Граница между моделью и аналитическим контуром проходит по телеметрии",
            "Реализация цифрового стенда",
            "Реализация цифрового стенда опирается на руководство CoppeliaSim [25], описание ZeroMQ Remote API [26], общий обзор удаленного API [27] и пакет Python-клиента CoppeliaSim ZMQ Remote API [28]. Это позволяет отделить модель движения от внешнего контура сбора и обработки телеметрии.",
        ),
        (
            "Для первичной реализации рассматриваются три класса моделей",
            "Для воспроизводимости обучения",
            "Для воспроизводимости обучения и оценки моделей используются материалы scikit-learn по общему пользовательскому руководству [29] и разделу supervised learning [30]. Настройка градиентного бустинга соотносится с документацией XGBoost [31], описанием параметров [32] и справочником Python API [33].",
        ),
        (
            "Хранилище должно поддерживать запись временных рядов с тегами",
            "Логика хранения временных рядов",
            "Логика хранения временных рядов согласуется с документацией InfluxDB [34], описанием ключевых понятий InfluxDB v2 [35] и общим пояснением назначения time series database [36]. Поэтому в проекте используются теги цикла, фазы, оси и типа сигнала, а измерения сохраняются как последовательность наблюдений во времени.",
        ),
        (
            "Визуальное предупреждение должно сопровождаться указанием привода",
            "Состав панели и правила предупреждений",
            "Состав панели и правила предупреждений проектируются с учетом документации Grafana [37], рекомендаций по построению dashboards [38] и механизма alerting [39]. Такой подход позволяет разделить оперативный контроль текущих значений и аналитическое наблюдение за трендом деградации.",
        ),
    ]

    inserted: list[str] = []
    for anchor, marker, text in additions:
        if add_after_if_absent(document, anchor, marker, text):
            inserted.append(marker)

    document.save(DOCX)

    requested = [
        1,
        3,
        5,
        6,
        12,
        14,
        15,
        16,
        17,
        18,
        19,
        21,
        22,
        23,
        24,
        25,
        26,
        27,
        28,
        29,
        30,
        31,
        32,
        33,
        34,
        35,
        36,
        37,
        38,
        39,
        42,
        43,
        44,
    ]
    updated = Document(DOCX)
    body = body_text_before_bibliography(updated)
    cited = sorted(set(int(n) for n in re.findall(r"\[(\d+)(?=[,\]\s])", body)))
    remaining = [n for n in requested if n not in cited]
    headings = {
        normalize_text(p.text)
        for p in updated.paragraphs
        if p.style is not None and p.style.name.startswith("Heading")
    }
    still_headings = sorted(SMALL_HEADINGS_TO_REMOVE & headings)

    print(f"backup={backup}")
    print(f"removed={len(removed)}")
    for item in removed:
        print(f"removed_heading={item}")
    print(f"inserted={len(inserted)}")
    for item in inserted:
        print(f"inserted_marker={item}")
    print(f"missing_requested_refs={remaining}")
    print(f"remaining_small_headings={still_headings}")


if __name__ == "__main__":
    main()
