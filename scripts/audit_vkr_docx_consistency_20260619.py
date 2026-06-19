from __future__ import annotations

import csv
import importlib.util
import json
import re
from collections import Counter, defaultdict
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


ROOT = Path(".")
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
OUT = ROOT / "scratch" / "vkr_consistency_audit_20260619.json"
HELPER_PATH = ROOT / "scripts" / "apply_vkr_formula_logic_cleanup_20260619.py"


def load_helper():
    spec = importlib.util.spec_from_file_location("helper", HELPER_PATH)
    helper = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(helper)
    return helper


def text_escape(s: str) -> str:
    return s.replace("\n", " ")[:240]


def all_doc_text(doc: Document) -> str:
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def iter_table_text(doc: Document):
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield ti, ri, ci, pi, p.text


def caption_numbers(paras_text: list[str], kind: str) -> list[tuple[int, int, str]]:
    if kind == "table":
        rx = re.compile(r"^\s*Таблица\s+(\d+)\b")
    elif kind == "figure":
        rx = re.compile(r"^\s*Рисунок\s+(\d+)\b")
    else:
        raise ValueError(kind)
    out = []
    for i, txt in enumerate(paras_text):
        m = rx.match(txt)
        if m:
            out.append((i, int(m.group(1)), txt))
    return out


def appendix_captions(paras_text: list[str]) -> dict[str, list[tuple[int, int, str]]]:
    rx = re.compile(r"^\s*(Таблица|Листинг)\s+([А-ЯЁ])\.(\d+)\b")
    result: dict[str, list[tuple[int, int, str]]] = defaultdict(list)
    for i, txt in enumerate(paras_text):
        m = rx.match(txt)
        if m:
            result[f"{m.group(1)} {m.group(2)}"].append((i, int(m.group(3)), txt))
    return result


def continuity(nums: list[int]) -> dict:
    if not nums:
        return {"count": 0, "min": None, "max": None, "missing": [], "duplicates": []}
    c = Counter(nums)
    return {
        "count": len(nums),
        "min": min(nums),
        "max": max(nums),
        "missing": [n for n in range(min(nums), max(nums) + 1) if n not in c],
        "duplicates": [n for n, count in c.items() if count > 1],
    }


def split_before_bib_and_appendix(paras_text: list[str]) -> tuple[int | None, int | None]:
    bib_candidates = []
    app_candidates = []
    for i, txt in enumerate(paras_text):
        stripped = txt.strip()
        # The static TOC contains entries such as "Список ...108" and
        # "Приложение А109"; use the later real headings for section bounds.
        if stripped.startswith("Список использованных источников"):
            bib_candidates.append(i)
        if stripped.startswith("Приложение А"):
            app_candidates.append(i)
    bib = bib_candidates[-1] if bib_candidates else None
    app_after_bib = [i for i in app_candidates if bib is None or i > bib]
    app = app_after_bib[0] if app_after_bib else (app_candidates[-1] if app_candidates else None)
    return bib, app


def find_refs(paras_text: list[str], label_set: set[int], kind: str, start: int, end: int | None):
    if kind == "formula":
        context_rx = re.compile(
            r"(формул|формуле|формулы|выражени|зависимост|метрик|MAE|RMSE|R²|RUL|HI|AТО|Kдан|Kфаз|Kпред|Qv|Tобн)",
            re.I,
        )
        num_rx = re.compile(r"\((\d{1,3})\)(?:\s*[–-]\s*\((\d{1,3})\))?")
    elif kind == "table":
        context_rx = re.compile(r"\bтаблиц[а-я]*\s+(\d+)(?:\s*[–-]\s*(\d+))?", re.I)
        num_rx = context_rx
    elif kind == "figure":
        context_rx = re.compile(r"\bрисунк[а-я]*\s+(\d+)(?:\s*[–-]\s*(\d+))?", re.I)
        num_rx = context_rx
    else:
        raise ValueError(kind)
    refs = []
    invalid = []
    end = len(paras_text) if end is None else end
    for i in range(start, end):
        txt = paras_text[i]
        if kind == "formula":
            if not context_rx.search(txt):
                continue
            matches = list(num_rx.finditer(txt))
        else:
            matches = list(num_rx.finditer(txt))
        for m in matches:
            a = int(m.group(1))
            b = int(m.group(2)) if m.group(2) else None
            refs.append({"paragraph": i, "match": m.group(0), "text": text_escape(txt)})
            nums = range(a, b + 1) if b is not None and b >= a else [a]
            for n in nums:
                if n not in label_set:
                    invalid.append({"paragraph": i, "number": n, "match": m.group(0), "text": text_escape(txt)})
    return {"count": len(refs), "refs": refs, "invalid": invalid}


def source_audit(paras_text: list[str], bib_idx: int | None, app_idx: int | None):
    if bib_idx is None:
        return {"error": "bibliography heading not found"}
    end = app_idx or len(paras_text)
    source_nums = []
    for txt in paras_text[bib_idx + 1 : end]:
        m = re.match(r"^\s*(\d+)\.\s+", txt)
        if m:
            source_nums.append(int(m.group(1)))
    main_text = "\n".join(paras_text[:bib_idx])
    cited = sorted({int(n) for n in re.findall(r"\[(\d{1,3})\]", main_text)})
    invalid = [n for n in cited if n not in set(source_nums)]
    uncited = [n for n in source_nums if n not in set(cited)]
    return {
        "sources": continuity(source_nums),
        "cited_count": len(cited),
        "cited_minmax": (min(cited), max(cited)) if cited else None,
        "invalid_citations": invalid,
        "uncited_sources": uncited,
    }


def style_audit(doc: Document):
    heading_large = []
    heading_not_14 = []
    table_non12 = []
    table_runs = 0
    for pi, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style is not None else ""
        if style_name.startswith("Heading") or re.match(r"^\s*((\d+(\.\d+)*)\.|Реферат|Введение|Заключение|Список использованных источников|Приложение)", p.text):
            for r in p.runs:
                if not r.text.strip():
                    continue
                size = r.font.size.pt if r.font.size is not None else None
                if size and size > 18:
                    heading_large.append((pi, size, text_escape(p.text)))
                if size and abs(size - 14) > 0.01:
                    heading_not_14.append((pi, size, text_escape(p.text)))
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.text.strip():
                            continue
                        table_runs += 1
                        size = r.font.size.pt if r.font.size is not None else None
                        if size and abs(size - 12) > 0.01:
                            table_non12.append((ti, size, text_escape(r.text)))
    return {
        "heading_large_count": len(heading_large),
        "heading_not_14_count": len(heading_not_14),
        "heading_large_examples": heading_large[:20],
        "heading_not_14_examples": heading_not_14[:20],
        "table_runs": table_runs,
        "table_non12_count": len(table_non12),
        "table_non12_examples": table_non12[:20],
    }


def data_audit(all_text_value: str):
    result = {}
    raw = ROOT / "data" / "telemetry" / "vkr_raw" / "long_live_01.jsonl"
    norm = ROOT / "data" / "telemetry" / "vkr_normalized" / "vkr_telemetry_normalized.csv"
    features = ROOT / "data" / "features" / "vkr_features.csv"
    degr = ROOT / "data" / "experiments" / "vkr_degradation_features.csv"
    nn_metrics = ROOT / "data" / "results" / "vkr_nn_rul_metrics.csv"
    summaries = [
        ROOT / "data" / "results" / "file_pipeline_run_summary.json",
        ROOT / "data" / "results" / "telemetry_validation_summary.json",
        ROOT / "data" / "results" / "feature_summary.json",
        ROOT / "data" / "results" / "degradation_summary.json",
        ROOT / "data" / "results" / "nn_rul_summary.json",
        ROOT / "data" / "results" / "rul_summary.json",
    ]
    if raw.exists():
        count = 0
        t_min = None
        t_max = None
        for line in raw.open("r", encoding="utf-8"):
            if not line.strip():
                continue
            count += 1
            try:
                d = json.loads(line)
            except json.JSONDecodeError:
                continue
            for key in ("t", "time", "timestamp", "sim_time", "simulation_time"):
                if key in d and isinstance(d[key], (int, float)):
                    val = float(d[key])
                    t_min = val if t_min is None else min(t_min, val)
                    t_max = val if t_max is None else max(t_max, val)
                    break
        result["raw_jsonl"] = {"count": count, "t_min": t_min, "t_max": t_max, "span": None if t_min is None or t_max is None else t_max - t_min}
    for path, key in [(norm, "normalized_csv"), (features, "features_csv"), (degr, "degradation_csv")]:
        if path.exists():
            with path.open("r", encoding="utf-8", newline="") as f:
                rows = sum(1 for _ in f) - 1
            result[key] = {"rows": rows}
    if nn_metrics.exists():
        with nn_metrics.open("r", encoding="utf-8", newline="") as f:
            rows = list(csv.DictReader(f))
        result["nn_metrics_rows"] = rows[:10]
    json_summaries = {}
    for p in summaries:
        if p.exists():
            try:
                json_summaries[p.name] = json.loads(p.read_text(encoding="utf-8"))
            except Exception as exc:
                json_summaries[p.name] = {"error": str(exc)}
    result["summaries"] = json_summaries
    expected_tokens = {
        "22174": "raw packets",
        "88696": "normalized rows",
        "600": "feature rows",
        "192000": "RUL/degradation rows",
        "153600": "train rows",
        "38400": "test rows",
        "1,441": "MAE comma",
        "2,144": "RMSE comma",
        "0,988": "R2 comma",
        "2059,05": "observation span",
        "0,0929": "mean sample step",
        "10,77": "observed frequency",
        "187": "cycle time",
        "231": "packages per hour",
        "14,55": "tonnes per hour",
        "58212": "tonnes per year",
        "450000": "economic effect",
    }
    result["expected_token_presence"] = {token: (token in all_text_value) for token in expected_tokens}
    stale_tokens = ["17920", "14336", "3584", "1,173", "1,442", "0,994", "23,6 / 472", "20 Гц", "final_scena", "pred_final"]
    result["stale_token_presence"] = {token: (token in all_text_value) for token in stale_tokens}
    return result


def formula_bodies(root, helper):
    formulas = []
    paras = root.xpath(".//w:body//w:p", namespaces=helper.NS)
    for i, p in enumerate(paras):
        if helper.is_formula_paragraph(p):
            lab = helper.formula_label(p)
            txt = helper.para_text(p)
            body = re.sub(r"\(\d+\)\s*$", "", txt)
            normalized = re.sub(r"\s+", "", body)
            formulas.append({"paragraph": i, "label": lab, "text": txt, "norm": normalized})
    duplicates = []
    by_norm = defaultdict(list)
    for f in formulas:
        by_norm[f["norm"]].append(f)
    for norm, items in by_norm.items():
        if norm and len(items) > 1:
            duplicates.append([{"label": x["label"], "text": x["text"]} for x in items])
    return formulas, duplicates


def main():
    helper = load_helper()
    doc = Document(DOCX)
    all_text_value = all_doc_text(doc)
    with ZipFile(DOCX) as z:
        zip_bad = z.testzip()
        root = etree.fromstring(z.read("word/document.xml"))
        rels = etree.fromstring(z.read("word/_rels/document.xml.rels"))
    paras = root.xpath(".//w:body//w:p", namespaces=helper.NS)
    paras_text = [helper.para_text(p) for p in paras]
    bib_idx, app_idx = split_before_bib_and_appendix(paras_text)

    formulas, dup_formula_bodies = formula_bodies(root, helper)
    formula_nums = [f["label"] for f in formulas if f["label"] is not None]
    tables = caption_numbers(paras_text[: app_idx or len(paras_text)], "table")
    figures = caption_numbers(paras_text[: app_idx or len(paras_text)], "figure")
    table_nums = [n for _, n, _ in tables]
    figure_nums = [n for _, n, _ in figures]
    image_rels = [
        r.get("Target")
        for r in rels.xpath(".//*[local-name()='Relationship']")
        if "image" in (r.get("Type") or "")
    ]

    report = {
        "docx": str(DOCX),
        "zip_bad": zip_bad,
        "paragraph_count_xml": len(paras_text),
        "paragraph_count_python_docx": len(doc.paragraphs),
        "table_count_python_docx": len(doc.tables),
        "formula_labels": continuity(formula_nums),
        "formula_duplicate_bodies": dup_formula_bodies[:20],
        "main_table_captions": continuity(table_nums),
        "main_figure_captions": continuity(figure_nums),
        "appendix_captions": {k: continuity([n for _, n, _ in v]) for k, v in appendix_captions(paras_text).items()},
        "image_relationship_count": len(image_rels),
        "image_relationship_targets": image_rels,
        "formula_refs": find_refs(paras_text, set(formula_nums), "formula", 0, bib_idx),
        "table_refs": find_refs(paras_text, set(table_nums), "table", 0, bib_idx),
        "figure_refs": find_refs(paras_text, set(figure_nums), "figure", 0, bib_idx),
        "source_audit": source_audit(paras_text, bib_idx, app_idx),
        "style_audit": style_audit(doc),
        "broken_word_refs": {
            "source_not_found": "Ошибка! Источник ссылки не найден" in all_text_value,
            "question_marks": "???" in all_text_value,
            "todo": bool(re.search(r"TODO|ВСТАВК|ДОБАВИТЬ|PLACEHOLDER", all_text_value, re.I)),
        },
        "scene_name_checks": {
            "vkr_scena.ttt": "vkr_scena.ttt" in all_text_value,
            "final_scena": "final_scena" in all_text_value,
            "pred_final": "pred_final" in all_text_value,
        },
        "data_audit": data_audit(all_text_value),
    }
    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
