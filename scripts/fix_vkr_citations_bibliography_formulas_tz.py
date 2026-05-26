from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from compress_vkr_filled_sections import (
    clear_paragraph,
    find_index,
    format_body,
    insert_paragraph_after,
    replace_section,
    set_num_pr,
    set_run_font,
    set_update_fields_on_open,
)


BIBLIOGRAPHY_ENTRIES = [
    "ГОСТ 27.002-2015. Надежность в технике. Термины и определения. М.: Стандартинформ, 2016.",
    "ГОСТ 27.003-2016. Надежность в технике. Состав и общие правила задания требований по надежности. М.: Стандартинформ, 2018.",
    "ГОСТ 27.301-95. Надежность в технике. Расчет надежности. Основные положения. Минск: Межгосударственный совет по стандартизации, метрологии и сертификации, 1995.",
    "ГОСТ 34.602-89. Информационная технология. Комплекс стандартов на автоматизированные системы. Техническое задание на создание автоматизированной системы. М.: Издательство стандартов, 1990.",
    "ГОСТ 34.602-2020. Информационные технологии. Комплекс стандартов на автоматизированные системы. Техническое задание на создание автоматизированной системы. М.: Российский институт стандартизации, 2021.",
    "Лаврищева Е.М., Зеленов С.В., Пакулин Н.В. Методы оценки надежности программных и технических систем // Труды ИСП РАН. 2019. Т. 31, вып. 5. С. 95-108. DOI: 10.15514/ISPRAS-2019-31(5)-7.",
    "Задиран К.С., Щербаков М.В., Сай Ван Квонг. Прогнозирование остаточного ресурса оборудования в условиях малой выборки данных // Управление большими системами. 2023. Вып. 102. С. 99-113.",
    "Равин А.А., Хруцкий О.В. Инженерные методы прогнозирования остаточного ресурса оборудования // Вестник Астраханского государственного технического университета. Серия: Морская техника и технология. 2018. № 1. С. 33-47. DOI: 10.24143/2073-1574-2018-1-33-47.",
    "Власов А.И., Григорьев П.В., Кривошеин А.И. Модель предиктивного обслуживания оборудования с применением беспроводных сенсорных сетей // Надежность и качество сложных систем. 2018. № 2(22). С. 26-35. DOI: 10.21685/2307-4205-2018-2-4.",
    "Койбагаров М.К., Исабеков Ж.Н., Курмангалиева Л.А., Байтурганова В.К., Рахметова П.М. Разработка системы предиктивного обслуживания на основе машинного обучения // Вестник Университета Шакарима. Технические науки. 2025. № 2(18). С. 121-128. DOI: 10.53360/2788-7995-2025-2(18)-14.",
    "Стародубцева С.А., Гусев А.С. Прогнозирование остаточного ресурса конструкций и деталей машин // Известия МГТУ «МАМИ». 2012. Т. 6, № 2-1. С. 355-360. DOI: 10.17816/2074-0530-68547.",
    "ABB Robotics. Product manual IRB 660 - 180/3.15, IRB 660 - 250/3.15. Document ID: 3HAC025755-001. Revision AK. ABB, 2026.",
    "ABB Robotics. IRB 660. Taking palletizing to new heights: datasheet. ABB, 2018.",
    "Kang Z., Catal C., Tekinerdogan B. Remaining Useful Life Prediction of Equipment in Production Lines Using Artificial Neural Networks // Sensors. 2021. Vol. 21, No. 3. Article 932. DOI: 10.3390/s21030932.",
    "Taşcı B., Nuhuoğlu E., Ardil C. Remaining useful lifetime prediction for equipment in production lines // International Journal of Mechanical and Mechatronics Engineering. 2023. Vol. 17, No. 4. P. 155-160.",
    "Baur M., Albertelli P., Monno M. A review of prognostics and health management of machine tools // The International Journal of Advanced Manufacturing Technology. 2020. Vol. 107. P. 2843-2863. DOI: 10.1007/s00170-020-05202-3.",
    "Gharib H., Kovács G. A Review of Prognostic and Health Management Methods for Industrial Equipment // Applied Sciences. 2023. Vol. 13, No. 24. Article 13173. DOI: 10.3390/app132413173.",
    "Liu Y. et al. A comprehensive overview of remaining useful life prediction // Mechanical Systems and Signal Processing. 2025. Vol. 222. Article 111780. DOI: 10.1016/j.ymssp.2024.111780.",
    "Kumar S. et al. A Comprehensive Review of Remaining Useful Life Prediction for Rotating Machines // Archives of Computational Methods in Engineering. 2024. Vol. 31. P. 3443-3476. DOI: 10.1007/s11831-024-10072-4.",
    "Kritzinger W. et al. Digital Twin in manufacturing: A categorical literature review and classification // IFAC-PapersOnLine. 2018. Vol. 51, No. 11. P. 1016-1022. DOI: 10.1016/j.ifacol.2018.08.474.",
    "Fuller A. et al. Digital Twin: Enabling Technologies, Challenges and Open Research // IEEE Access. 2020. Vol. 8. P. 108952-108971. DOI: 10.1109/ACCESS.2020.2998358.",
    "Sharma A. et al. Digital Twins: State of the art theory and practice, challenges, and open research questions // Journal of Industrial Information Integration. 2022. Vol. 30. Article 100383. DOI: 10.1016/j.jii.2022.100383.",
    "Soori M. et al. Digital twin for smart manufacturing, A review // Sustainable Manufacturing and Service Economics. 2023. Vol. 2. Article 100017. DOI: 10.1016/j.smse.2023.100017.",
    "Zhang L. et al. Digital Twins for Additive Manufacturing: A State-of-the-Art Review // Applied Sciences. 2020. Vol. 10, No. 23. Article 8350. DOI: 10.3390/app10238350.",
    "Coppelia Robotics. CoppeliaSim User Manual [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/ (дата обращения: 05.05.2026).",
    "Coppelia Robotics. ZeroMQ Remote API [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/en/zmqRemoteApiOverview.htm (дата обращения: 05.05.2026).",
    "Coppelia Robotics. Remote API overview [Электронный ресурс]. URL: https://manual.coppeliarobotics.com/en/remoteApiOverview.htm (дата обращения: 05.05.2026).",
    "CoppeliaSim ZMQ Remote API client. Python package [Электронный ресурс]. URL: https://pypi.org/project/coppeliasim-zmqremoteapi-client/ (дата обращения: 05.05.2026).",
    "Scikit-learn developers. Scikit-learn User Guide [Электронный ресурс]. URL: https://scikit-learn.org/stable/user_guide.html (дата обращения: 05.05.2026).",
    "Scikit-learn developers. Supervised learning [Электронный ресурс]. URL: https://scikit-learn.org/stable/supervised_learning.html (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Documentation [Электронный ресурс]. URL: https://xgboost.readthedocs.io/ (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Parameters [Электронный ресурс]. URL: https://xgboost.readthedocs.io/en/stable/parameter.html (дата обращения: 05.05.2026).",
    "XGBoost developers. XGBoost Python API Reference [Электронный ресурс]. URL: https://xgboost.readthedocs.io/en/stable/python/python_api.html (дата обращения: 05.05.2026).",
    "InfluxData. InfluxDB documentation [Электронный ресурс]. URL: https://docs.influxdata.com/ (дата обращения: 05.05.2026).",
    "InfluxData. InfluxDB key concepts v2 [Электронный ресурс]. URL: https://docs.influxdata.com/influxdb/v2/reference/key-concepts/ (дата обращения: 05.05.2026).",
    "InfluxData. Time Series Database explained [Электронный ресурс]. URL: https://www.influxdata.com/time-series-database/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Grafana documentation [Электронный ресурс]. URL: https://grafana.com/docs/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Dashboards documentation [Электронный ресурс]. URL: https://grafana.com/docs/grafana/latest/visualizations/dashboards/ (дата обращения: 05.05.2026).",
    "Grafana Labs. Alerting documentation [Электронный ресурс]. URL: https://grafana.com/docs/grafana/latest/alerting/ (дата обращения: 05.05.2026).",
    "Kumar P., Khalid S., Kim H.S. Prognostics and Health Management of Rotating Machinery of Industrial Robot with Deep Learning Applications - A Review // Mathematics. 2023. Vol. 11, No. 13. Article 3008. DOI: 10.3390/math11133008.",
    "Xiao B., Zhong J., Bao X., Chen L., Bao J., Zheng Y. Digital twin-driven prognostics and health management for industrial assets // Scientific Reports. 2024. Vol. 14. Article 13443. DOI: 10.1038/s41598-024-63990-0.",
    "Hu Y., Liu S., Lu H., Zhang H. Remaining Useful Life Model and Assessment of Mechanical Products: A Brief Review and a Note on the State Space Model Method // Chinese Journal of Mechanical Engineering. 2019. Vol. 32. Article 15. DOI: 10.1186/s10033-019-0317-y.",
    "Wang X., Wang T., Ming A., Han Q., Chu F., Zhang W., Li A. Deep Spatiotemporal Convolutional-Neural-Network-Based Remaining Useful Life Estimation of Bearings // Chinese Journal of Mechanical Engineering. 2021. Vol. 34. Article 62. DOI: 10.1186/s10033-021-00576-1.",
    "Tanveer M., Yazdani M.H., Khan R.T.A., Kim H.S. Real-Time AI-Driven Prognostics and Health Management in Robotics // Applied Sciences. 2026. Vol. 16, No. 7. Article 3441. DOI: 10.3390/app16073441.",
    "Wojtulewicz A., Chaber P. Industrial Robot Control System with a Predictive Maintenance Module Using IIoT Technology // Sensors. 2025. Vol. 25, No. 4. Article 1154. DOI: 10.3390/s25041154.",
]


CITATION_APPEND_RULES = [
    ("Традиционное планово-предупредительное обслуживание задает ремонтные воздействия", "[9, с. 28]"),
    ("Перспективным направлением является предиктивное обслуживание", "[7, с. 100]"),
    ("В качестве технической основы рассматривается промышленный робот ABB IRB 660-180/3.15", "[13, с. 1]"),
    ("Цифровая модель в CoppeliaSim используется как безопасная среда", "[41, с. 1]"),
    ("Методическую основу работы составляют системный анализ", "[6, с. 95]"),
    ("Линия розлива рассматривается как последовательная система", "[2, с. 3]"),
    ("Повторяемые траектории робота позволяют сравнивать телеметрию одинаковых фаз цикла", "[8, с. 34]"),
    ("Для задач PdM производственная функция дополняется диагностическим контуром", "[9, с. 28]"),
    ("ABB IRB 660-180/3.15 относится к специализированным паллетизирующим роботам", "[13, с. 1]"),
    ("Для робота-паллетизатора характерны износ зубчатых передач", "[11, с. 355]"),
    ("Для рассматриваемого участка возможны реактивное обслуживание", "[9, с. 28]"),
    ("Цифровая модель не заменяет промышленную диагностику", "[41, с. 1]"),
    ("PdM выбран для робота-паллетизатора по трем причинам", "[9, с. 28]"),
    ("Система строится как программно-аппаратный комплекс с пятью логическими уровнями", "[45, с. 1]"),
    ("Для каждого окна наблюдения W = [t0, t0 + Δt] рассчитываются статистические", "[8, с. 34]"),
    ("Поток данных строится по цепочке: CoppeliaSim -> Remote API", "[45, с. 1]"),
    ("На концептуальном уровне задача прогнозирования задается как регрессия", "[7, с. 101]"),
]


FORMULA_REPLACEMENTS = {
    "Cпр = tпр · Cч + Cбр + Cлог,": "C_пр = t_пр · C_ч + C_бр + C_лог",
    "RUL(t) = tпред - t,": "RUL(t) = t_пред - t",
    "Nпал = nсл · nгр,": "N_пал = n_сл · n_гр",
    "Qуп = 3600 · Nпал / Tпал,": "Q_уп = 3600 · N_пал / T_пал",
    "Tпал = tп + Σ(j=1..Nпал)(tзахв,j + tпер,j + tукл,j) + Σ(k=1..nсл-1)tлист,k + tотв,": "T_пал = t_п + Σ_{j=1}^{N_пал}(t_захв,j + t_пер,j + t_укл,j) + Σ_{k=1}^{n_сл-1}t_лист,k + t_отв",
    "F_i,s = {Mср, Mmax, σM, ωср, amax, E},": "F_{i,s} = {M_ср, M_max, σ_M, ω_ср, a_max, E}",
    "E_i,s = Σ |M_i(t) · ω_i(t)| · Δt,": "E_{i,s} = Σ |M_i(t) · ω_i(t)| · Δt",
    "Draw -> Dclean -> F -> HI -> RUL -> Aто,": "D_raw → D_clean → F → HI → RUL → A_ТО",
    "P_i(t) = M_i(t) · ω_i(t),": "P_i(t) = M_i(t) · ω_i(t)",
    "Kз,i = Mmax,i / Mдоп,i,": "K_з,i = M_max,i / M_доп,i",
    "L_i = Σ |M_i(t)| · Δt,": "L_i = Σ |M_i(t)| · Δt",
    "HI(t) = 1 - Σ w_j · f_j(t),": "HI(t) = 1 - Σ_{j=1}^{p} w_j · f_j(t)",
    "HI(t) <= HIкр или x_j(t) >= xкр,j.": "HI(t) ≤ HI_кр или x_j(t) ≥ x_кр,j",
    "RUL_N = Nкр - Nтек,": "RUL_N = N_кр - N_тек",
    "CΣ = Cто + Pотк · Cпр + Cинф,": "C_Σ = C_ТО + P_отк · C_пр + C_инф",
    "если RUL < Tпл + Tрез, то планировать ТО,": "если RUL < T_пл + T_рез, то планировать ТО",
    "S = {D, F, HI, RUL, Aто},": "S = {D, F, HI, RUL, A_ТО}",
    "MAE = (1/n)Σ|RUL_i - RUL^_i|,": "MAE = (1/n) · Σ_{i=1}^{n}|RUL_i - RUL̂_i|",
    "RMSE = sqrt((1/n)Σ(RUL_i - RUL^_i)^2).": "RMSE = √[(1/n) · Σ_{i=1}^{n}(RUL_i - RUL̂_i)^2]",
    "d_k = {t_k, c_k, s_k, i, q_i, ω_i, a_i, M_i, event_k},": "d_k = {t_k, c_k, s_k, i, q_i, ω_i, a_i, M_i, event_k}",
    "x_norm = (x - xmin) / (xmax - xmin),": "x_norm = (x - x_min) / (x_max - x_min)",
    "F_W = {mean(x), max(x), std(x), rms(x), slope(x), E},": "F_W = {mean(x), max(x), std(x), rms(x), slope(x), E}",
    "slope(x) = (x(t0 + Δt) - x(t0)) / Δt.": "slope(x) = [x(t_0 + Δt) - x(t_0)] / Δt",
    "V = Ns · Np · fs · b,": "V = N_s · N_p · f_s · b",
    "Vсут = V · Tсм · Kзм,": "V_сут = V · T_см · K_зм",
    "R = Σ w_j · r_j,": "R = Σ_{j=1}^{m} w_j · r_j",
    "RUL^ = g(F, θ),": "RUL̂ = g(F, θ)",
    "Kпред = Nсвоевр / Nпред.": "K_пред = N_своевр / N_пред",
    "Tобн <= Δtдоп,": "T_обн ≤ Δt_доп",
    "Pпотерь <= Pдоп,": "P_потерь ≤ P_доп",
    "Kготовн = Tработ / (Tработ + Tпрост),": "K_готовн = T_работ / (T_работ + T_прост)",
    "ΔCпр = Cпр,до - Cпр,после,": "ΔC_пр = C_пр,до - C_пр,после",
    "Kпред = Nсвоевр / Nпред,": "K_пред = N_своевр / N_пред",
    "Pпотерь = Nпот / Nобщ <= Pдоп.": "P_потерь = N_пот / N_общ ≤ P_доп",
    "RUL^ = g(F_W, HI, θ).": "RUL̂ = g(F_W, HI, θ)",
    "MAE = (1/n)Σ|RUL_i - RUL^_i|,": "MAE = (1/n) · Σ_{i=1}^{n}|RUL_i - RUL̂_i|",
    "RMSE = sqrt((1/n)Σ(RUL_i - RUL^_i)^2).": "RMSE = √[(1/n) · Σ_{i=1}^{n}(RUL_i - RUL̂_i)^2]",
    "RUL < Tпл + Tрез или HI <= HIкр,": "RUL < T_пл + T_рез или HI ≤ HI_кр",
    "Accept = 1, если все K_i = 1; иначе Accept = 0.": "Accept = 1, если K_i = 1 для всех i = 1...n; иначе Accept = 0",
}


NUMERIC_CITATION_RE = re.compile(r"\s*\[(?=[0-9])[\d,\-\sсcC\.]+\]")
VALID_CITATION_RE = re.compile(r"\[\d+,\s*с\.\s*\d+\]")


def set_cambria_math(run, size: int = 14) -> None:
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:cs"), "Cambria Math")
    run.font.size = Pt(size)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, 14)


def append_single_page_citation(paragraph: Paragraph, citation: str) -> bool:
    text = (paragraph.text or "").strip()
    if citation in text or VALID_CITATION_RE.search(text):
        return False
    if text.endswith("."):
        text = f"{text[:-1]} {citation}."
    elif text.endswith((":", ";", ",")):
        text = f"{text} {citation}"
    else:
        text = f"{text} {citation}"
    set_paragraph_text(paragraph, text)
    format_body(paragraph)
    return True


def strip_old_numeric_citations(doc: Document) -> int:
    bibliography_idx = find_index(doc, "Список литературы")
    changed = 0
    for paragraph in doc.paragraphs[:bibliography_idx]:
        text = paragraph.text
        if not text or "[" not in text:
            continue
        new_text = NUMERIC_CITATION_RE.sub("", text)
        new_text = re.sub(r"\s{2,}", " ", new_text).strip()
        if new_text != text:
            set_paragraph_text(paragraph, new_text)
            if not paragraph.style.name.startswith("Heading"):
                format_body(paragraph)
            changed += 1
    return changed


def append_literature_citations(doc: Document) -> int:
    changed = 0
    bibliography_idx = find_index(doc, "Список литературы")
    for needle, citation in CITATION_APPEND_RULES:
        for paragraph in doc.paragraphs[:bibliography_idx]:
            if needle in (paragraph.text or ""):
                if append_single_page_citation(paragraph, citation):
                    changed += 1
                break
    return changed


def rebuild_bibliography(doc: Document) -> int:
    start_idx = find_index(doc, "Список литературы")
    end_idx = find_index(doc, "Приложения")
    anchor = doc.paragraphs[start_idx]
    end = doc.paragraphs[end_idx]
    parent = anchor._element.getparent()
    start_pos = parent.index(anchor._element)
    end_pos = parent.index(end._element)
    for element in list(parent)[start_pos + 1 : end_pos]:
        parent.remove(element)

    current = anchor
    for entry in BIBLIOGRAPHY_ENTRIES:
        current = insert_paragraph_after(current, entry, "Normal")
        set_num_pr(current, "6", "0")
        format_body(current, first_line=False)
    return len(BIBLIOGRAPHY_ENTRIES)


def is_formula_text(text: str) -> bool:
    if not text or len(text) > 180:
        return False
    if text.startswith("[") or text.endswith(";"):
        return False
    formula_markers = ["=", "Σ", "√", "≤", "≥", "→", "{", "}"]
    if not any(marker in text for marker in formula_markers):
        return False
    prose_markers = ["где ", "если ", "или ", "Согласно ", "Для ", "При "]
    if any(text.startswith(marker) for marker in prose_markers):
        return False
    return True


def improve_formulas(doc: Document) -> int:
    bibliography_idx = find_index(doc, "Список литературы")
    changed = 0
    for paragraph in doc.paragraphs[:bibliography_idx]:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        if text in FORMULA_REPLACEMENTS:
            text = FORMULA_REPLACEMENTS[text]
            clear_paragraph(paragraph)
            paragraph.add_run(text)
            changed += 1
        if is_formula_text(text) and paragraph.style.name in {"Normal", "List Paragraph"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = paragraph.paragraph_format
            pf.first_line_indent = None
            pf.line_spacing = 1.15
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            for run in paragraph.runs:
                set_cambria_math(run, 14)
    return changed


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: fix_vkr_citations_bibliography_formulas_tz.py <docx> <chapter3.md>")

    docx_path = Path(sys.argv[1])
    chapter3_path = Path(sys.argv[2])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = docx_path.with_name(
        f"{docx_path.stem}.backup_before_gost_citation_formula_fix_{timestamp}{docx_path.suffix}"
    )
    shutil.copy2(docx_path, backup)

    doc = Document(str(docx_path))
    stripped = strip_old_numeric_citations(doc)
    appended = append_literature_citations(doc)
    replace_section(doc, "Техническое задание", "Техническое проектирование", chapter3_path, 1, True)
    doc.paragraphs[find_index(doc, "Техническое задание")].paragraph_format.page_break_before = True
    bibliography_count = rebuild_bibliography(doc)
    formula_changes = improve_formulas(doc)
    doc.save(str(docx_path))
    set_update_fields_on_open(docx_path)

    print(f"Updated VKR RPZ: {docx_path}")
    print(f"Backup: {backup}")
    print(f"Old numeric citations stripped: {stripped}")
    print(f"Single-source page citations appended: {appended}")
    print(f"Bibliography entries: {bibliography_count}")
    print(f"Formula replacements: {formula_changes}")


if __name__ == "__main__":
    main()
