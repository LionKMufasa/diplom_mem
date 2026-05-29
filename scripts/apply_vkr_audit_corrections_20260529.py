from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"


BIBLIOGRAPHY_FIXES = {
    "Taşcı": (
        "Taşcı B., Omar A., Ayvaz S. Remaining useful lifetime prediction for "
        "predictive maintenance in manufacturing // Computers & Industrial Engineering. "
        "2023. Vol. 184. Article 109566. DOI: 10.1016/j.cie.2023.109566."
    ),
    "Gharib": (
        "Gharib H., Kovács G. A Review of Prognostic and Health Management (PHM) "
        "Methods and Limitations for Marine Diesel Engines: New Research Directions // "
        "Machines. 2023. Vol. 11, No. 7. Article 695. DOI: 10.3390/machines11070695."
    ),
    "Liu Y.": (
        "Liu Y., Wen J., Wang G. A comprehensive overview of remaining useful life "
        "prediction: From traditional literature review to scientometric analysis // "
        "Machine Learning with Applications. 2025. Vol. 21. Article 100704. "
        "DOI: 10.1016/j.mlwa.2025.100704."
    ),
    "Kumar S. et al. A Comprehensive Review of Remaining Useful Life Prediction": (
        "Kumar S., Raj K.K., Cirrincione M., Cirrincione G., Franzitta V., Kumar R.R. "
        "A Comprehensive Review of Remaining Useful Life Estimation Approaches for "
        "Rotating Machinery // Energies. 2024. Vol. 17, No. 22. Article 5538. "
        "DOI: 10.3390/en17225538."
    ),
}


RUN_REPLACEMENTS = {
    "56 строк фазовых признаков": "600 строк фазовых признаков",
    "56 строк признаков": "600 строк признаков",
    "56 строк": "600 строк",
    "17920 RUL-оценок": "192000 RUL-оценок",
    "17920 нейросетевых прогнозов": "192000 нейросетевых прогнозов",
    "17920 строк RUL-оценок": "192000 строк RUL-оценок",
    "17920 строк нейросетевых прогнозов": "192000 строк нейросетевых прогнозов",
    "17920 прогнозных строк": "192000 прогнозных строк",
    "17920 RUL-строк": "192000 RUL-строк",
    "17920 RUL‑строк": "192000 RUL-строк",
    "17920 RUL‑оценок": "192000 RUL-оценок",
    "17920 расчетных строк": "192000 расчетных строк",
    "17920 прогнозов": "192000 прогнозов",
    "17920 строк": "192000 строк",
    "14336 строках": "153600 строках",
    "14336 строк": "153600 строк",
    "3584 строках": "38400 строках",
    "3584 строки": "38400 строк",
    "38400 строки": "38400 строк",
    "3584": "38400",
    "MAE = 1,173": "MAE = 1,441",
    "RMSE = 1,442": "RMSE = 2,144",
    "R² = 0,994": "R² = 0,988",
    "1,173 цикла": "1,441 цикла",
    "1,173": "1,441",
    "1,442 цикла": "2,144 цикла",
    "1,442": "2,144",
    "0,994": "0,988",
    "MAE=1,173; RMSE=1,442; R2=0,994": "MAE=1,441; RMSE=2,144; R2=0,988",
}


TABLE_EXACT_REPLACEMENTS = {
    "56": "600",
    "17920": "192000",
    "3584": "38400",
    "1,173": "1,441",
    "1,442": "2,144",
    "0,994": "0,988",
    "204": "72",
}


SCENARIO_METRIC_ROWS = {
    "S0": ("9600", "0,814", "0,850", "1,000"),
    "S1": ("9600", "2,544", "4,056", "0,964"),
    "S2": ("9600", "1,395", "2,002", "0,991"),
    "S3": ("9600", "1,013", "1,668", "0,996"),
    "Среднее": ("38400", "1,441", "2,144", "0,988"),
}


CITATION_WITH_PAGE_RE = re.compile(r"\[(\d+)\s*,\s*[сc]\.?\s*[^\]]+\]", re.IGNORECASE)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def format_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(1.25)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    format_body_paragraph(paragraph)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = paragraph.part.document.styles["Normal"]
    result.add_run(text)
    format_body_paragraph(result)
    return result


def replace_in_runs(paragraph: Paragraph, replacements: dict[str, str]) -> int:
    changed = 0
    for run in paragraph.runs:
        text = run.text
        new_text = CITATION_WITH_PAGE_RE.sub(r"[\1]", text)
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            run.text = new_text
            changed += 1
    return changed


def all_paragraphs(document: Document) -> list[Paragraph]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def bibliography_index(document: Document) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if normalize_text(paragraph.text).startswith("Список использованных источников"):
            return index
    return len(document.paragraphs)


def apply_citation_and_number_replacements(document: Document) -> tuple[int, int]:
    biblio_at = bibliography_index(document)
    citation_changes = 0
    numeric_changes = 0

    for index, paragraph in enumerate(document.paragraphs):
        replacements = RUN_REPLACEMENTS
        before = paragraph.text
        if index < biblio_at:
            replace_in_runs(paragraph, replacements)
        else:
            replace_in_runs(paragraph, {})
        after = paragraph.text
        if before != after:
            numeric_changes += int(any(old in before for old in replacements))
            citation_changes += int(CITATION_WITH_PAGE_RE.search(before) is not None)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = normalize_text(cell.text)
                full_replaced = cell_text
                for old, new in RUN_REPLACEMENTS.items():
                    full_replaced = full_replaced.replace(old, new)
                if cell_text in TABLE_EXACT_REPLACEMENTS:
                    set_cell_text(cell, TABLE_EXACT_REPLACEMENTS[cell_text])
                    numeric_changes += 1
                    continue
                if full_replaced != cell_text:
                    set_cell_text(cell, full_replaced)
                    numeric_changes += 1
                    continue
                for paragraph in cell.paragraphs:
                    before = paragraph.text
                    replace_in_runs(paragraph, RUN_REPLACEMENTS)
                    after = paragraph.text
                    if before != after:
                        numeric_changes += int(any(old in before for old in RUN_REPLACEMENTS))
                        citation_changes += int(CITATION_WITH_PAGE_RE.search(before) is not None)

    return citation_changes, numeric_changes


def fix_bibliography(document: Document) -> int:
    changes = 0
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        for marker, replacement in BIBLIOGRAPHY_FIXES.items():
            if text.startswith(marker):
                set_paragraph_text(paragraph, replacement)
                changes += 1
                break
    return changes


def apply_heading_levels(document: Document) -> int:
    changes = 0
    heading1 = document.styles["Heading 1"]
    heading2 = document.styles["Heading 2"]
    heading3 = document.styles["Heading 3"]
    special_level1 = {
        "Реферат",
        "Перечень принятых сокращений",
        "Введение",
        "Заключение",
        "Список использованных источников",
    }
    appendix_re = re.compile(r"^Приложение\s+[А-ЯЁ]$")
    h3_re = re.compile(r"^\d+\.\d+\.\d+\.?\s+")
    h2_re = re.compile(r"^\d+\.\d+\.?\s+")
    h1_re = re.compile(r"^\d+\.?\s+")

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        target = None
        if text in special_level1 or appendix_re.match(text):
            target = heading1
        elif h3_re.match(text):
            target = heading3
        elif h2_re.match(text):
            target = heading2
        elif h1_re.match(text):
            target = heading1
        if target is not None and paragraph.style.style_id != target.style_id:
            paragraph.style = target
            changes += 1
    return changes


def rewrite_target_paragraphs(document: Document) -> int:
    changes = 0
    replacements_by_anchor = [
        (
            "На основе этих фаз сформировано",
            "Средний шаг записи по исходным пакетам составил 0,0929 с, что соответствует частоте 10,77 Гц. "
            "В данных выделено 14 фаз технологического цикла и восстановлено 12 циклов паллетизации; "
            "дополнительно зафиксирован 1121 фазовый сегмент, что позволяет отделять повторяющиеся операции "
            "внутри длинного прогона. На основе фаз, циклов и четырех осей сформировано 600 строк признаков "
            "для последующего моделирования деградации.",
        ),
        (
            "Нейросетевая модель MLPRegressor обучалась",
            "Нейросетевая модель MLPRegressor обучалась на 153600 строках и проверялась на 38400 строках. "
            "Разделение выполнено по синтетическим циклам, поэтому строки одного и того же расчетного цикла "
            "не попадают одновременно в обучающую и тестовую части. Средняя ошибка на тестовой выборке "
            "составила MAE = 1,441 цикла, RMSE = 2,144 цикла, средний коэффициент детерминации R² = 0,988. "
            "Для сравнения детерминированная трендовая оценка на том же наборе дает MAE около 6,335 цикла "
            "и R² около 0,807, поэтому нейросетевая модель точнее описывает нелинейные сценарии S1...S3.",
        ),
        (
            "В рабочем варианте целесообразно использовать два уровня моделей.",
            "В рабочем варианте используются два уровня моделей. Первый уровень - простая базовая модель тренда, "
            "которая обеспечивает интерпретируемую проверку расчетов. Второй уровень - MLPRegressor из scikit-learn: "
            "эта компактная нейросетевая регрессионная модель выбрана для апробации нелинейной зависимости между "
            "признаками деградации и RUL на сформированных сценариях. Random Forest и XGBoost остаются сравнительными "
            "или резервными вариантами для будущего расширения набора данных.",
        ),
        (
            "Апробация выполнена на наборе long_live_01.",
            "Апробация выполнена на наборе long_live_01. Получено 22174 пакета телеметрии и 88696 нормализованных строк "
            "по четырем осям; коэффициенты полноты данных и фазовой разметки составили 1,000. На основе 600 строк "
            "фазовых признаков сформировано 192000 RUL-оценок и 192000 нейросетевых прогнозов. Средняя тестовая ошибка "
            "MLPRegressor равна MAE = 1,441 цикла, RMSE = 2,144 цикла, R² = 0,988, что подтверждает пригодность выбранного "
            "подхода для учебного прототипа предиктивного обслуживания.",
        ),
        (
            "Численные результаты показывают, что коэффициенты полноты данных",
            "Численные результаты показывают, что коэффициенты полноты данных и фазовой разметки равны 1,000, "
            "средняя тестовая ошибка нейросетевой модели составляет 1,441 цикла при R² = 0,988, а для расчетного "
            "экономического сценария годовой эффект равен 450000 руб. при сроке окупаемости 1,0 год. Практическая "
            "часть подтверждает работоспособность разработанного контура, но точность RUL должна рассматриваться "
            "как результат модельной апробации, а не как готовая промышленная гарантия прогноза.",
        ),
    ]
    for paragraph in document.paragraphs:
        text = paragraph.text
        for anchor, replacement in replacements_by_anchor:
            if anchor in text:
                set_paragraph_text(paragraph, replacement)
                changes += 1
                break
    return changes


def insert_limitations(document: Document) -> int:
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    changes = 0
    limitation = (
        "Важно учитывать ограничение апробации: сценарии S0...S3 и коэффициент деградации alpha заданы модельно. "
        "Поэтому полученные значения HI/RUL и метрики качества подтверждают работоспособность алгоритмического "
        "контура на синтетически сформированных деградационных рядах, но не переносятся напрямую на реальный "
        "редуктор или привод без калибровки по истории ТОиР, отказов и диагностических измерений."
    )
    if "синтетически сформированных деградационных рядах" not in body:
        for paragraph in document.paragraphs:
            if "В данном подразделе повторно не приводятся формулы RUL" in paragraph.text:
                insert_paragraph_after(paragraph, limitation)
                changes += 1
                break

    conclusion = (
        "Ограничением выполненной апробации является синтетический характер деградационных сценариев. "
        "Для промышленного внедрения необходимо собрать реальные эксплуатационные ряды, уточнить веса признаков "
        "HI, проверить пороги предупреждений и переобучить модель RUL на данных фактических отказов и ремонтов."
    )
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "Ограничением выполненной апробации является синтетический характер" not in body:
        for paragraph in document.paragraphs:
            if paragraph.text.startswith("Полученный результат может быть использован"):
                insert_paragraph_after(paragraph, conclusion)
                changes += 1
                break
    return changes


def update_model_table(document: Document) -> int:
    changes = 0
    for table in document.tables:
        if not table.rows:
            continue
        header = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if header[:4] != ["Модель", "Преимущество", "Ограничение", "Роль в ВКР"]:
            continue
        has_mlp = False
        for row in table.rows[1:]:
            cells = row.cells
            if normalize_text(cells[0].text) == "XGBoost":
                set_cell_text(cells[3], "Резервная сравнительная модель")
                changes += 1
            if normalize_text(cells[0].text) == "MLPRegressor":
                has_mlp = True
        if not has_mlp:
            row = table.add_row()
            values = [
                "MLPRegressor",
                "Нелинейная регрессия при компактной структуре",
                "Чувствителен к синтетической разметке",
                "Рабочая модель апробации",
            ]
            for cell, value in zip(row.cells, values):
                set_cell_text(cell, value)
            changes += 1
        break
    return changes


def update_metric_table(document: Document) -> int:
    changes = 0
    for table in document.tables:
        if not table.rows:
            continue
        header = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if header[:6] != ["Модель", "Сценарий", "Строк", "MAE, цикл", "RMSE, цикл", "R2"]:
            continue
        for row in table.rows[1:]:
            scenario = normalize_text(row.cells[1].text)
            if scenario not in SCENARIO_METRIC_ROWS:
                continue
            count, mae, rmse, r2 = SCENARIO_METRIC_ROWS[scenario]
            for cell, value in zip(row.cells[2:6], [count, mae, rmse, r2]):
                if normalize_text(cell.text) != value:
                    set_cell_text(cell, value)
                    changes += 1
        break
    return changes


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def verify(document: Document) -> dict:
    biblio_at = bibliography_index(document)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs[:biblio_at])
    page_citation_left = CITATION_WITH_PAGE_RE.findall(body_text)
    all_text = "\n".join(paragraph.text for paragraph in all_paragraphs(document))
    old_tokens = [
        token
        for token in ["56 строк признаков", "17920", "14336", "3584", "1,173", "1,442", "0,994"]
        if token in all_text
    ]
    biblio_ok = {
        "tasci": "Computers & Industrial Engineering" in all_text and "10.1016/j.cie.2023.109566" in all_text,
        "gharib": "Machines. 2023. Vol. 11" in all_text and "10.3390/machines11070695" in all_text,
        "liu": "Machine Learning with Applications" in all_text and "10.1016/j.mlwa.2025.100704" in all_text,
        "kumar": "Energies. 2024. Vol. 17" in all_text and "10.3390/en17225538" in all_text,
    }
    heading_counts = {"Heading 1": 0, "Heading 2": 0, "Heading 3": 0}
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name in heading_counts:
            heading_counts[paragraph.style.name] += 1
    return {
        "page_citation_left": len(page_citation_left),
        "old_tokens": old_tokens,
        "biblio_ok": biblio_ok,
        "heading_counts": heading_counts,
    }


def main() -> None:
    backup = DOCX.with_name(f"{DOCX.stem}.backup_before_audit_corrections_20260529_{datetime.now():%H%M%S}.docx")
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    citation_changes, numeric_changes = apply_citation_and_number_replacements(document)
    biblio_changes = fix_bibliography(document)
    paragraph_changes = rewrite_target_paragraphs(document)
    limitation_changes = insert_limitations(document)
    model_table_changes = update_model_table(document)
    metric_table_changes = update_metric_table(document)
    heading_changes = apply_heading_levels(document)
    document.save(DOCX)

    updated = Document(DOCX)
    report = verify(updated)
    print(f"backup={backup}")
    print(f"citation_changes={citation_changes}")
    print(f"numeric_changes={numeric_changes}")
    print(f"biblio_changes={biblio_changes}")
    print(f"paragraph_changes={paragraph_changes}")
    print(f"limitation_changes={limitation_changes}")
    print(f"model_table_changes={model_table_changes}")
    print(f"metric_table_changes={metric_table_changes}")
    print(f"heading_changes={heading_changes}")
    print(f"verification={report}")


if __name__ == "__main__":
    main()
