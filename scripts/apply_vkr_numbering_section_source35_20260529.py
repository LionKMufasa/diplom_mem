from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"
FONT_NAME = "Times New Roman"


def set_run_font(run, *, size_pt: int = 14) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(key), FONT_NAME)


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        set_run_font(paragraph.runs[0])
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def insert_paragraph_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if style is not None:
        new_paragraph.style = style
    run = new_paragraph.add_run(text)
    set_run_font(run)
    return new_paragraph


def get_main_borders(doc: Document) -> tuple[int, int]:
    bib_start = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Список использованных источников")
    )
    app_start = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Приложение")
    )
    return bib_start, app_start


def renumber_formula_labels(doc: Document) -> tuple[dict[int, int], int]:
    old_to_new: dict[int, int] = {}
    changed = 0
    next_num = 1
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        m = re.fullmatch(r"\((\d+)\)", txt)
        if not m:
            continue
        old_num = int(m.group(1))
        new_num = next_num
        old_to_new.setdefault(old_num, new_num)
        label = f"({new_num})"
        for run in paragraph.runs:
            if re.search(r"\(\d+\)", run.text):
                run.text = re.sub(r"\(\d+\)", label, run.text)
                set_run_font(run)
                changed += 1
                break
        next_num += 1
    return old_to_new, changed


def update_formula_references(doc: Document, old_to_new: dict[int, int]) -> int:
    changed = 0

    def repl_range(match: re.Match[str]) -> str:
        a = int(match.group(1))
        b = int(match.group(2))
        return f"({old_to_new.get(a, a)})-({old_to_new.get(b, b)})"

    def repl_single(match: re.Match[str]) -> str:
        n = int(match.group(1))
        return f"({old_to_new.get(n, n)})"

    for paragraph in doc.paragraphs:
        txt = paragraph.text
        if "прогноз остаточного ресурса рассчитывается по формулам" in txt:
            replace_paragraph_text(
                paragraph,
                "В данном подразделе повторно не приводятся формулы RUL и метрик качества: прогноз "
                "остаточного ресурса рассчитывается по формулам (85)-(87), а MAE, RMSE и R² - по "
                "формулам (88)-(90). Далее используются только численные результаты апробации. "
                "Логика программной реализации расчета HI/RUL кратко приведена в приложении Г.",
            )
            changed += 1
            continue
        if "формул" not in txt.lower() and "формуле" not in txt.lower():
            continue
        new_txt = re.sub(r"\((\d+)\)-\((\d+)\)", repl_range, txt)
        new_txt = re.sub(r"\((\d+)\)", repl_single, new_txt)
        if new_txt != txt:
            replace_paragraph_text(paragraph, new_txt)
            changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                replacements = {
                    "Формулы (81)-(85)": "Формулы (80)-(84)",
                    "Формулы (74)-(77)": "Формулы (74)-(76)",
                    "Формулы (86)-(88)": "Формулы (85)-(87)",
                    "Формула (107)": "Формула (106)",
                }
                if txt in replacements:
                    cell.text = replacements[txt]
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run)
                    changed += 1
    return changed


def renumber_table_captions_and_refs(doc: Document) -> tuple[dict[int, int], int, int]:
    table_word = "Таблица"
    old_to_new: dict[int, int] = {}
    changed_caps = 0
    next_num = 1
    bib_start, app_start = get_main_borders(doc)

    for idx, paragraph in enumerate(doc.paragraphs[:app_start]):
        txt = paragraph.text.strip()
        m = re.match(rf"^{table_word}\s+(\d+)\s*[-–]?\s*(.*)", txt)
        if not m:
            continue
        old_num = int(m.group(1))
        rest = m.group(2).strip()
        if rest.startswith("-") or rest.startswith("–"):
            rest = rest[1:].strip()
        old_to_new[old_num] = next_num
        new_txt = f"{table_word} {next_num} - {rest}" if rest else f"{table_word} {next_num}"
        if new_txt != txt:
            replace_paragraph_text(paragraph, new_txt)
            changed_caps += 1
        next_num += 1

    changed_refs = 0
    for idx, paragraph in enumerate(doc.paragraphs[:app_start]):
        txt = paragraph.text
        if table_word in txt:
            continue
        new_txt = txt
        # Specific mismatch introduced by earlier table renumbering.
        new_txt = re.sub(r"таблице\s+14\b", "таблице 13", new_txt)
        new_txt = re.sub(r"таблицу\s+14\b", "таблицу 13", new_txt)
        # General updates for later table numbers, where captions shifted 39 -> 38, etc.
        for old_num, new_num in sorted(old_to_new.items(), key=lambda x: -x[0]):
            if old_num == new_num:
                continue
            new_txt = re.sub(rf"\bтаблице\s+{old_num}\b", f"таблице {new_num}", new_txt)
            new_txt = re.sub(rf"\bтаблицу\s+{old_num}\b", f"таблицу {new_num}", new_txt)
            new_txt = re.sub(rf"\bтаблицы\s+{old_num}\b", f"таблицы {new_num}", new_txt)
        if new_txt != txt:
            replace_paragraph_text(paragraph, new_txt)
            changed_refs += 1

    return old_to_new, changed_caps, changed_refs


def expand_section_361(doc: Document) -> int:
    inserted = 0
    marker = "Для учебного прототипа приняты три группы испытаний"
    if any(marker in p.text for p in doc.paragraphs):
        return 0

    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("Контроль выполнения ТЗ проводится поэтапно"):
            insert_paragraph_after(
                paragraph,
                "Для учебного прототипа приняты три группы испытаний: предварительные, функциональные и "
                "расчетно-аналитические. Предварительные испытания подтверждают корректность запуска цифровой "
                "сцены, наличие контролируемых объектов и воспроизводимость паллетизационного цикла. "
                "Функциональные испытания проверяют получение данных по осям motor1...motor4, запись времени, "
                "цикла и фазы, а также отсутствие критических потерь при нормализации телеметрии.",
                style=paragraph.style,
            )
            inserted += 1
            insert_paragraph_after(
                doc.paragraphs[idx + 1],
                "Расчетно-аналитические испытания выполняются после формирования признаков. На этом этапе "
                "проверяется расчет HI, RUL, уровня риска, рекомендаций по обслуживанию и метрик качества "
                "прогноза. Отдельно контролируется, что диагностический контур не передает управляющие команды "
                "роботу и используется только для мониторинга, анализа и поддержки решений по ТОиР.",
                style=paragraph.style,
            )
            inserted += 1
            break
    return inserted


def add_source35_reference(doc: Document) -> int:
    bib_start, _ = get_main_borders(doc)
    if any("[35]" in p.text for p in doc.paragraphs[:bib_start]):
        return 0
    for paragraph in doc.paragraphs:
        if "Логика хранения временных рядов согласуется" in paragraph.text:
            txt = paragraph.text
            txt = txt.replace(
                "описанием ключевых понятий InfluxDB v2 [34].",
                "описанием ключевых понятий InfluxDB v2 [34] и общими принципами баз данных временных рядов [35].",
            )
            replace_paragraph_text(paragraph, txt)
            return 1
    return 0


def ooxml_renumber_formula_labels() -> int:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    tmp = DOCX_PATH.with_suffix(".tmp.docx")
    changed = 0
    with zipfile.ZipFile(DOCX_PATH, "r") as zin:
        root = etree.fromstring(zin.read("word/document.xml"))
        next_num = 1
        for paragraph in root.xpath(".//w:p", namespaces=ns):
            if not paragraph.xpath(".//m:oMath | .//m:oMathPara", namespaces=ns):
                continue
            nodes = paragraph.xpath(".//w:t | .//m:t", namespaces=ns)
            full = "".join(node.text or "" for node in nodes)
            match = re.search(r"\((\d+)\)\s*$", full)
            if not match:
                continue
            old_label = match.group(0).strip()
            new_label = f"({next_num})"
            if old_label != new_label:
                suffix_nodes = []
                suffix_text = ""
                for node in reversed(nodes):
                    suffix_nodes.append(node)
                    suffix_text = (node.text or "") + suffix_text
                    if len(suffix_text) >= len(old_label):
                        break
                label_nodes = list(reversed(suffix_nodes))
                total = "".join(node.text or "" for node in label_nodes)
                if total.endswith(old_label):
                    new_total = total[: -len(old_label)] + new_label
                    label_nodes[0].text = new_total
                    for node in label_nodes[1:]:
                        node.text = ""
                    changed += 1
            next_num += 1

        new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = new_xml
                zout.writestr(item, data)
    tmp.replace(DOCX_PATH)
    return changed


def ooxml_fix_formula_reference_texts() -> int:
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    tmp = DOCX_PATH.with_suffix(".tmp.docx")
    changed = 0
    replacements = {
        "Срок окупаемости вычислен по формуле (113)": "Срок окупаемости вычислен по формуле (112)",
    }
    with zipfile.ZipFile(DOCX_PATH, "r") as zin:
        root = etree.fromstring(zin.read("word/document.xml"))
        for paragraph in root.xpath(".//w:p", namespaces=ns):
            nodes = paragraph.xpath(".//w:t | .//m:t", namespaces=ns)
            full = "".join(node.text or "" for node in nodes)
            if full in replacements and nodes:
                nodes[0].text = replacements[full]
                for node in nodes[1:]:
                    node.text = ""
                changed += 1
        new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = new_xml
                zout.writestr(item, data)
    tmp.replace(DOCX_PATH)
    return changed


def audit_docx(doc: Document) -> dict[str, object]:
    report: dict[str, object] = {}
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        report["zip_bad"] = str(zf.testzip())

    formula_nums = []
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    with zipfile.ZipFile(DOCX_PATH, "r") as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    for paragraph in root.xpath(".//w:p", namespaces=ns):
        if not paragraph.xpath(".//m:oMath | .//m:oMathPara", namespaces=ns):
            continue
        full = "".join(node.text or "" for node in paragraph.xpath(".//w:t | .//m:t", namespaces=ns)).strip()
        m = re.search(r"\((\d+)\)$", full)
        if m:
            formula_nums.append(int(m.group(1)))
    report["formula_count"] = len(formula_nums)
    report["formula_max"] = max(formula_nums) if formula_nums else 0
    report["formula_missing"] = [
        n for n in range(1, max(formula_nums) + 1) if n not in set(formula_nums)
    ] if formula_nums else []
    report["formula_duplicates"] = sorted({n for n in formula_nums if formula_nums.count(n) > 1})

    _, app_start = get_main_borders(doc)
    table_nums = []
    for paragraph in doc.paragraphs[:app_start]:
        m = re.match(r"^Таблица\s+(\d+)\s*[-–]", paragraph.text.strip())
        if m:
            table_nums.append(int(m.group(1)))
    report["table_count"] = len(table_nums)
    report["table_max"] = max(table_nums) if table_nums else 0
    report["table_missing"] = [
        n for n in range(1, max(table_nums) + 1) if n not in set(table_nums)
    ] if table_nums else []
    report["table_duplicates"] = sorted({n for n in table_nums if table_nums.count(n) > 1})
    report["source35_cited"] = int(any("[35]" in p.text for p in doc.paragraphs[: get_main_borders(doc)[0]]))
    report["section361_expanded"] = int(any("Для учебного прототипа приняты три группы испытаний" in p.text for p in doc.paragraphs))
    report["table13_ref_ok"] = int(any("приведены в таблице 13" in p.text for p in doc.paragraphs))
    return report


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    backup = DOCX_PATH.with_name(
        f"{DOCX_PATH.stem}.backup_before_numbering_section_source35_{datetime.now():%Y%m%d_%H%M%S}{DOCX_PATH.suffix}"
    )
    shutil.copy2(DOCX_PATH, backup)
    doc = Document(DOCX_PATH)
    formula_map, formula_labels_changed = renumber_formula_labels(doc)
    formula_refs_changed = update_formula_references(doc, formula_map)
    table_map, table_caps_changed, table_refs_changed = renumber_table_captions_and_refs(doc)
    inserted_361 = expand_section_361(doc)
    source35_added = add_source35_reference(doc)
    doc.save(DOCX_PATH)
    ooxml_formula_labels_changed = ooxml_renumber_formula_labels()
    ooxml_reference_texts_changed = ooxml_fix_formula_reference_texts()
    audited = Document(DOCX_PATH)
    print(f"backup={backup}")
    print(
        "report="
        + repr(
            {
                "formula_labels_changed": formula_labels_changed,
                "formula_refs_changed": formula_refs_changed,
                "table_caps_changed": table_caps_changed,
                "table_refs_changed": table_refs_changed,
                "section361_paragraphs_inserted": inserted_361,
                "source35_refs_added": source35_added,
                "ooxml_formula_labels_changed": ooxml_formula_labels_changed,
                "ooxml_reference_texts_changed": ooxml_reference_texts_changed,
                "formula_map_sample": {k: formula_map[k] for k in sorted(formula_map) if k != formula_map[k]},
                "table_map_changes": {k: table_map[k] for k in sorted(table_map) if k != table_map[k]},
            }
        )
    )
    print("audit=" + repr(audit_docx(audited)))


if __name__ == "__main__":
    main()
