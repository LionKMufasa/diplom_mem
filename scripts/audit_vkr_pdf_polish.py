from __future__ import annotations

import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "вкр" / "ВКР 2026 Миронов Егор Максимович.docx"


def all_text(doc: Document) -> list[str]:
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts


def main() -> None:
    with zipfile.ZipFile(DOCX) as zf:
        bad_member = zf.testzip()
    doc = Document(DOCX)
    texts = all_text(doc)
    flat = "\n".join(texts)

    table_captions = [
        " ".join(p.text.split())
        for p in doc.paragraphs
        if re.match(r"^Таблица\s+\d+", " ".join(p.text.split()))
    ]
    fig_captions = [
        " ".join(p.text.split())
        for p in doc.paragraphs
        if re.match(r"^Рисунок\s+\d+", " ".join(p.text.split()))
    ]
    fig_numbers = [
        int(match.group(1))
        for caption in fig_captions
        if (match := re.match(r"^Рисунок\s+(\d+)", caption))
    ]

    empty_tables = [
        caption
        for caption in table_captions
        if re.match(r"^Таблица\s+\d+\s*[-–]\s*$", caption)
    ]
    duplicate_figures = sorted(
        number for number, count in Counter(fig_numbers).items() if count > 1
    )
    non_sequential_figures = fig_numbers != list(range(1, len(fig_numbers) + 1))
    stale_tokens = [
        token
        for token in [
            "Цлевое значение",
            "Будущая вставка",
            "Будут вставлены после финальных прогонов",
            "формулам (88)-(90)",
            "формулам (91)-(93)",
            "final_scena_diplom.ttt",
            "pred_final.ttt",
        ]
        if token in flat
    ]
    formula_caption_combined = any(
        "(94)" in p.text and "Таблица 31" in p.text for p in doc.paragraphs
    )
    scene_mentions = sorted(set(re.findall(r"[\w./\\-]*\.ttt", flat)))
    unexpected_scene_mentions = [
        mention for mention in scene_mentions if not mention.endswith("vkr_scena.ttt")
    ]
    appendix_ok = "Приложение А. Дополнительные материалы по программной реализации ПАК" in flat
    appendix_tables = [
        caption
        for caption in flat.splitlines()
        if caption.startswith("Таблица А.")
    ]

    print(f"docx={DOCX}")
    print(f"zip_ok={bad_member is None}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"table_captions={len(table_captions)}")
    print(f"empty_table_captions={empty_tables}")
    print(f"figure_captions={len(fig_captions)}")
    print(f"duplicate_figures={duplicate_figures}")
    print(f"figures_sequential={not non_sequential_figures}")
    print(f"stale_tokens={stale_tokens}")
    print(f"formula_caption_combined={formula_caption_combined}")
    print(f"scene_mentions={scene_mentions}")
    print(f"unexpected_scene_mentions={unexpected_scene_mentions}")
    print(f"appendix_ok={appendix_ok}")
    print(f"appendix_tables={appendix_tables}")


if __name__ == "__main__":
    main()
