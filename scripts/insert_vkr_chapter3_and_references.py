from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

from compress_vkr_filled_sections import (
    find_index,
    format_body,
    insert_paragraph_after,
    replace_section,
    set_num_pr,
    set_run_font,
    set_update_fields_on_open,
)


BIBLIOGRAPHY_ADDITIONS = [
    "ГОСТ 34.602-89. Информационная технология. Комплекс стандартов на автоматизированные системы. Техническое задание на создание автоматизированной системы. М.: Издательство стандартов, 1990.",
    "Kumar P., Khalid S., Kim H.S. Prognostics and Health Management of Rotating Machinery of Industrial Robot with Deep Learning Applications - A Review // Mathematics. 2023. Vol. 11, No. 13. Article 3008. DOI: 10.3390/math11133008.",
    "Xiao B., Zhong J., Bao X., Chen L., Bao J., Zheng Y. Digital twin-driven prognostics and health management for industrial assets // Scientific Reports. 2024. Vol. 14. Article 13443. DOI: 10.1038/s41598-024-63990-0.",
    "Hu Y., Liu S., Lu H., Zhang H. Remaining Useful Life Model and Assessment of Mechanical Products: A Brief Review and a Note on the State Space Model Method // Chinese Journal of Mechanical Engineering. 2019. Vol. 32. Article 15. DOI: 10.1186/s10033-019-0317-y.",
    "Wang X., Wang T., Ming A., Han Q., Chu F., Zhang W., Li A. Deep Spatiotemporal Convolutional-Neural-Network-Based Remaining Useful Life Estimation of Bearings // Chinese Journal of Mechanical Engineering. 2021. Vol. 34. Article 62. DOI: 10.1186/s10033-021-00576-1.",
    "Tanveer M., Yazdani M.H., Khan R.T.A., Kim H.S. Real-Time AI-Driven Prognostics and Health Management in Robotics // Applied Sciences. 2026. Vol. 16, No. 7. Article 3441. DOI: 10.3390/app16073441.",
    "Wojtulewicz A., Chaber P. Industrial Robot Control System with a Predictive Maintenance Module Using IIoT Technology // Sensors. 2025. Vol. 25, No. 4. Article 1154. DOI: 10.3390/s25041154.",
]


CITATION_APPEND_RULES = [
    (
        "Традиционное планово-предупредительное обслуживание задает ремонтные воздействия",
        " Нормативные термины надежности и подход к заданию требований принимаются по [1, 2].",
    ),
    (
        "Перспективным направлением является предиктивное обслуживание",
        " Обзоры RUL/PHM выделяют модельно-ориентированные, data-driven и гибридные подходы к такой задаче [7-12, 34, 36, 38].",
    ),
    (
        "В качестве технической основы рассматривается промышленный робот ABB IRB 660-180/3.15",
        " Технические параметры и назначение робота подтверждаются эксплуатационными материалами ABB [5, 6].",
    ),
    (
        "Цифровая модель в CoppeliaSim используется как безопасная среда",
        " Использование цифровых моделей и digital twin в PHM производственных активов соответствует [13-16, 35].",
    ),
    (
        "Методическую основу работы составляют системный анализ",
        " Выбор методов согласован с литературой по надежности, PHM, RUL и цифровым двойникам [1-4, 7-17, 34-39].",
    ),
    (
        "Линия розлива рассматривается как последовательная система",
        " Анализ отказов и готовности далее выполняется в терминах надежности по [1-3].",
    ),
    (
        "Повторяемые траектории робота позволяют сравнивать телеметрию одинаковых фаз цикла",
        " Для промышленных роботов такой подход поддерживается исследованиями PHM вращающихся компонентов и IIoT-мониторинга [34, 39].",
    ),
    (
        "Для задач PdM производственная функция дополняется диагностическим контуром",
        " Связка телеметрии, диагностики, прогноза и решения о ТО соответствует общей логике PHM [10-12, 35, 38].",
    ),
    (
        "Для робота-паллетизатора характерны износ зубчатых передач",
        " Для механических узлов и подшипников подобные признаки используются в задачах RUL и построения health indicator [36, 37].",
    ),
    (
        "Для рассматриваемого участка возможны реактивное обслуживание",
        " Сравнение стратегий обслуживания опирается на современные обзоры predictive maintenance и RUL [7-12].",
    ),
    (
        "Цифровая модель не заменяет промышленную диагностику",
        " Применение цифровой модели для диагностики и поддержки решений соответствует подходам digital twin в производстве [13-16, 35].",
    ),
    (
        "PdM выбран для робота-паллетизатора по трем причинам",
        " Для робототехнических систем этот выбор дополнительно подтверждается работами по PHM промышленных роботов и IIoT-мониторингу [34, 38, 39].",
    ),
    (
        "Система строится как программно-аппаратный комплекс с пятью логическими уровнями",
        " Выделение уровней моделирования, сбора, хранения, аналитики и визуализации согласуется с практикой цифровых двойников и систем мониторинга [13-16, 27-32, 35].",
    ),
    (
        "Для каждого окна наблюдения W = [t0, t0 + Δt] рассчитываются статистические",
        " Использование признаков временных рядов для RUL соответствует data-driven подходам к механическим объектам [11, 12, 36, 37].",
    ),
    (
        "Поток данных строится по цепочке: CoppeliaSim -> Remote API",
        " Аналогичная логика обмена процессными данными и визуализации используется в IIoT-решениях для промышленных роботов [18-20, 27-32, 39].",
    ),
    (
        "CoppeliaSim выбран для имитации роботизированной ячейки",
        " Выбор инструментов связан с наличием открытой документации и воспроизводимых API для моделирования, хранения, визуализации и машинного обучения [18-20, 22-32].",
    ),
    (
        "На концептуальном уровне задача прогнозирования задается как регрессия",
        " Такая постановка соответствует распространенной формализации RUL как задачи прогноза по признаковому описанию состояния [7-12, 34, 36, 38].",
    ),
]


def append_bibliography_entries(doc: Document) -> int:
    appendix_idx = find_index(doc, "Приложения")
    bibliography_idx = find_index(doc, "Список литературы")
    existing = "\n".join(p.text for p in doc.paragraphs[bibliography_idx:appendix_idx])
    current = doc.paragraphs[appendix_idx - 1]
    added = 0
    for entry in BIBLIOGRAPHY_ADDITIONS:
        if entry[:80] in existing:
            continue
        current = insert_paragraph_after(current, entry, "Normal")
        set_num_pr(current, "6", "0")
        format_body(current, first_line=False)
        added += 1
    return added


def append_literature_citations(doc: Document) -> int:
    changed = 0
    for needle, suffix in CITATION_APPEND_RULES:
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if needle in text and suffix.strip() not in text:
                run = paragraph.add_run(suffix)
                set_run_font(run, 14)
                changed += 1
                break
    return changed


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: insert_vkr_chapter3_and_references.py <docx> <chapter3.md>")

    docx_path = Path(sys.argv[1])
    chapter3_path = Path(sys.argv[2])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = docx_path.with_name(f"{docx_path.stem}.backup_before_chapter3_tz_{timestamp}{docx_path.suffix}")
    shutil.copy2(docx_path, backup)

    doc = Document(str(docx_path))
    replace_section(doc, "Техническое задание", "Техническое проектирование", chapter3_path, 1, True)
    doc.paragraphs[find_index(doc, "Техническое задание")].paragraph_format.page_break_before = True
    added_bib = append_bibliography_entries(doc)
    added_citations = append_literature_citations(doc)
    doc.save(str(docx_path))
    set_update_fields_on_open(docx_path)

    print(f"Inserted chapter 3 into {docx_path}")
    print(f"Backup: {backup}")
    print(f"Added bibliography entries: {added_bib}")
    print(f"Added citation appends: {added_citations}")


if __name__ == "__main__":
    main()

