from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = PROJECT_ROOT / "вкр"


def wtag(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def mtag(name: str) -> str:
    return f"{{{M_NS}}}{name}"


def find_working_docx() -> Path:
    candidates = [
        p
        for p in DOCX_DIR.glob("*.docx")
        if "2026" in p.name and "backup" not in p.name.lower() and not p.name.startswith("~$")
    ]
    if not candidates:
        raise FileNotFoundError("Working VKR DOCX was not found in the vkr folder")
    return candidates[0]


def para_text(p: etree._Element) -> str:
    return "".join(p.xpath(".//w:t/text()|.//m:t/text()", namespaces=NS)).strip()


def formula_label(p: etree._Element) -> int | None:
    for t in reversed(p.xpath(".//w:t|.//m:t", namespaces=NS)):
        value = t.text or ""
        m = re.fullmatch(r"\((\d{1,3})\)", value.strip())
        if m:
            return int(m.group(1))
    matches = re.findall(r"\((\d{1,3})\)", para_text(p))
    if matches:
        return int(matches[-1])
    return None


def has_math(p: etree._Element) -> bool:
    return bool(p.xpath(".//m:oMath|.//m:oMathPara", namespaces=NS))


def is_formula_paragraph(p: etree._Element) -> bool:
    if not has_math(p):
        return False
    return formula_label(p) is not None


def math_run(text: str) -> etree._Element:
    r = etree.Element(mtag("r"))
    t = etree.SubElement(r, mtag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def math_parts(parent: etree._Element, parts: list[str | etree._Element]) -> None:
    for part in parts:
        if isinstance(part, str):
            parent.append(math_run(part))
        else:
            parent.append(part)


def sub(base: str, index: str) -> etree._Element:
    node = etree.Element(mtag("sSub"))
    e = etree.SubElement(node, mtag("e"))
    e.append(math_run(base))
    sub_node = etree.SubElement(node, mtag("sub"))
    sub_node.append(math_run(index))
    return node


def sub_sup(base: str, lower: str, upper: str) -> etree._Element:
    node = etree.Element(mtag("sSubSup"))
    e = etree.SubElement(node, mtag("e"))
    e.append(math_run(base))
    sub_node = etree.SubElement(node, mtag("sub"))
    sub_node.append(math_run(lower))
    sup_node = etree.SubElement(node, mtag("sup"))
    sup_node.append(math_run(upper))
    return node


def frac(num_parts: list[str | etree._Element], den_parts: list[str | etree._Element]) -> etree._Element:
    node = etree.Element(mtag("f"))
    num = etree.SubElement(node, mtag("num"))
    math_parts(num, num_parts)
    den = etree.SubElement(node, mtag("den"))
    math_parts(den, den_parts)
    return node


def word_run_tab() -> etree._Element:
    r = etree.Element(wtag("r"))
    rpr = etree.SubElement(r, wtag("rPr"))
    sz = etree.SubElement(rpr, wtag("sz"))
    sz.set(f"{{{W_NS}}}val", "28")
    szcs = etree.SubElement(rpr, wtag("szCs"))
    szcs.set(f"{{{W_NS}}}val", "28")
    etree.SubElement(r, wtag("tab"))
    return r


def word_run_text(text: str) -> etree._Element:
    r = etree.Element(wtag("r"))
    rpr = etree.SubElement(r, wtag("rPr"))
    fonts = etree.SubElement(rpr, wtag("rFonts"))
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(f"{{{W_NS}}}{attr}", "Times New Roman")
    sz = etree.SubElement(rpr, wtag("sz"))
    sz.set(f"{{{W_NS}}}val", "28")
    szcs = etree.SubElement(rpr, wtag("szCs"))
    szcs.set(f"{{{W_NS}}}val", "28")
    t = etree.SubElement(r, wtag("t"))
    t.text = text
    return r


def replace_formula_paragraph(p: etree._Element, parts: list[str | etree._Element], label: int) -> None:
    ppr = p.find(wtag("pPr"))
    for child in list(p):
        if child is not ppr:
            p.remove(child)
    p.append(word_run_tab())
    omath = etree.Element(mtag("oMath"))
    math_parts(omath, parts)
    p.append(omath)
    p.append(word_run_tab())
    p.append(word_run_text(f"({label})"))


def formula_parts() -> dict[int, list[str | etree._Element]]:
    return {
        82: [
            sub("N", "пер"),
            " = ",
            sub("N", "сл"),
            " · (1 + ",
            sub("N", "уп,сл"),
            ") = 4 · (1 + 3) = 16,",
        ],
        83: [sub("N", "созд"), " = ", sub("N", "пер"), " + 1 = 17,"],
        84: [sub("T", "набл"), " = ", sub("t", "max"), " - ", sub("t", "min"), " = 2059,05 с,"],
        85: [
            sub("Δt", "ср"),
            " = ",
            frac([sub("T", "набл")], ["n - 1"]),
            " = ",
            frac(["2059,05"], ["22173"]),
            " = 0,0929 с,",
        ],
        86: [sub("f", "набл"), " = ", frac(["1"], [sub("Δt", "ср")]), " = 10,77 Гц,"],
        87: [sub("K", "дан"), " = ", frac([sub("N", "корр")], [sub("N", "общ")]), ","],
        88: [sub("K", "фаз"), " = ", frac([sub("N", "зап,фаз")], [sub("N", "общ")]), ","],
        89: [
            sub("C", "пр"),
            " = ",
            sub("t", "пр"),
            " · ",
            sub("C", "ч"),
            " + ",
            sub("C", "рем"),
            " + ",
            sub("C", "бр"),
            ",",
        ],
        90: [
            sub("E", "год"),
            " = ",
            sub("N", "ав"),
            " · ",
            sub("P", "пред"),
            " · ",
            sub("C", "пр"),
            " - ",
            sub("C", "экспл"),
            ",",
        ],
        91: [sub("T", "ок"), " = ", frac([sub("C", "вн")], [sub("E", "год")]), ","],
        92: [
            sub("Q", "v"),
            " = ",
            sub_sup("Σ", "j=1", "m"),
            " ",
            sub("w", "j"),
            " · ",
            sub("r", "v,j"),
            ",",
        ],
    }


COMPACT_SUBSCRIPTS = {
    "Nзап,phase": ("N", "зап,фаз"),
    "Nуп,сл": ("N", "уп,сл"),
    "Nпер": ("N", "пер"),
    "Nсл": ("N", "сл"),
    "Nсозд": ("N", "созд"),
    "Tнабл": ("T", "набл"),
    "tmax": ("t", "max"),
    "tmin": ("t", "min"),
    "Δtср": ("Δt", "ср"),
    "fнабл": ("f", "набл"),
    "Kдан": ("K", "дан"),
    "Kфаз": ("K", "фаз"),
    "Kпред": ("K", "пред"),
    "Nкорр": ("N", "корр"),
    "Nобщ": ("N", "общ"),
    "Nзап,фаз": ("N", "зап,фаз"),
    "Cпр": ("C", "пр"),
    "tпр": ("t", "пр"),
    "Cч": ("C", "ч"),
    "Cрем": ("C", "рем"),
    "Cбр": ("C", "бр"),
    "Eгод": ("E", "год"),
    "Nав": ("N", "ав"),
    "Pпред": ("P", "пред"),
    "Cэкспл": ("C", "экспл"),
    "Tок": ("T", "ок"),
    "Cвн": ("C", "вн"),
    "Tобн": ("T", "обн"),
    "Qv": ("Q", "v"),
    "wj": ("w", "j"),
    "rv,j": ("r", "v,j"),
    "αi": ("α", "i"),
    "HIi": ("HI", "i"),
    "HIкр": ("HI", "кр"),
}

TOKEN_RE = re.compile(
    r"(?<![\w])("
    + "|".join(re.escape(token) for token in sorted(COMPACT_SUBSCRIPTS, key=len, reverse=True))
    + r")(?![\w])"
)


def set_run_font(run, *, subscript: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.subscript = subscript
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def add_text(paragraph, text: str, *, subscript: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, subscript=subscript)


def rebuild_text_paragraph(paragraph) -> int:
    text = paragraph.text
    if not TOKEN_RE.search(text):
        return 0
    paragraph.clear()
    pos = 0
    conversions = 0
    for match in TOKEN_RE.finditer(text):
        add_text(paragraph, text[pos : match.start()])
        base, index = COMPACT_SUBSCRIPTS[match.group(1)]
        add_text(paragraph, base)
        add_text(paragraph, index, subscript=True)
        conversions += 1
        pos = match.end()
    add_text(paragraph, text[pos:])
    return conversions


def find_paragraph_range(doc: Document) -> tuple[int, int]:
    start = None
    end = len(doc.paragraphs)
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("6.2."):
            start = idx
        elif start is not None and text.startswith("Заключение"):
            end = idx
            break
    if start is None:
        raise RuntimeError("Could not find chapter 6.2 in the DOCX")
    return start, end


def patch_formula_xml(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    body_paragraphs = root.xpath(".//w:body/w:p", namespaces=NS)
    by_label = {formula_label(p): p for p in body_paragraphs if is_formula_paragraph(p)}

    changed = 0
    for label, parts in formula_parts().items():
        if label not in by_label:
            raise RuntimeError(f"Formula ({label}) was not found")
        replace_formula_paragraph(by_label[label], parts, label)
        changed += 1

    files["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    tmp_path = docx_path.with_suffix(".tmp_chapter6_formulas.docx")
    with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    tmp_path.replace(docx_path)
    return changed


def patch_text_notation(docx_path: Path) -> tuple[int, int, int, int]:
    doc = Document(docx_path)
    start, end = find_paragraph_range(doc)
    changed_paragraphs = 0
    conversions = 0
    for idx in range(start, end):
        paragraph = doc.paragraphs[idx]
        if paragraph.style.name.startswith("Heading"):
            continue
        count = rebuild_text_paragraph(paragraph)
        if count:
            changed_paragraphs += 1
            conversions += count
    doc.save(docx_path)
    return start, end, changed_paragraphs, conversions


def main() -> None:
    docx_path = find_working_docx()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = docx_path.with_name(f"{docx_path.stem}.backup_before_chapter6_formula_notation_{timestamp}.docx")
    shutil.copy2(docx_path, backup_path)

    formula_changes = patch_formula_xml(docx_path)
    start, end, changed_paragraphs, conversions = patch_text_notation(docx_path)

    print(f"docx={docx_path}")
    print(f"backup={backup_path}")
    print(f"formula_changes={formula_changes}")
    print(f"paragraph_range={start}:{end}")
    print(f"changed_paragraphs={changed_paragraphs}")
    print(f"text_subscript_conversions={conversions}")


if __name__ == "__main__":
    main()
